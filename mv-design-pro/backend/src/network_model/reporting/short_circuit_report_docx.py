"""
Reporting/export layer - DOCX report generator for ShortCircuitResult.

This module provides functions to export IEC 60909 short-circuit calculation
results to DOCX format. It uses the stable Result API contract
(ShortCircuitResult.to_dict() + white_box_trace) without modifying any solver logic.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from network_model.solvers.short_circuit_iec60909 import ShortCircuitResult

# Check for python-docx availability at import time
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def _format_complex(value: dict | complex | str | Any) -> str:
    """
    Format a complex number for display.

    Handles dict with re/im keys, complex numbers, and strings.
    Returns a readable string like "1.23 + j*4.56" or "1.23 - j*4.56".
    """
    if isinstance(value, dict) and "re" in value and "im" in value:
        re_part = value["re"]
        im_part = value["im"]
    elif isinstance(value, complex):
        re_part = value.real
        im_part = value.imag
    elif isinstance(value, str):
        return value
    else:
        return str(value)

    if im_part >= 0:
        return f"{re_part:.6g} + j*{im_part:.6g}"
    else:
        return f"{re_part:.6g} - j*{abs(im_part):.6g}"


def _format_value(value: Any) -> str:
    """Format a value for display in the report."""
    if value is None:
        return "—"
    if isinstance(value, dict) and "re" in value and "im" in value:
        return _format_complex(value)
    if isinstance(value, complex):
        return _format_complex(value)
    if isinstance(value, float):
        return f"{value:.6g}"
    if isinstance(value, bool):
        return "Tak" if value else "Nie"
    if isinstance(value, list):
        return f"[{len(value)} elementów]"
    return str(value)


def export_short_circuit_result_to_docx(
    result: ShortCircuitResult,
    path: str | Path,
    *,
    title: str | None = None,
    include_white_box: bool = True,
) -> Path:
    """
    Export a single ShortCircuitResult to a DOCX file.

    Creates a human-readable engineering report with:
    - Title page/header with fault parameters
    - Results section with key calculation outputs
    - White Box section with calculation trace (optional)

    Args:
        result: The ShortCircuitResult instance to export.
        path: Target file path (str or Path). Parent directories are created
              if they do not exist.
        title: Custom report title. Defaults to "Raport zwarciowy IEC 60909".
        include_white_box: If True, include the white box trace section
                          showing calculation steps. Defaults to True.

    Returns:
        Path to the written DOCX file.

    Raises:
        ImportError: If python-docx is not installed in the environment.
        ValueError: If result.to_dict() does not return a dict.

    Example:
        >>> from network_model.solvers.short_circuit_iec60909 import (
        ...     ShortCircuitIEC60909Solver
        ... )
        >>> result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
        ...     graph=graph, fault_node_id="B", c_factor=1.0, tk_s=1.0
        ... )
        >>> path = export_short_circuit_result_to_docx(result, "output/report.docx")
    """
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX export but is not installed. "
            "Install it with: pip install python-docx"
        )

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Validate that to_dict() returns a dict
    data = result.to_dict()
    if not isinstance(data, dict):
        raise ValueError(
            f"result.to_dict() must return a dict, got {type(data).__name__}"
        )

    # Create document
    doc = Document()

    # Configure default font (minimal styling)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # 1) Title page / header
    report_title = title if title else "Raport zwarciowy IEC 60909"
    heading = doc.add_heading(report_title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with fault parameters
    subtitle_parts = [
        f"Typ zwarcia: {data.get('short_circuit_type', '—')}",
        f"Węzeł: {data.get('fault_node_id', '—')}",
        f"Un: {_format_value(data.get('un_v'))} V",
        f"c: {_format_value(data.get('c_factor'))}",
        f"tk: {_format_value(data.get('tk_s'))} s",
        f"tb: {_format_value(data.get('tb_s'))} s",
    ]
    subtitle = doc.add_paragraph(" | ".join(subtitle_parts))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # 2) Results section
    doc.add_heading("Wyniki", level=1)

    # Results table (2 columns: Name, Value)
    results_table = doc.add_table(rows=1, cols=2)
    results_table.style = "Table Grid"

    # Header row
    hdr_cells = results_table.rows[0].cells
    hdr_cells[0].text = "Nazwa"
    hdr_cells[1].text = "Wartość"

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Result fields to display
    result_fields = [
        ("Ik'' [A]", "ikss_a"),
        ("Ip [A]", "ip_a"),
        ("Ib [A]", "ib_a"),
        ("Ith [A]", "ith_a"),
        ("Sk [MVA]", "sk_mva"),
        ("κ [-]", "kappa"),
        ("R/X [-]", "rx_ratio"),
        ("Zk [Ω]", "zkk_ohm"),
    ]

    for label, key in result_fields:
        row_cells = results_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = _format_value(data.get(key))

    doc.add_paragraph()  # Spacer

    # 3) White Box section (if requested)
    if include_white_box:
        doc.add_heading("White Box", level=1)

        white_box_trace = data.get("white_box_trace", [])

        if not white_box_trace:
            doc.add_paragraph("Brak śladu obliczeń.")
        else:
            for step in white_box_trace:
                _add_white_box_step(doc, step)

    # Save document
    doc.save(str(output_path))

    return output_path


def _add_white_box_step(doc: Document, step: dict) -> None:
    """
    Add a single white box step to the document.

    Args:
        doc: The Document instance to add content to.
        step: A dict with keys: key, title, formula_latex, inputs,
              substitution, result, notes (optional).
    """
    step_key = step.get("key", "")
    step_title = step.get("title", "")

    # Step header
    header_text = f"{step_key}: {step_title}" if step_title else step_key
    doc.add_heading(header_text, level=2)

    # Formula
    formula = step.get("formula_latex", "")
    if formula:
        p = doc.add_paragraph()
        p.add_run("Wzór: ").bold = True
        p.add_run(formula)

    # Inputs table
    inputs = step.get("inputs", {})
    if inputs and isinstance(inputs, dict):
        doc.add_paragraph().add_run("Dane wejściowe:").bold = True
        inputs_table = doc.add_table(rows=1, cols=2)
        inputs_table.style = "Table Grid"

        hdr = inputs_table.rows[0].cells
        hdr[0].text = "Parametr"
        hdr[1].text = "Wartość"

        for k, v in inputs.items():
            row = inputs_table.add_row().cells
            row[0].text = str(k)
            row[1].text = _format_value(v)

    # Substitution
    substitution = step.get("substitution", "")
    if substitution:
        p = doc.add_paragraph()
        p.add_run("Podstawienie: ").bold = True
        p.add_run(str(substitution))

    # Result table
    result_data = step.get("result", {})
    if result_data and isinstance(result_data, dict):
        doc.add_paragraph().add_run("Wynik:").bold = True
        result_table = doc.add_table(rows=1, cols=2)
        result_table.style = "Table Grid"

        hdr = result_table.rows[0].cells
        hdr[0].text = "Parametr"
        hdr[1].text = "Wartość"

        for k, v in result_data.items():
            row = result_table.add_row().cells
            row[0].text = str(k)
            row[1].text = _format_value(v)

    # Notes (if any)
    notes = step.get("notes")
    if notes:
        p = doc.add_paragraph()
        p.add_run("Uwagi: ").bold = True
        p.add_run(str(notes))

    doc.add_paragraph()  # Spacer between steps
