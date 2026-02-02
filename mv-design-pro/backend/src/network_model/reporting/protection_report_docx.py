"""
FIX-12: Protection Coordination DOCX Report Generator

Creates Word document report with:
- Summary with overall verdict
- Device settings table
- Sensitivity/selectivity/overload check tables
- TCC data reference
- Binary determinism (same input → identical bytes)

CANONICAL ALIGNMENT:
- NOT-A-SOLVER: Only formatting, no physics calculations
- 100% Polish labels
- Deterministic output (uses shared docx_determinism module)
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

# Import shared determinism module
from network_model.reporting.docx_determinism import make_docx_bytes_deterministic

# Check for python-docx availability
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


# =============================================================================
# POLISH LABELS
# =============================================================================

VERDICT_COLORS_RGB = {
    "PASS": RGBColor(22, 163, 74) if _DOCX_AVAILABLE else None,  # green
    "MARGINAL": RGBColor(217, 119, 6) if _DOCX_AVAILABLE else None,  # amber
    "FAIL": RGBColor(220, 38, 38) if _DOCX_AVAILABLE else None,  # red
    "ERROR": RGBColor(107, 114, 128) if _DOCX_AVAILABLE else None,  # gray
}

VERDICT_LABELS_PL = {
    "PASS": "Prawidłowa",
    "MARGINAL": "Margines niski",
    "FAIL": "Nieskoordynowane",
    "ERROR": "Błąd analizy",
}


def _format_value(value: Any) -> str:
    """Format a value for display."""
    if value is None:
        return "—"
    if isinstance(value, float):
        if value == float("inf") or value > 900:
            return "∞"
        return f"{value:.3f}"
    if isinstance(value, bool):
        return "Tak" if value else "Nie"
    return str(value)



def export_protection_coordination_to_docx(
    result: dict[str, Any],
    path: str | Path,
    *,
    title: str | None = None,
    metadata: dict[str, Any] | None = None,
    deterministic: bool = True,
) -> Path:
    """
    Export protection coordination result to DOCX.

    Args:
        result: Coordination analysis result dict
        path: Target file path
        title: Custom report title
        metadata: Optional metadata (project_name, created_at, etc.)
        deterministic: If True, ensure binary reproducibility

    Returns:
        Path to the written DOCX file

    Raises:
        ImportError: If python-docx is not installed
    """
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "DOCX export requires python-docx. Install with: pip install python-docx"
        )

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Create document
    doc = Document()

    # ==========================================================================
    # TITLE
    # ==========================================================================
    report_title = title if title else "Raport koordynacji zabezpieczeń nadprądowych"
    title_para = doc.add_heading(report_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ==========================================================================
    # OVERALL VERDICT
    # ==========================================================================
    overall_verdict = result.get("overall_verdict", "ERROR")
    verdict_pl = VERDICT_LABELS_PL.get(overall_verdict, overall_verdict)
    verdict_color = VERDICT_COLORS_RGB.get(overall_verdict)

    verdict_para = doc.add_paragraph()
    verdict_para.add_run("Wynik analizy: ").bold = True
    verdict_run = verdict_para.add_run(verdict_pl)
    verdict_run.bold = True
    if verdict_color:
        verdict_run.font.color.rgb = verdict_color

    # Metadata
    if metadata:
        meta_para = doc.add_paragraph()
        if metadata.get("project_name"):
            meta_para.add_run(f"Projekt: {metadata['project_name']}  |  ")
        if metadata.get("created_at"):
            meta_para.add_run(f"Data: {metadata['created_at'][:19]}")
        meta_para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    # ==========================================================================
    # SUMMARY
    # ==========================================================================
    doc.add_heading("Podsumowanie", level=1)

    summary = result.get("summary", {})
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"

    # Header
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = "Parametr"
    hdr_cells[1].text = "Wartość"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True

    # Data rows
    summary_data = [
        ("Liczba urządzeń", str(summary.get("total_devices", 0))),
        ("Łączna liczba sprawdzeń", str(summary.get("total_checks", 0))),
        ("Czułość - prawidłowe", str(summary.get("sensitivity", {}).get("pass", 0))),
        ("Czułość - nieprawidłowe", str(summary.get("sensitivity", {}).get("fail", 0))),
        ("Selektywność - prawidłowe", str(summary.get("selectivity", {}).get("pass", 0))),
        ("Selektywność - nieprawidłowe", str(summary.get("selectivity", {}).get("fail", 0))),
        ("Przeciążalność - prawidłowe", str(summary.get("overload", {}).get("pass", 0))),
        ("Przeciążalność - nieprawidłowe", str(summary.get("overload", {}).get("fail", 0))),
    ]

    for label, value in summary_data:
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_paragraph()

    # ==========================================================================
    # DEVICE SETTINGS TABLE
    # ==========================================================================
    doc.add_heading("Tabela urządzeń i nastaw", level=1)

    devices = result.get("devices", [])
    if devices:
        # Sort by device name for deterministic order
        sorted_devices = sorted(devices, key=lambda d: d.get("name", d.get("id", "")))

        dev_table = doc.add_table(rows=1, cols=5)
        dev_table.style = "Table Grid"

        # Header
        dev_headers = ["Nazwa", "Typ", "I_pickup [A]", "TMS", "Krzywa"]
        hdr_cells = dev_table.rows[0].cells
        for i, h in enumerate(dev_headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # Data rows
        for dev in sorted_devices:
            row = dev_table.add_row().cells
            settings = dev.get("settings", {})
            stage_51 = settings.get("stage_51", {})
            curve_settings = stage_51.get("curve_settings", {})

            row[0].text = dev.get("name", "—")[:20]
            row[1].text = dev.get("device_type", "—")[:15]
            row[2].text = _format_value(stage_51.get("pickup_current_a"))
            row[3].text = _format_value(curve_settings.get("time_multiplier"))
            row[4].text = curve_settings.get("variant", "—")
    else:
        doc.add_paragraph("Brak urządzeń", style="No Spacing")

    doc.add_paragraph()

    # ==========================================================================
    # SENSITIVITY CHECKS
    # ==========================================================================
    doc.add_heading("Sprawdzenie czułości (I_min / I_pickup)", level=1)

    sensitivity_checks = result.get("sensitivity_checks", [])
    if sensitivity_checks:
        # Sort by device_id for deterministic order
        sorted_sens_checks = sorted(sensitivity_checks, key=lambda c: c.get("device_id", ""))

        sens_table = doc.add_table(rows=1, cols=5)
        sens_table.style = "Table Grid"

        # Header
        headers = ["Urządzenie", "I_min [A]", "I_pickup [A]", "Margines [%]", "Werdykt"]
        hdr_cells = sens_table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # Data
        for check in sorted_sens_checks:
            row = sens_table.add_row().cells
            row[0].text = check.get("device_id", "—")[:12]
            row[1].text = _format_value(check.get("i_fault_min_a"))
            row[2].text = _format_value(check.get("i_pickup_a"))
            row[3].text = _format_value(check.get("margin_percent"))

            verdict = check.get("verdict", "ERROR")
            verdict_text = VERDICT_LABELS_PL.get(verdict, verdict)
            row[4].text = verdict_text
            verdict_color = VERDICT_COLORS_RGB.get(verdict)
            if verdict_color and row[4].paragraphs[0].runs:
                row[4].paragraphs[0].runs[0].font.color.rgb = verdict_color
    else:
        doc.add_paragraph("Brak danych", style="No Spacing")

    doc.add_paragraph()

    # ==========================================================================
    # SELECTIVITY CHECKS
    # ==========================================================================
    doc.add_heading("Sprawdzenie selektywności czasowej (Δt)", level=1)

    selectivity_checks = result.get("selectivity_checks", [])
    if selectivity_checks:
        # Sort by (downstream_device_id, upstream_device_id) for deterministic order
        sorted_sel_checks = sorted(
            selectivity_checks,
            key=lambda c: (c.get("downstream_device_id", ""), c.get("upstream_device_id", ""))
        )

        sel_table = doc.add_table(rows=1, cols=6)
        sel_table.style = "Table Grid"

        # Header
        headers = ["Podrzędne", "Nadrzędne", "t_pod [s]", "t_nad [s]", "Δt [s]", "Werdykt"]
        hdr_cells = sel_table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # Data
        for check in sorted_sel_checks:
            row = sel_table.add_row().cells
            row[0].text = check.get("downstream_device_id", "—")[:12]
            row[1].text = check.get("upstream_device_id", "—")[:12]
            row[2].text = _format_value(check.get("t_downstream_s"))
            row[3].text = _format_value(check.get("t_upstream_s"))
            row[4].text = _format_value(check.get("margin_s"))

            verdict = check.get("verdict", "ERROR")
            verdict_text = VERDICT_LABELS_PL.get(verdict, verdict)
            row[5].text = verdict_text
            verdict_color = VERDICT_COLORS_RGB.get(verdict)
            if verdict_color and row[5].paragraphs[0].runs:
                row[5].paragraphs[0].runs[0].font.color.rgb = verdict_color
    else:
        doc.add_paragraph("Brak danych (wymaga min. 2 urządzeń)", style="No Spacing")

    doc.add_paragraph()

    # ==========================================================================
    # OVERLOAD CHECKS
    # ==========================================================================
    doc.add_heading("Sprawdzenie przeciążalności (I_pickup / I_rob)", level=1)

    overload_checks = result.get("overload_checks", [])
    if overload_checks:
        # Sort by device_id for deterministic order
        sorted_ovl_checks = sorted(overload_checks, key=lambda c: c.get("device_id", ""))

        ovl_table = doc.add_table(rows=1, cols=5)
        ovl_table.style = "Table Grid"

        # Header
        headers = ["Urządzenie", "I_rob [A]", "I_pickup [A]", "Margines [%]", "Werdykt"]
        hdr_cells = ovl_table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # Data
        for check in sorted_ovl_checks:
            row = ovl_table.add_row().cells
            row[0].text = check.get("device_id", "—")[:12]
            row[1].text = _format_value(check.get("i_operating_a"))
            row[2].text = _format_value(check.get("i_pickup_a"))
            row[3].text = _format_value(check.get("margin_percent"))

            verdict = check.get("verdict", "ERROR")
            verdict_text = VERDICT_LABELS_PL.get(verdict, verdict)
            row[4].text = verdict_text
            verdict_color = VERDICT_COLORS_RGB.get(verdict)
            if verdict_color and row[4].paragraphs[0].runs:
                row[4].paragraphs[0].runs[0].font.color.rgb = verdict_color
    else:
        doc.add_paragraph("Brak danych", style="No Spacing")

    doc.add_paragraph()

    # ==========================================================================
    # TCC CURVES INFO
    # ==========================================================================
    doc.add_heading("Krzywe czasowo-prądowe (TCC)", level=1)

    tcc_curves = result.get("tcc_curves", [])
    if tcc_curves:
        # Sort by device_name for deterministic order
        sorted_tcc_curves = sorted(tcc_curves, key=lambda c: c.get("device_name", ""))

        tcc_table = doc.add_table(rows=1, cols=4)
        tcc_table.style = "Table Grid"

        # Header
        headers = ["Urządzenie", "Typ krzywej", "I_pickup [A]", "TMS"]
        hdr_cells = tcc_table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # Data
        for curve in sorted_tcc_curves:
            row = tcc_table.add_row().cells
            row[0].text = curve.get("device_name", "—")[:20]
            row[1].text = curve.get("curve_type", "—")
            row[2].text = _format_value(curve.get("pickup_current_a"))
            row[3].text = _format_value(curve.get("time_multiplier"))

        doc.add_paragraph()
        note = doc.add_paragraph()
        note.add_run("Wykres TCC dostępny w eksporcie interaktywnym").italic = True
    else:
        doc.add_paragraph("Brak krzywych TCC", style="No Spacing")

    doc.add_paragraph()

    # ==========================================================================
    # FOOTER
    # ==========================================================================
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"Run ID: {result.get('run_id', '—')}").font.size = Pt(8)
    footer.add_run(f"  |  Wygenerowano: {result.get('created_at', '—')[:19]}").font.size = Pt(8)

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    # Make deterministic if requested (use shared module)
    if deterministic:
        docx_bytes = make_docx_bytes_deterministic(docx_bytes)

    # Write to file
    output_path.write_bytes(docx_bytes)
    return output_path
