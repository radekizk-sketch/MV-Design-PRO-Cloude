"""
P20d: Reporting/export layer - DOCX report generator for PowerFlowResult.

This module provides functions to export Power Flow calculation results
to DOCX format. It uses the stable Result API contract
(PowerFlowResultV1.to_dict() + PowerFlowTrace.to_dict()) without modifying
any solver logic.

CANONICAL ALIGNMENT:
- NOT-A-SOLVER: No physics calculations, only formatting
- 100% Polish labels
- UTF-8 encoding
- Binary determinism: same input -> identical bytes (SHA256)
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from network_model.solvers.power_flow_result import PowerFlowResultV1
    from network_model.solvers.power_flow_trace import PowerFlowTrace

# Check for python-docx availability at import time
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from network_model.reporting.docx_determinism import make_docx_deterministic


def _format_value(value: Any) -> str:
    """Format a value for display in the report."""
    if value is None:
        return "—"
    if isinstance(value, float):
        return f"{value:.6g}"
    if isinstance(value, bool):
        return "Tak" if value else "Nie"
    if isinstance(value, list):
        return f"[{len(value)} elementow]"
    return str(value)


def export_power_flow_result_to_docx(
    result: PowerFlowResultV1,
    path: str | Path,
    *,
    trace: PowerFlowTrace | None = None,
    metadata: dict[str, Any] | None = None,
    title: str | None = None,
    include_trace: bool = True,
) -> Path:
    """
    Export a single PowerFlowResultV1 to a DOCX file.

    Creates a human-readable engineering report with:
    - Title page/header with run parameters
    - Summary section with key metrics (convergence, losses, voltage range)
    - Bus results table
    - Branch results table
    - Trace section with Newton-Raphson iterations (optional)

    Args:
        result: The PowerFlowResultV1 instance to export.
        path: Target file path (str or Path). Parent directories are created
              if they do not exist.
        trace: Optional PowerFlowTrace to include as trace section.
        metadata: Optional metadata dict (project_name, case_name, run_id, etc.).
        title: Custom report title. Defaults to "Raport rozplywu mocy".
        include_trace: If True, include the trace section showing NR iterations.
                       Defaults to True.

    Returns:
        Path to the written DOCX file.

    Raises:
        ImportError: If python-docx is not installed in the environment.
        ValueError: If result.to_dict() does not return a dict.
    """
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX export but is not installed. "
            "Install it with: pip install python-docx"
        )

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Validate that to_dict() returns a dict
    data = result.to_dict()
    if not isinstance(data, dict):
        raise ValueError(
            f"result.to_dict() must return a dict, got {type(data).__name__}"
        )

    # Create document
    doc = Document()

    # Configure default font (minimal styling)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # 1) Title page / header
    report_title = title if title else "Raport rozplywu mocy"
    heading = doc.add_heading(report_title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with run parameters
    subtitle_parts = [
        f"Status: {'Zbiezny' if data.get('converged') else 'Niezbiezny'}",
        f"Iteracje: {data.get('iterations_count', '—')}",
        f"Tolerancja: {_format_value(data.get('tolerance_used'))}",
        f"Moc bazowa: {data.get('base_mva', 100)} MVA",
    ]
    subtitle = doc.add_paragraph(" | ".join(subtitle_parts))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata if provided
    if metadata:
        meta_parts = []
        if metadata.get("project_name"):
            meta_parts.append(f"Projekt: {metadata['project_name']}")
        if metadata.get("case_name"):
            meta_parts.append(f"Przypadek: {metadata['case_name']}")
        if metadata.get("run_id"):
            meta_parts.append(f"Run: {metadata['run_id'][:8]}...")
        if metadata.get("created_at"):
            meta_parts.append(f"Data: {metadata['created_at']}")
        if meta_parts:
            meta_p = doc.add_paragraph(" | ".join(meta_parts))
            meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # 2) Summary section (Podsumowanie)
    doc.add_heading("Podsumowanie", level=1)

    summary = data.get("summary", {})
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"

    # Header row
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = "Parametr"
    hdr_cells[1].text = "Wartosc"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Summary fields
    summary_fields = [
        ("Status zbieznosci", "Zbiezny" if data.get("converged") else "Niezbiezny"),
        ("Liczba iteracji", data.get("iterations_count", "—")),
        ("Wezel bilansujacy", data.get("slack_bus_id", "—")),
        ("Calkowite straty P [MW]", _format_value(summary.get("total_losses_p_mw"))),
        ("Calkowite straty Q [Mvar]", _format_value(summary.get("total_losses_q_mvar"))),
        ("Moc czynna slack [MW]", _format_value(summary.get("slack_p_mw"))),
        ("Moc bierna slack [Mvar]", _format_value(summary.get("slack_q_mvar"))),
        ("Min. napiecie [pu]", _format_value(summary.get("min_v_pu"))),
        ("Max. napiecie [pu]", _format_value(summary.get("max_v_pu"))),
    ]

    for label, value in summary_fields:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)

    doc.add_paragraph()  # Spacer

    # 3) Bus results section (Wyniki wezlowe)
    doc.add_heading("Wyniki wezlowe (szyny)", level=1)

    bus_results = data.get("bus_results", [])
    if bus_results:
        bus_table = doc.add_table(rows=1, cols=5)
        bus_table.style = "Table Grid"

        # Header
        bus_hdr = bus_table.rows[0].cells
        bus_hdr[0].text = "ID szyny"
        bus_hdr[1].text = "V [pu]"
        bus_hdr[2].text = "Kat [deg]"
        bus_hdr[3].text = "P_inj [MW]"
        bus_hdr[4].text = "Q_inj [Mvar]"
        for cell in bus_hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows (limit to 50 for readability)
        for bus in bus_results[:50]:
            row = bus_table.add_row().cells
            row[0].text = str(bus.get("bus_id", "—"))[:20]
            row[1].text = _format_value(bus.get("v_pu"))
            row[2].text = _format_value(bus.get("angle_deg"))
            row[3].text = _format_value(bus.get("p_injected_mw"))
            row[4].text = _format_value(bus.get("q_injected_mvar"))

        if len(bus_results) > 50:
            doc.add_paragraph(f"... oraz {len(bus_results) - 50} dodatkowych wezlow (patrz pelny eksport JSON)")
    else:
        doc.add_paragraph("Brak wynikow wezlowych.")

    doc.add_paragraph()  # Spacer

    # 4) Branch results section (Wyniki galeziowe)
    doc.add_heading("Wyniki galeziowe", level=1)

    branch_results = data.get("branch_results", [])
    if branch_results:
        branch_table = doc.add_table(rows=1, cols=5)
        branch_table.style = "Table Grid"

        # Header
        branch_hdr = branch_table.rows[0].cells
        branch_hdr[0].text = "ID galezi"
        branch_hdr[1].text = "P_from [MW]"
        branch_hdr[2].text = "Q_from [Mvar]"
        branch_hdr[3].text = "Straty P [MW]"
        branch_hdr[4].text = "Straty Q [Mvar]"
        for cell in branch_hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows (limit to 50 for readability)
        for branch in branch_results[:50]:
            row = branch_table.add_row().cells
            row[0].text = str(branch.get("branch_id", "—"))[:20]
            row[1].text = _format_value(branch.get("p_from_mw"))
            row[2].text = _format_value(branch.get("q_from_mvar"))
            row[3].text = _format_value(branch.get("losses_p_mw"))
            row[4].text = _format_value(branch.get("losses_q_mvar"))

        if len(branch_results) > 50:
            doc.add_paragraph(f"... oraz {len(branch_results) - 50} dodatkowych galezi (patrz pelny eksport JSON)")
    else:
        doc.add_paragraph("Brak wynikow galeziowych.")

    doc.add_paragraph()  # Spacer

    # 5) Trace section (optional)
    if include_trace and trace is not None:
        doc.add_heading("Slad obliczen (Newton-Raphson)", level=1)

        trace_data = trace.to_dict()

        # Trace metadata
        doc.add_paragraph().add_run("Metadane solvera:").bold = True
        trace_meta = doc.add_table(rows=1, cols=2)
        trace_meta.style = "Table Grid"
        hdr = trace_meta.rows[0].cells
        hdr[0].text = "Parametr"
        hdr[1].text = "Wartosc"

        trace_meta_fields = [
            ("Wersja solvera", trace_data.get("solver_version", "—")),
            ("Hash wejscia", trace_data.get("input_hash", "—")[:16] + "..."),
            ("Metoda startu", trace_data.get("init_method", "—")),
            ("Tolerancja", _format_value(trace_data.get("tolerance"))),
            ("Max iteracji", trace_data.get("max_iterations", "—")),
            ("Wynik koncowy", "Zbiezny" if trace_data.get("converged") else "Niezbiezny"),
            ("Wykonane iteracje", trace_data.get("final_iterations_count", "—")),
        ]
        for label, value in trace_meta_fields:
            row = trace_meta.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()

        # Iteration trace table
        iterations = trace_data.get("iterations", [])
        if iterations:
            doc.add_paragraph().add_run("Iteracje:").bold = True
            iter_table = doc.add_table(rows=1, cols=4)
            iter_table.style = "Table Grid"

            iter_hdr = iter_table.rows[0].cells
            iter_hdr[0].text = "k"
            iter_hdr[1].text = "Norma mismatch"
            iter_hdr[2].text = "Max mismatch [pu]"
            iter_hdr[3].text = "Status"
            for cell in iter_hdr:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for it in iterations:
                row = iter_table.add_row().cells
                row[0].text = str(it.get("k", "—"))
                row[1].text = f"{it.get('norm_mismatch', 0):.2e}"
                row[2].text = f"{it.get('max_mismatch_pu', 0):.2e}"
                row[3].text = it.get("cause_if_failed", "OK") or "OK"

        doc.add_paragraph()
        doc.add_paragraph("Uwaga: Pelne macierze Jacobiego dostepne w eksporcie JSON/trace.")

    # Save document
    doc.save(str(output_path))

    # Normalize for binary determinism (fixed timestamps, sorted ZIP entries)
    make_docx_deterministic(output_path)

    return output_path


def export_power_flow_comparison_to_docx(
    comparison: dict[str, Any],
    path: str | Path,
    *,
    title: str | None = None,
) -> Path:
    """
    Export a PowerFlowComparisonResult to a DOCX file.

    Args:
        comparison: The comparison result dict (from to_dict()).
        path: Target file path (str or Path).
        title: Custom report title. Defaults to "Raport porownania rozplywu mocy".

    Returns:
        Path to the written DOCX file.

    Raises:
        ImportError: If python-docx is not installed.
        ValueError: If comparison is not a dict.
    """
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX export but is not installed. "
            "Install it with: pip install python-docx"
        )

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not isinstance(comparison, dict):
        raise ValueError(
            f"comparison must be a dict, got {type(comparison).__name__}"
        )

    # Create document
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # 1) Title
    report_title = title if title else "Raport porownania rozplywu mocy"
    heading = doc.add_heading(report_title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle_parts = [
        f"Run A: {comparison.get('run_a_id', '—')[:8]}...",
        f"Run B: {comparison.get('run_b_id', '—')[:8]}...",
        f"Data: {comparison.get('created_at', '—')}",
    ]
    subtitle = doc.add_paragraph(" | ".join(subtitle_parts))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 2) Summary
    doc.add_heading("Podsumowanie porownania", level=1)

    summary = comparison.get("summary", {})
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"

    hdr = summary_table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "Wartosc"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    summary_fields = [
        ("Liczba szyn", summary.get("total_buses", "—")),
        ("Liczba galezi", summary.get("total_branches", "—")),
        ("Zbieznosc A", "Tak" if summary.get("converged_a") else "Nie"),
        ("Zbieznosc B", "Tak" if summary.get("converged_b") else "Nie"),
        ("Straty P (A) [MW]", _format_value(summary.get("total_losses_p_mw_a"))),
        ("Straty P (B) [MW]", _format_value(summary.get("total_losses_p_mw_b"))),
        ("Delta strat P [MW]", _format_value(summary.get("delta_total_losses_p_mw"))),
        ("Max delta V [pu]", _format_value(summary.get("max_delta_v_pu"))),
        ("Max delta kat [deg]", _format_value(summary.get("max_delta_angle_deg"))),
        ("Liczba problemow", summary.get("total_issues", "—")),
        ("Krytyczne", summary.get("critical_issues", "—")),
        ("Powazne", summary.get("major_issues", "—")),
        ("Srednie", summary.get("moderate_issues", "—")),
        ("Drobne", summary.get("minor_issues", "—")),
    ]

    for label, value in summary_fields:
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    doc.add_paragraph()

    # 3) Ranking problemow
    doc.add_heading("Ranking problemow", level=1)

    ranking = comparison.get("ranking", [])
    if ranking:
        rank_table = doc.add_table(rows=1, cols=4)
        rank_table.style = "Table Grid"

        rank_hdr = rank_table.rows[0].cells
        rank_hdr[0].text = "Priorytet"
        rank_hdr[1].text = "Kod"
        rank_hdr[2].text = "Element"
        rank_hdr[3].text = "Opis"
        for cell in rank_hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        severity_labels = {5: "Krytyczny", 4: "Powazny", 3: "Sredni", 2: "Drobny", 1: "Informacyjny"}
        for issue in ranking[:30]:  # Limit to 30
            row = rank_table.add_row().cells
            row[0].text = severity_labels.get(issue.get("severity", 1), "?")
            row[1].text = issue.get("issue_code", "—")
            row[2].text = str(issue.get("element_ref", "—"))[:20]
            row[3].text = issue.get("description_pl", "—")

        if len(ranking) > 30:
            doc.add_paragraph(f"... oraz {len(ranking) - 30} dodatkowych problemow")
    else:
        doc.add_paragraph("Brak wykrytych problemow.")

    doc.add_paragraph()

    # 4) Top bus differences
    doc.add_heading("Top roznice napieciowe (szyny)", level=1)

    bus_diffs = comparison.get("bus_diffs", [])
    if bus_diffs:
        # Sort by absolute delta_v_pu descending
        sorted_diffs = sorted(bus_diffs, key=lambda x: abs(x.get("delta_v_pu", 0)), reverse=True)[:20]

        diff_table = doc.add_table(rows=1, cols=5)
        diff_table.style = "Table Grid"

        diff_hdr = diff_table.rows[0].cells
        diff_hdr[0].text = "ID szyny"
        diff_hdr[1].text = "V_A [pu]"
        diff_hdr[2].text = "V_B [pu]"
        diff_hdr[3].text = "Delta V [pu]"
        diff_hdr[4].text = "Delta kat [deg]"
        for cell in diff_hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        for diff in sorted_diffs:
            row = diff_table.add_row().cells
            row[0].text = str(diff.get("bus_id", "—"))[:16]
            row[1].text = _format_value(diff.get("v_pu_a"))
            row[2].text = _format_value(diff.get("v_pu_b"))
            row[3].text = _format_value(diff.get("delta_v_pu"))
            row[4].text = _format_value(diff.get("delta_angle_deg"))
    else:
        doc.add_paragraph("Brak roznic szyn do wyswietlenia.")

    # Save document
    doc.save(str(output_path))

    # Normalize for binary determinism (fixed timestamps, sorted ZIP entries)
    make_docx_deterministic(output_path)

    return output_path
