"""
Reporting/export layer - DOCX report generator for AnalysisRun bundles.

This module exports read-only AnalysisRun report data to DOCX format using the
shared reporting dependencies (python-docx) without invoking solvers.

CANONICAL ALIGNMENT:
- Binary determinism: same input -> identical bytes (SHA256)
"""

from __future__ import annotations

import json
from io import BytesIO
from importlib.util import find_spec
from typing import Any

_DOCX_AVAILABLE = find_spec("docx") is not None

if _DOCX_AVAILABLE:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

from network_model.reporting.docx_determinism import make_docx_bytes_deterministic


def export_analysis_run_to_docx(bundle: dict[str, Any]) -> bytes:
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX export but is not installed. "
            "Install it with: pip install python-docx"
        )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_title(doc, "Raport z uruchomienia AnalysisRun")
    _add_header_section(doc, bundle)
    _add_input_snapshot_section(doc, bundle)
    _add_result_summary_section(doc, bundle)
    _add_white_box_section(doc, bundle)
    _add_results_section(doc, bundle)
    _add_overlay_section(doc, bundle)
    _add_overlay_appendix(doc, bundle)

    output = BytesIO()
    doc.save(output)
    # Normalize for binary determinism (fixed timestamps, sorted ZIP entries)
    return make_docx_bytes_deterministic(output.getvalue())


def _add_title(doc: Document, title: str) -> None:
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _add_header_section(doc: Document, bundle: dict[str, Any]) -> None:
    doc.add_heading("A. Nagłówek", level=1)
    project = bundle.get("project", {})
    operating_case = bundle.get("operating_case", {})
    run = bundle.get("run", {})
    rows = [
        ("Projekt", project.get("name")),
        ("Project ID", project.get("id")),
        ("OperatingCase", operating_case.get("name")),
        ("OperatingCase ID", operating_case.get("id")),
        ("AnalysisRun ID", run.get("id")),
        ("Deterministic ID", run.get("deterministic_id")),
        ("Analysis type", run.get("analysis_type")),
        ("Status", run.get("status")),
        ("Created at", run.get("created_at")),
        ("Started at", run.get("started_at")),
        ("Finished at", run.get("finished_at")),
        ("Duration [s]", _format_value(run.get("duration_seconds"))),
        ("Queue time [s]", _format_value(run.get("queue_seconds"))),
    ]
    _add_key_value_table(doc, rows)
    doc.add_paragraph()


def _add_input_snapshot_section(doc: Document, bundle: dict[str, Any]) -> None:
    doc.add_heading("B. Input snapshot", level=1)
    run = bundle.get("run", {})
    doc.add_paragraph(f"Input hash: {run.get('input_hash') or '—'}")
    snapshot = bundle.get("input_snapshot", {})
    doc.add_paragraph(_pretty_json(snapshot))
    doc.add_paragraph()


def _add_result_summary_section(doc: Document, bundle: dict[str, Any]) -> None:
    doc.add_heading("C. Result summary", level=1)
    summary = bundle.get("result_summary", {})
    rows = []
    for key in sorted(summary.keys()):
        label = _label_for_summary_key(key)
        rows.append((label, _format_value(summary.get(key))))
    if rows:
        _add_key_value_table(doc, rows)
    else:
        doc.add_paragraph("Brak danych.")
    doc.add_paragraph()


def _add_white_box_section(doc: Document, bundle: dict[str, Any]) -> None:
    doc.add_heading("D. White-box trace", level=1)
    steps = bundle.get("white_box_trace", [])
    if not steps:
        doc.add_paragraph("Brak śladu obliczeń.")
        doc.add_paragraph()
        return
    for step in steps:
        key = step.get("key") or ""
        title = step.get("title") or ""
        heading = key if not title else f"{key}: {title}"
        doc.add_heading(heading, level=2)
        if step.get("severity") is not None:
            doc.add_paragraph(f"Severity: {step.get('severity')}")
        if step.get("notes") is not None:
            doc.add_paragraph(f"Notes: {step.get('notes')}")
        if "metrics" in step:
            doc.add_paragraph("Metrics:")
            doc.add_paragraph(_pretty_json(step.get("metrics")), style="No Spacing")
        if "data" in step:
            doc.add_paragraph("Data:")
            doc.add_paragraph(_pretty_json(step.get("data")), style="No Spacing")
        doc.add_paragraph()


def _add_results_section(doc: Document, bundle: dict[str, Any]) -> None:
    doc.add_heading("E. Wyniki", level=1)
    results = bundle.get("results", [])
    if not results:
        doc.add_paragraph("Brak zapisanych wyników.")
        doc.add_paragraph()
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Typ"
    header[1].text = "Utworzono"
    header[2].text = "Rozmiar payloadu"
    header[3].text = "Kluczowe parametry"
    for cell in header:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for result in results:
        row = table.add_row().cells
        row[0].text = str(result.get("result_type") or "—")
        row[1].text = str(result.get("created_at") or "—")
        row[2].text = f"{result.get('payload_size_bytes') or 0} B"
        row[3].text = _format_value(result.get("payload_summary"))
    doc.add_paragraph()


def _add_overlay_section(doc: Document, bundle: dict[str, Any]) -> None:
    doc.add_heading("F. SLD overlay payload (JSON)", level=1)
    overlay = bundle.get("overlay")
    if not overlay:
        doc.add_paragraph("Brak danych SLD overlay.")
        doc.add_paragraph()
        return
    diagram = overlay.get("diagram", {})
    summary = overlay.get("summary", {})
    doc.add_paragraph(
        f"Diagram: {diagram.get('name') or '—'} ({diagram.get('id') or '—'})"
    )
    doc.add_paragraph(
        f"Nodes: {summary.get('node_count', 0)} | Branches: {summary.get('branch_count', 0)}"
    )
    payload = overlay.get("payload", {})
    doc.add_paragraph("Skrócony JSON:")
    doc.add_paragraph(_truncate_json(payload), style="No Spacing")
    doc.add_paragraph()


def _add_overlay_appendix(doc: Document, bundle: dict[str, Any]) -> None:
    overlay = bundle.get("overlay")
    if not overlay:
        return
    doc.add_heading("Aneks: pełny SLD overlay payload", level=1)
    payload = overlay.get("payload", {})
    doc.add_paragraph(_pretty_json(payload), style="No Spacing")


def _add_key_value_table(doc: Document, rows: list[tuple[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Pole"
    header_cells[1].text = "Wartość"
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(label)
        row_cells[1].text = _format_value(value)


def _format_value(value: Any) -> str:
    if value is None:
        return "—"
    if isinstance(value, bool):
        return "Tak" if value else "Nie"
    if isinstance(value, float):
        return f"{value:.6g}"
    if isinstance(value, (dict, list)):
        return _truncate_json(value)
    return str(value)


def _pretty_json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True)


def _truncate_json(value: Any, max_chars: int = 1400) -> str:
    payload = _pretty_json(value)
    if len(payload) <= max_chars:
        return payload
    return f"{payload[:max_chars]}\n... (skrócono)"


def _label_for_summary_key(key: str) -> str:
    if key == "pcc_node_id":
        return "PCC – punkt wspólnego przyłączenia"
    return key
