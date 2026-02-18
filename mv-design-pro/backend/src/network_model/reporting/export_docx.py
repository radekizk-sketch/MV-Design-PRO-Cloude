"""
Reporting/export layer - Professional DOCX report generators for SC and PF results.

This module provides generate_sc_report_docx and generate_pf_report_docx functions
that create comprehensive, human-readable engineering reports with:
- Title page with project info
- Input data tables (network topology, impedances)
- Equations section with LaTeX-rendered formulas (as text in DOCX)
- Calculation trace (white-box steps)
- Results tables
- Source contributions

CANONICAL ALIGNMENT:
- NOT-A-SOLVER: No physics calculations, only formatting
- 100% Polish labels
- UTF-8 encoding
- Binary determinism: same input -> identical bytes (SHA256)
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from network_model.solvers.short_circuit_iec60909 import ShortCircuitResult
    from network_model.solvers.power_flow_result import PowerFlowResultV1
    from network_model.solvers.power_flow_trace import PowerFlowTrace

# Check for python-docx availability at import time
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from network_model.reporting.docx_determinism import make_docx_deterministic


def _format_complex(value: dict | complex | str | Any) -> str:
    """Format a complex number for display."""
    if isinstance(value, dict) and "re" in value and "im" in value:
        re_part = value["re"]
        im_part = value["im"]
    elif isinstance(value, complex):
        re_part = value.real
        im_part = value.imag
    elif isinstance(value, str):
        return value
    else:
        return str(value)

    if im_part >= 0:
        return f"{re_part:.6g} + j*{im_part:.6g}"
    else:
        return f"{re_part:.6g} - j*{abs(im_part):.6g}"


def _format_value(value: Any) -> str:
    """Format a value for display in the report."""
    if value is None:
        return "\u2014"
    if isinstance(value, dict) and "re" in value and "im" in value:
        return _format_complex(value)
    if isinstance(value, complex):
        return _format_complex(value)
    if isinstance(value, float):
        return f"{value:.6g}"
    if isinstance(value, bool):
        return "Tak" if value else "Nie"
    if isinstance(value, list):
        return f"[{len(value)} elementow]"
    return str(value)


def _make_header_bold(table: Any) -> None:
    """Make the header row of a table bold."""
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _add_white_box_step(doc: Any, step: dict) -> None:
    """Add a single white box calculation step to the document."""
    step_key = step.get("key", "")
    step_title = step.get("title", "")

    header_text = f"{step_key}: {step_title}" if step_title else step_key
    doc.add_heading(header_text, level=3)

    # Formula (LaTeX as text)
    formula = step.get("formula_latex", "")
    if formula:
        p = doc.add_paragraph()
        p.add_run("Wzor: ").bold = True
        p.add_run(formula)

    # Inputs table
    inputs = step.get("inputs", {})
    if inputs and isinstance(inputs, dict):
        doc.add_paragraph().add_run("Dane wejsciowe:").bold = True
        inputs_table = doc.add_table(rows=1, cols=2)
        inputs_table.style = "Table Grid"
        hdr = inputs_table.rows[0].cells
        hdr[0].text = "Parametr"
        hdr[1].text = "Wartosc"
        _make_header_bold(inputs_table)
        for k, v in inputs.items():
            row = inputs_table.add_row().cells
            row[0].text = str(k)
            row[1].text = _format_value(v)

    # Substitution
    substitution = step.get("substitution", "")
    if substitution:
        p = doc.add_paragraph()
        p.add_run("Podstawienie: ").bold = True
        p.add_run(str(substitution))

    # Result
    result_data = step.get("result", {})
    if result_data and isinstance(result_data, dict):
        doc.add_paragraph().add_run("Wynik:").bold = True
        result_table = doc.add_table(rows=1, cols=2)
        result_table.style = "Table Grid"
        hdr = result_table.rows[0].cells
        hdr[0].text = "Parametr"
        hdr[1].text = "Wartosc"
        _make_header_bold(result_table)
        for k, v in result_data.items():
            row = result_table.add_row().cells
            row[0].text = str(k)
            row[1].text = _format_value(v)

    # Notes
    notes = step.get("notes")
    if notes:
        p = doc.add_paragraph()
        p.add_run("Uwagi: ").bold = True
        p.add_run(str(notes))

    doc.add_paragraph()  # Spacer


def generate_sc_report_docx(
    result: ShortCircuitResult,
    snapshot: dict[str, Any] | None,
    output_path: str | Path,
    *,
    title: str | None = None,
    project_info: dict[str, str] | None = None,
) -> Path:
    """
    Generate a professional DOCX report for IEC 60909 short-circuit results.

    Creates a comprehensive engineering report with:
    - Title page with project info
    - Input data table (network topology, impedances from snapshot)
    - Equations section with LaTeX-rendered formulas (as text)
    - Calculation trace (white-box steps)
    - Results table (Ik'', ip, Ith, Sk)
    - Source contributions table

    Args:
        result: ShortCircuitResult instance to export.
        snapshot: Network snapshot dict with topology/impedances (optional).
        output_path: Target file path.
        title: Custom report title. Defaults to "Raport zwarciowy IEC 60909".
        project_info: Optional project metadata dict (name, version, author, date).

    Returns:
        Path to the written DOCX file.

    Raises:
        ImportError: If python-docx is not installed.
        ValueError: If result.to_dict() does not return a dict.
    """
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX export but is not installed. "
            "Install it with: pip install python-docx"
        )

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    data = result.to_dict()
    if not isinstance(data, dict):
        raise ValueError(
            f"result.to_dict() must return a dict, got {type(data).__name__}"
        )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # =========================================================================
    # 1) Title page
    # =========================================================================
    report_title = title if title else "Raport zwarciowy IEC 60909"
    heading = doc.add_heading(report_title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project info
    if project_info:
        info_parts = []
        for key in ("name", "version", "author", "date"):
            if project_info.get(key):
                label_map = {
                    "name": "Projekt",
                    "version": "Wersja",
                    "author": "Autor",
                    "date": "Data",
                }
                info_parts.append(f"{label_map.get(key, key)}: {project_info[key]}")
        if info_parts:
            info_p = doc.add_paragraph(" | ".join(info_parts))
            info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fault parameters subtitle
    _dash = "\u2014"
    subtitle_parts = [
        f"Typ zwarcia: {data.get('short_circuit_type', _dash)}",
        f"Wezel: {data.get('fault_node_id', _dash)}",
        f"Un: {_format_value(data.get('un_v'))} V",
        f"c: {_format_value(data.get('c_factor'))}",
        f"tk: {_format_value(data.get('tk_s'))} s",
        f"tb: {_format_value(data.get('tb_s'))} s",
    ]
    subtitle = doc.add_paragraph(" | ".join(subtitle_parts))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # =========================================================================
    # 2) Input data section (from snapshot)
    # =========================================================================
    doc.add_heading("Dane wejsciowe sieci", level=1)

    if snapshot:
        # Network nodes
        nodes = snapshot.get("nodes", [])
        if nodes:
            doc.add_heading("Wezly sieci", level=2)
            node_table = doc.add_table(rows=1, cols=3)
            node_table.style = "Table Grid"
            hdr = node_table.rows[0].cells
            hdr[0].text = "ID"
            hdr[1].text = "Typ"
            hdr[2].text = "Un [kV]"
            _make_header_bold(node_table)
            for node in nodes[:50]:
                row = node_table.add_row().cells
                row[0].text = str(node.get("id", "\u2014"))[:20]
                row[1].text = str(node.get("node_type", "\u2014"))
                row[2].text = _format_value(node.get("voltage_level"))
            doc.add_paragraph()

        # Network branches / impedances
        branches = snapshot.get("branches", [])
        if branches:
            doc.add_heading("Galezi i impedancje", level=2)
            branch_table = doc.add_table(rows=1, cols=5)
            branch_table.style = "Table Grid"
            hdr = branch_table.rows[0].cells
            hdr[0].text = "ID"
            hdr[1].text = "Typ"
            hdr[2].text = "Od"
            hdr[3].text = "Do"
            hdr[4].text = "Impedancja"
            _make_header_bold(branch_table)
            for branch in branches[:50]:
                row = branch_table.add_row().cells
                row[0].text = str(branch.get("id", "\u2014"))[:20]
                row[1].text = str(branch.get("branch_type", "\u2014"))
                row[2].text = str(branch.get("from_node_id", "\u2014"))[:12]
                row[3].text = str(branch.get("to_node_id", "\u2014"))[:12]
                z_info = branch.get("z_ohm") or branch.get("uk_percent", "\u2014")
                row[4].text = _format_value(z_info)
            doc.add_paragraph()
    else:
        doc.add_paragraph("Brak danych migawki sieci (snapshot).")
        doc.add_paragraph()

    # =========================================================================
    # 3) Equations section
    # =========================================================================
    doc.add_heading("Rownania IEC 60909", level=1)

    equations = [
        ("Impedancja zastpcza", "Z_k = f(Z_1, Z_2, Z_0)"),
        ("Prad zwarciowy poczatkowy", "I_k'' = (c * U_n * k_U) / |Z_k|"),
        ("Wspolczynnik udaru", "kappa = 1.02 + 0.98 * exp(-3 * R/X)"),
        ("Prad udarowy", "I_p = kappa * sqrt(2) * I_k''"),
        ("Prad zastepczy cieplny", "I_th = I_k'' * sqrt(t_k)"),
        ("Moc zwarciowa", "S_k = sqrt(3) * U_n * I_k'' / 10^6"),
    ]

    eq_table = doc.add_table(rows=1, cols=2)
    eq_table.style = "Table Grid"
    hdr = eq_table.rows[0].cells
    hdr[0].text = "Nazwa"
    hdr[1].text = "Wzor"
    _make_header_bold(eq_table)
    for name, formula in equations:
        row = eq_table.add_row().cells
        row[0].text = name
        row[1].text = formula
    doc.add_paragraph()

    # =========================================================================
    # 4) Calculation trace (white-box steps)
    # =========================================================================
    doc.add_heading("Slad obliczen (White Box)", level=1)

    white_box_trace = data.get("white_box_trace", [])
    if not white_box_trace:
        doc.add_paragraph("Brak sladu obliczen.")
    else:
        for step in white_box_trace:
            _add_white_box_step(doc, step)

    doc.add_paragraph()

    # =========================================================================
    # 5) Results table (Ik'', ip, Ith, Sk)
    # =========================================================================
    doc.add_heading("Wyniki", level=1)

    results_table = doc.add_table(rows=1, cols=2)
    results_table.style = "Table Grid"
    hdr_cells = results_table.rows[0].cells
    hdr_cells[0].text = "Nazwa"
    hdr_cells[1].text = "Wartosc"
    _make_header_bold(results_table)

    result_fields = [
        ("Ik'' [A]", "ikss_a"),
        ("Ip [A]", "ip_a"),
        ("Ib [A]", "ib_a"),
        ("Ith [A]", "ith_a"),
        ("Sk [MVA]", "sk_mva"),
        ("kappa [-]", "kappa"),
        ("R/X [-]", "rx_ratio"),
        ("Zk [ohm]", "zkk_ohm"),
        ("Ik Thevenin [A]", "ik_thevenin_a"),
        ("Ik falowniki [A]", "ik_inverters_a"),
        ("Ik calkowity [A]", "ik_total_a"),
    ]

    for label, key in result_fields:
        row_cells = results_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = _format_value(data.get(key))

    doc.add_paragraph()

    # =========================================================================
    # 6) Source contributions table
    # =========================================================================
    doc.add_heading("Udzialy zrodel", level=1)

    contributions = data.get("contributions", [])
    if contributions:
        contrib_table = doc.add_table(rows=1, cols=5)
        contrib_table.style = "Table Grid"
        hdr = contrib_table.rows[0].cells
        hdr[0].text = "ID zrodla"
        hdr[1].text = "Nazwa"
        hdr[2].text = "Typ"
        hdr[3].text = "I wklad [A]"
        hdr[4].text = "Udzial [%]"
        _make_header_bold(contrib_table)
        for contrib in contributions:
            row = contrib_table.add_row().cells
            row[0].text = str(contrib.get("source_id", "\u2014"))[:16]
            row[1].text = str(contrib.get("source_name", "\u2014"))[:20]
            row[2].text = str(contrib.get("source_type", "\u2014"))
            row[3].text = _format_value(contrib.get("i_contrib_a"))
            share = contrib.get("share")
            row[4].text = f"{share * 100:.2f}" if isinstance(share, (int, float)) else "\u2014"
    else:
        doc.add_paragraph("Brak danych o udzialach zrodel.")

    doc.add_paragraph()

    # Branch contributions (if present)
    branch_contributions = data.get("branch_contributions")
    if branch_contributions:
        doc.add_heading("Udzialy galeziowe", level=2)
        bc_table = doc.add_table(rows=1, cols=5)
        bc_table.style = "Table Grid"
        hdr = bc_table.rows[0].cells
        hdr[0].text = "Zrodlo"
        hdr[1].text = "Galaz"
        hdr[2].text = "Od"
        hdr[3].text = "Do"
        hdr[4].text = "I wklad [A]"
        _make_header_bold(bc_table)
        for bc in branch_contributions[:50]:
            row = bc_table.add_row().cells
            row[0].text = str(bc.get("source_id", "\u2014"))[:12]
            row[1].text = str(bc.get("branch_id", "\u2014"))[:12]
            row[2].text = str(bc.get("from_node_id", "\u2014"))[:12]
            row[3].text = str(bc.get("to_node_id", "\u2014"))[:12]
            row[4].text = _format_value(bc.get("i_contrib_a"))

    # Save document
    doc.save(str(out))
    make_docx_deterministic(out)

    return out


def generate_pf_report_docx(
    result: PowerFlowResultV1,
    snapshot: dict[str, Any] | None,
    output_path: str | Path,
    *,
    trace: PowerFlowTrace | None = None,
    title: str | None = None,
    project_info: dict[str, str] | None = None,
) -> Path:
    """
    Generate a professional DOCX report for Power Flow results.

    Creates a comprehensive engineering report with:
    - Bus voltages table
    - Branch flows table
    - Losses summary
    - Violations section

    Args:
        result: PowerFlowResultV1 instance to export.
        snapshot: Network snapshot dict (optional).
        output_path: Target file path.
        trace: Optional PowerFlowTrace for iteration details.
        title: Custom report title. Defaults to "Raport rozplywu mocy".
        project_info: Optional project metadata dict (name, version, author, date).

    Returns:
        Path to the written DOCX file.

    Raises:
        ImportError: If python-docx is not installed.
        ValueError: If result.to_dict() does not return a dict.
    """
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX export but is not installed. "
            "Install it with: pip install python-docx"
        )

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    data = result.to_dict()
    if not isinstance(data, dict):
        raise ValueError(
            f"result.to_dict() must return a dict, got {type(data).__name__}"
        )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # =========================================================================
    # 1) Title page
    # =========================================================================
    report_title = title if title else "Raport rozplywu mocy"
    heading = doc.add_heading(report_title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project info
    if project_info:
        info_parts = []
        for key in ("name", "version", "author", "date"):
            if project_info.get(key):
                label_map = {
                    "name": "Projekt",
                    "version": "Wersja",
                    "author": "Autor",
                    "date": "Data",
                }
                info_parts.append(f"{label_map.get(key, key)}: {project_info[key]}")
        if info_parts:
            info_p = doc.add_paragraph(" | ".join(info_parts))
            info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parameters subtitle
    _dash = "\u2014"
    subtitle_parts = [
        f"Status: {'Zbiezny' if data.get('converged') else 'Niezbiezny'}",
        f"Iteracje: {data.get('iterations_count', _dash)}",
        f"Tolerancja: {_format_value(data.get('tolerance_used'))}",
        f"Moc bazowa: {data.get('base_mva', 100)} MVA",
    ]
    subtitle = doc.add_paragraph(" | ".join(subtitle_parts))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # =========================================================================
    # 2) Summary / Losses
    # =========================================================================
    doc.add_heading("Podsumowanie i straty", level=1)

    summary = data.get("summary", {})
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "Wartosc"
    _make_header_bold(summary_table)

    summary_fields = [
        ("Status zbieznosci", "Zbiezny" if data.get("converged") else "Niezbiezny"),
        ("Liczba iteracji", str(data.get("iterations_count", "\u2014"))),
        ("Wezel bilansujacy", str(data.get("slack_bus_id", "\u2014"))),
        ("Calkowite straty P [MW]", _format_value(summary.get("total_losses_p_mw"))),
        ("Calkowite straty Q [Mvar]", _format_value(summary.get("total_losses_q_mvar"))),
        ("Moc czynna slack [MW]", _format_value(summary.get("slack_p_mw"))),
        ("Moc bierna slack [Mvar]", _format_value(summary.get("slack_q_mvar"))),
        ("Min. napiecie [pu]", _format_value(summary.get("min_v_pu"))),
        ("Max. napiecie [pu]", _format_value(summary.get("max_v_pu"))),
    ]

    for label, value in summary_fields:
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = str(value)
    doc.add_paragraph()

    # =========================================================================
    # 3) Bus voltages table
    # =========================================================================
    doc.add_heading("Napiecia wezlowe", level=1)

    bus_results = data.get("bus_results", [])
    if bus_results:
        bus_table = doc.add_table(rows=1, cols=5)
        bus_table.style = "Table Grid"
        hdr = bus_table.rows[0].cells
        hdr[0].text = "ID szyny"
        hdr[1].text = "V [pu]"
        hdr[2].text = "Kat [deg]"
        hdr[3].text = "P_inj [MW]"
        hdr[4].text = "Q_inj [Mvar]"
        _make_header_bold(bus_table)

        for bus in bus_results[:50]:
            row = bus_table.add_row().cells
            row[0].text = str(bus.get("bus_id", "\u2014"))[:20]
            row[1].text = _format_value(bus.get("v_pu"))
            row[2].text = _format_value(bus.get("angle_deg"))
            row[3].text = _format_value(bus.get("p_injected_mw"))
            row[4].text = _format_value(bus.get("q_injected_mvar"))

        if len(bus_results) > 50:
            doc.add_paragraph(
                f"... oraz {len(bus_results) - 50} dodatkowych wezlow"
            )
    else:
        doc.add_paragraph("Brak wynikow wezlowych.")
    doc.add_paragraph()

    # =========================================================================
    # 4) Branch flows table
    # =========================================================================
    doc.add_heading("Przeplywy galeziowe", level=1)

    branch_results = data.get("branch_results", [])
    if branch_results:
        branch_table = doc.add_table(rows=1, cols=5)
        branch_table.style = "Table Grid"
        hdr = branch_table.rows[0].cells
        hdr[0].text = "ID galezi"
        hdr[1].text = "P_from [MW]"
        hdr[2].text = "Q_from [Mvar]"
        hdr[3].text = "Straty P [MW]"
        hdr[4].text = "Straty Q [Mvar]"
        _make_header_bold(branch_table)

        for branch in branch_results[:50]:
            row = branch_table.add_row().cells
            row[0].text = str(branch.get("branch_id", "\u2014"))[:20]
            row[1].text = _format_value(branch.get("p_from_mw"))
            row[2].text = _format_value(branch.get("q_from_mvar"))
            row[3].text = _format_value(branch.get("losses_p_mw"))
            row[4].text = _format_value(branch.get("losses_q_mvar"))

        if len(branch_results) > 50:
            doc.add_paragraph(
                f"... oraz {len(branch_results) - 50} dodatkowych galezi"
            )
    else:
        doc.add_paragraph("Brak wynikow galeziowych.")
    doc.add_paragraph()

    # =========================================================================
    # 5) Violations section
    # =========================================================================
    doc.add_heading("Naruszenia granic", level=1)

    violations_found = False

    # Voltage violations
    v_min_limit = 0.95
    v_max_limit = 1.05
    v_violations = []
    for bus in bus_results:
        v_pu = bus.get("v_pu")
        if v_pu is not None and (v_pu < v_min_limit or v_pu > v_max_limit):
            v_violations.append(bus)

    if v_violations:
        violations_found = True
        doc.add_heading("Naruszenia napieciowe", level=2)
        vv_table = doc.add_table(rows=1, cols=3)
        vv_table.style = "Table Grid"
        hdr = vv_table.rows[0].cells
        hdr[0].text = "Szyna"
        hdr[1].text = "V [pu]"
        hdr[2].text = "Status"
        _make_header_bold(vv_table)
        for bus in v_violations:
            row = vv_table.add_row().cells
            row[0].text = str(bus.get("bus_id", "\u2014"))[:20]
            v_val = bus.get("v_pu", 0.0)
            row[1].text = _format_value(v_val)
            if v_val < v_min_limit:
                row[2].text = "Podnapiecie"
            else:
                row[2].text = "Przepiecie"
        doc.add_paragraph()

    # Convergence violation
    if not data.get("converged", True):
        violations_found = True
        p = doc.add_paragraph()
        p.add_run("UWAGA: Obliczenia NIE zbiegly. ").bold = True
        p.add_run("Wyniki moga byc niedokladne.")
        doc.add_paragraph()

    if not violations_found:
        doc.add_paragraph("Brak naruszen. Wszystkie parametry w granicach dopuszczalnych.")
    doc.add_paragraph()

    # =========================================================================
    # 6) NR Trace (optional)
    # =========================================================================
    if trace is not None:
        doc.add_heading("Slad obliczen (Newton-Raphson)", level=1)
        trace_data = trace.to_dict()

        trace_meta = doc.add_table(rows=1, cols=2)
        trace_meta.style = "Table Grid"
        hdr = trace_meta.rows[0].cells
        hdr[0].text = "Parametr"
        hdr[1].text = "Wartosc"
        _make_header_bold(trace_meta)

        trace_fields = [
            ("Wersja solvera", trace_data.get("solver_version", "\u2014")),
            ("Hash wejscia", str(trace_data.get("input_hash", "\u2014"))[:16] + "..."),
            ("Metoda startu", trace_data.get("init_method", "\u2014")),
            ("Tolerancja", _format_value(trace_data.get("tolerance"))),
            ("Max iteracji", str(trace_data.get("max_iterations", "\u2014"))),
            (
                "Wynik koncowy",
                "Zbiezny" if trace_data.get("converged") else "Niezbiezny",
            ),
            ("Wykonane iteracje", str(trace_data.get("final_iterations_count", "\u2014"))),
        ]
        for label, value in trace_fields:
            row = trace_meta.add_row().cells
            row[0].text = label
            row[1].text = str(value)
        doc.add_paragraph()

        iterations = trace_data.get("iterations", [])
        if iterations:
            doc.add_paragraph().add_run("Iteracje:").bold = True
            iter_table = doc.add_table(rows=1, cols=4)
            iter_table.style = "Table Grid"
            iter_hdr = iter_table.rows[0].cells
            iter_hdr[0].text = "k"
            iter_hdr[1].text = "Norma mismatch"
            iter_hdr[2].text = "Max mismatch [pu]"
            iter_hdr[3].text = "Status"
            _make_header_bold(iter_table)
            for it in iterations:
                row = iter_table.add_row().cells
                row[0].text = str(it.get("k", "\u2014"))
                row[1].text = f"{it.get('norm_mismatch', 0):.2e}"
                row[2].text = f"{it.get('max_mismatch_pu', 0):.2e}"
                row[3].text = it.get("cause_if_failed", "OK") or "OK"

    # Save document
    doc.save(str(out))
    make_docx_deterministic(out)

    return out
