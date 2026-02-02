"""
Reference Patterns API — Wzorce odniesienia

READ-ONLY API for running and viewing reference pattern validations.
Pattern A: Dobór I>> dla linii SN (selektywność, czułość, cieplne, SPZ)

CANONICAL ALIGNMENT:
- NOT-A-SOLVER: Reference patterns are INTERPRETATION layer (no physics)
- WHITE BOX: Full trace of validation steps
- DETERMINISM: Same inputs → identical outputs
- Polish labels (100% PL UI)

NO CODENAMES IN UI.
"""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any
from uuid import uuid4

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel

from application.reference_patterns import (
    PATTERN_ID,
    PATTERN_NAME_PL,
    PATTERN_A_FIXTURES_SUBDIR,
    ReferencePatternResult,
    run_pattern_a,
    load_fixture,
    fixture_to_input,
    get_pattern_a_fixtures_dir,
)

router = APIRouter(
    prefix="/api/reference-patterns",
    tags=["reference-patterns"],
)


# =============================================================================
# Request/Response Models
# =============================================================================


class PatternMetadata(BaseModel):
    """Pattern metadata for listing."""

    pattern_id: str
    name_pl: str
    description_pl: str


class FixtureMetadata(BaseModel):
    """Fixture metadata for listing."""

    fixture_id: str
    filename: str
    description: str | None = None
    expected_verdict: str | None = None
    notes_pl: str | None = None


class PatternRunRequest(BaseModel):
    """Request body for running a pattern."""

    pattern_id: str
    fixture_file: str | None = None
    input_override: dict[str, Any] | None = None


class PatternRunResponse(BaseModel):
    """Response for pattern run."""

    run_id: str
    pattern_id: str
    name_pl: str
    verdict: str
    verdict_description_pl: str
    summary_pl: str
    checks: list[dict[str, Any]]
    trace: list[dict[str, Any]]
    artifacts: dict[str, Any]


class PatternListResponse(BaseModel):
    """Response for pattern list."""

    patterns: list[PatternMetadata]


class FixtureListResponse(BaseModel):
    """Response for fixture list."""

    pattern_id: str
    fixtures: list[FixtureMetadata]


# =============================================================================
# Helper Functions
# =============================================================================


def result_to_response(result: ReferencePatternResult, run_id: str) -> PatternRunResponse:
    """Convert ReferencePatternResult to API response."""
    data = result.to_dict()
    return PatternRunResponse(
        run_id=run_id,
        pattern_id=data["pattern_id"],
        name_pl=data["name_pl"],
        verdict=data["verdict"],
        verdict_description_pl=data["verdict_description_pl"],
        summary_pl=data["summary_pl"],
        checks=data["checks"],
        trace=data["trace"],
        artifacts=data["artifacts"],
    )


def list_pattern_a_fixtures() -> list[FixtureMetadata]:
    """List available fixtures for Pattern A."""
    fixtures_dir = get_pattern_a_fixtures_dir()
    fixtures: list[FixtureMetadata] = []

    if not fixtures_dir.exists():
        return fixtures

    for filepath in sorted(fixtures_dir.glob("*.json")):
        try:
            with open(filepath, encoding="utf-8") as f:
                data = json.load(f)

            fixture_id = filepath.stem
            fixtures.append(
                FixtureMetadata(
                    fixture_id=fixture_id,
                    filename=filepath.name,
                    description=data.get("_description"),
                    expected_verdict=data.get("_expected_verdict"),
                    notes_pl=data.get("_notes_pl"),
                )
            )
        except (json.JSONDecodeError, OSError):
            # Skip invalid files
            continue

    return fixtures


# =============================================================================
# Endpoints
# =============================================================================


@router.get("/patterns", response_model=PatternListResponse)
async def list_patterns() -> PatternListResponse:
    """
    List available reference patterns.

    Currently only Pattern A is available.
    """
    patterns = [
        PatternMetadata(
            pattern_id=PATTERN_ID,
            name_pl=PATTERN_NAME_PL,
            description_pl="Wzorzec walidacji doboru nastaw I>> dla linii SN — "
            "selektywność, czułość, wytrzymałość cieplna, blokada SPZ",
        )
    ]
    return PatternListResponse(patterns=patterns)


@router.get("/patterns/{pattern_id}/fixtures", response_model=FixtureListResponse)
async def list_pattern_fixtures(pattern_id: str) -> FixtureListResponse:
    """
    List available fixtures (test cases) for a pattern.

    Pattern A fixtures: case_A_zgodne, case_B_niezgodne, case_C_graniczne
    """
    if pattern_id != PATTERN_ID:
        raise HTTPException(
            status_code=404,
            detail=f"Wzorzec '{pattern_id}' nie istnieje. Dostępny: {PATTERN_ID}",
        )

    fixtures = list_pattern_a_fixtures()
    return FixtureListResponse(pattern_id=pattern_id, fixtures=fixtures)


@router.post("/run", response_model=PatternRunResponse)
async def run_pattern(request: PatternRunRequest) -> PatternRunResponse:
    """
    Run a reference pattern validation.

    Args:
        request: Pattern run request with pattern_id and optional fixture/input override

    Returns:
        Pattern validation result with verdict, checks, trace, and artifacts
    """
    if request.pattern_id != PATTERN_ID:
        raise HTTPException(
            status_code=400,
            detail=f"Nieznany wzorzec: '{request.pattern_id}'. Dostępny: {PATTERN_ID}",
        )

    run_id = str(uuid4())

    try:
        if request.fixture_file:
            # Run with fixture
            result = run_pattern_a(fixture_file=request.fixture_file)
        elif request.input_override:
            # Run with custom input
            # For now, we require a fixture base — override functionality can be added later
            raise HTTPException(
                status_code=400,
                detail="Uruchomienie z input_override wymaga bazowego fixture_file",
            )
        else:
            raise HTTPException(
                status_code=400,
                detail="Wymagany fixture_file lub input_override",
            )

        return result_to_response(result, run_id)

    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Błąd wykonania wzorca: {str(e)}",
        )


@router.get("/fixtures/{fixture_file}", response_model=PatternRunResponse)
async def run_pattern_with_fixture(fixture_file: str) -> PatternRunResponse:
    """
    Run Pattern A with a specific fixture file (convenience GET endpoint).

    Args:
        fixture_file: Fixture filename (e.g., "case_A_zgodne.json")

    Returns:
        Pattern validation result
    """
    run_id = str(uuid4())

    try:
        result = run_pattern_a(fixture_file=fixture_file)
        return result_to_response(result, run_id)
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Błąd wykonania wzorca: {str(e)}",
        )


# =============================================================================
# Export Endpoints — PDF and DOCX
# =============================================================================


@router.get("/fixtures/{fixture_file}/export/pdf")
async def export_pattern_result_pdf(fixture_file: str):
    """
    Export reference pattern result to PDF.

    Args:
        fixture_file: Fixture filename (e.g., "case_A_zgodne.json")

    Returns:
        PDF file as attachment
    """
    from fastapi.responses import Response
    import io

    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="Eksport PDF wymaga biblioteki reportlab. Zainstaluj: pip install reportlab",
        )

    try:
        result = run_pattern_a(fixture_file=fixture_file)
        data = result.to_dict()
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Błąd wykonania wzorca: {str(e)}",
        )

    # Create PDF in memory
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    page_width, page_height = A4

    left_margin = 25 * mm
    top_margin = page_height - 25 * mm
    y = top_margin
    line_height = 5 * mm

    # Title
    c.setFont("Helvetica-Bold", 16)
    title = f"Raport wzorca odniesienia: {data['name_pl']}"
    c.drawString(left_margin, y, title)
    y -= 10 * mm

    # Verdict banner
    c.setFont("Helvetica-Bold", 14)
    verdict_text = f"Werdykt: {data['verdict']}"
    c.drawString(left_margin, y, verdict_text)
    y -= 6 * mm

    c.setFont("Helvetica", 10)
    c.drawString(left_margin, y, data["verdict_description_pl"])
    y -= 6 * mm
    c.drawString(left_margin, y, data["summary_pl"])
    y -= 10 * mm

    # Checks section
    c.setFont("Helvetica-Bold", 12)
    c.drawString(left_margin, y, "Sprawdzenia")
    y -= 6 * mm

    c.setFont("Helvetica", 9)
    for check in data.get("checks", []):
        status_text = f"[{check['status']}] {check['name_pl']}: {check['description_pl']}"
        # Truncate if too long
        if len(status_text) > 80:
            status_text = status_text[:77] + "..."
        c.drawString(left_margin, y, status_text)
        y -= line_height
        if y < 30 * mm:
            c.showPage()
            y = top_margin

    y -= 5 * mm

    # Artifacts section
    c.setFont("Helvetica-Bold", 12)
    c.drawString(left_margin, y, "Wartosci posrednie")
    y -= 6 * mm

    c.setFont("Helvetica", 9)
    artifacts = data.get("artifacts", {})
    artifact_items = [
        ("Czas zwarcia sumaryczny (tk)", artifacts.get("tk_total_s"), "s"),
        ("Prad znamionowy cieplny (Ithn)", artifacts.get("ithn_a"), "A"),
        ("Prad dopuszczalny cieplnie (Ithdop)", artifacts.get("ithdop_a"), "A"),
        ("I_min okna nastaw (pierwotna)", artifacts.get("window_i_min_primary_a"), "A"),
        ("I_max okna nastaw (pierwotna)", artifacts.get("window_i_max_primary_a"), "A"),
        ("Zalecana nastawa (wtorna)", artifacts.get("recommended_setting_secondary_a"), "A"),
        ("Okno nastaw prawidlowe", "Tak" if artifacts.get("window_valid") else "Nie", ""),
        ("Kryterium limitujace I_min", artifacts.get("limiting_criterion_min"), ""),
        ("Kryterium limitujace I_max", artifacts.get("limiting_criterion_max"), ""),
    ]

    for label, value, unit in artifact_items:
        if value is not None:
            if isinstance(value, float):
                text = f"{label}: {value:.4g} {unit}".strip()
            else:
                text = f"{label}: {value} {unit}".strip()
            c.drawString(left_margin, y, text)
            y -= line_height
            if y < 30 * mm:
                c.showPage()
                y = top_margin

    y -= 5 * mm

    # Trace section (summary)
    trace = data.get("trace", [])
    if trace:
        c.setFont("Helvetica-Bold", 12)
        c.drawString(left_margin, y, f"Slad obliczen ({len(trace)} krokow)")
        y -= 6 * mm

        c.setFont("Helvetica", 8)
        for step in trace[:20]:  # Limit to first 20 steps
            step_text = f"{step['step']}: {step['description_pl']}"
            if len(step_text) > 90:
                step_text = step_text[:87] + "..."
            c.drawString(left_margin, y, step_text)
            y -= line_height * 0.8
            if y < 30 * mm:
                c.showPage()
                y = top_margin

        if len(trace) > 20:
            c.drawString(left_margin, y, f"... oraz {len(trace) - 20} dodatkowych krokow")
            y -= line_height

    # Footer
    y = 15 * mm
    c.setFont("Helvetica", 8)
    c.drawString(left_margin, y, "Wygenerowano przez MV-DESIGN PRO — Wzorce odniesienia")

    c.save()
    buffer.seek(0)

    # Generate filename
    fixture_name = fixture_file.replace(".json", "")
    filename = f"wzorzec_odniesienia_{fixture_name}.pdf"

    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/fixtures/{fixture_file}/export/docx")
async def export_pattern_result_docx(fixture_file: str):
    """
    Export reference pattern result to DOCX.

    Args:
        fixture_file: Fixture filename (e.g., "case_A_zgodne.json")

    Returns:
        DOCX file as attachment
    """
    from fastapi.responses import Response
    import io

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="Eksport DOCX wymaga biblioteki python-docx. Zainstaluj: pip install python-docx",
        )

    try:
        result = run_pattern_a(fixture_file=fixture_file)
        data = result.to_dict()
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Błąd wykonania wzorca: {str(e)}",
        )

    # Create DOCX document
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    heading = doc.add_heading(f"Raport wzorca odniesienia: {data['name_pl']}", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Verdict section
    verdict_para = doc.add_paragraph()
    verdict_run = verdict_para.add_run(f"Werdykt: {data['verdict']}")
    verdict_run.bold = True
    verdict_run.font.size = Pt(14)

    doc.add_paragraph(data["verdict_description_pl"])
    doc.add_paragraph(data["summary_pl"])

    doc.add_paragraph()

    # Checks section
    doc.add_heading("Sprawdzenia", level=1)

    checks_table = doc.add_table(rows=1, cols=3)
    checks_table.style = "Table Grid"

    hdr = checks_table.rows[0].cells
    hdr[0].text = "Kryterium"
    hdr[1].text = "Status"
    hdr[2].text = "Opis"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for check in data.get("checks", []):
        row = checks_table.add_row().cells
        row[0].text = check["name_pl"]
        row[1].text = check["status_pl"]
        row[2].text = check["description_pl"]

    doc.add_paragraph()

    # Artifacts section
    doc.add_heading("Wartosci posrednie", level=1)

    artifacts = data.get("artifacts", {})

    # Time and thermal
    doc.add_heading("Czas zwarcia i wytrzymalosc cieplna", level=2)
    thermal_table = doc.add_table(rows=1, cols=2)
    thermal_table.style = "Table Grid"
    hdr = thermal_table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "Wartosc"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    def add_artifact_row(table, label, value, unit=""):
        if value is not None:
            row = table.add_row().cells
            row[0].text = label
            if isinstance(value, float):
                row[1].text = f"{value:.4g} {unit}".strip()
            else:
                row[1].text = f"{value} {unit}".strip()

    add_artifact_row(thermal_table, "Czas zwarcia sumaryczny (tk)", artifacts.get("tk_total_s"), "s")
    add_artifact_row(thermal_table, "Prad znamionowy cieplny (Ithn)", artifacts.get("ithn_a"), "A")
    add_artifact_row(thermal_table, "Prad dopuszczalny cieplnie (Ithdop)", artifacts.get("ithdop_a"), "A")

    doc.add_paragraph()

    # Window settings
    doc.add_heading("Okno nastaw I>>", level=2)
    window_table = doc.add_table(rows=1, cols=2)
    window_table.style = "Table Grid"
    hdr = window_table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "Wartosc"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    add_artifact_row(window_table, "I_min (selektywnosc, pierwotna)", artifacts.get("window_i_min_primary_a"), "A")
    add_artifact_row(window_table, "I_max (okno, pierwotna)", artifacts.get("window_i_max_primary_a"), "A")
    add_artifact_row(window_table, "Zalecana nastawa (wtorna)", artifacts.get("recommended_setting_secondary_a"), "A")
    add_artifact_row(window_table, "Okno nastaw prawidlowe", "Tak" if artifacts.get("window_valid") else "Nie")
    add_artifact_row(window_table, "Kryterium limitujace I_min", artifacts.get("limiting_criterion_min"))
    add_artifact_row(window_table, "Kryterium limitujace I_max", artifacts.get("limiting_criterion_max"))

    doc.add_paragraph()

    # Trace section
    trace = data.get("trace", [])
    if trace:
        doc.add_heading(f"Slad obliczen ({len(trace)} krokow)", level=1)

        for step in trace[:30]:  # Limit to first 30 steps
            step_para = doc.add_paragraph()
            step_run = step_para.add_run(f"{step['step']}: ")
            step_run.bold = True
            step_para.add_run(step["description_pl"])

        if len(trace) > 30:
            doc.add_paragraph(f"... oraz {len(trace) - 30} dodatkowych krokow")

    doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run("Wygenerowano przez MV-DESIGN PRO — Wzorce odniesienia")
    footer_run.font.size = Pt(9)
    footer_run.italic = True

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Generate filename
    fixture_name = fixture_file.replace(".json", "")
    filename = f"wzorzec_odniesienia_{fixture_name}.docx"

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
