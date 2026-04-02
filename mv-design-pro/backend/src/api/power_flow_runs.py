"""Canonical-only API endpoints for power flow runs."""

from __future__ import annotations

from typing import Any
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Query, status
from pydantic import BaseModel, Field

from api.canonical_run_views import (
    build_power_flow_export_bundle,
    build_power_flow_interpretation,
    build_power_flow_run_header,
    get_power_flow_result as get_canonical_power_flow_result,
    get_power_flow_trace as get_canonical_power_flow_trace,
)
from api.dependencies import get_uow_factory
from application.analysis_run.read_model import canonicalize_json
from enm.canonical_analysis import (
    CanonicalRun,
    create_run as create_canonical_run,
    execute_run as execute_canonical_run,
    get_run as get_canonical_run,
    list_runs_for_project as list_canonical_runs_for_project,
)


router = APIRouter(tags=["power-flow"])


class PowerFlowRunCreateRequest(BaseModel):
    operating_case_id: UUID | None = Field(
        default=None,
        description="ID przypadku operacyjnego. Jeżeli None, użyje Active Case.",
    )
    options: dict[str, Any] | None = Field(
        default=None,
        description="Opcje solvera (tolerance, max_iter, trace_level, etc.)",
    )


class PowerFlowRunResponse(BaseModel):
    id: str
    deterministic_id: str
    project_id: str
    operating_case_id: str
    analysis_type: str
    status: str
    result_status: str
    created_at: str
    started_at: str | None
    finished_at: str | None
    input_hash: str
    converged: bool | None = None
    iterations: int | None = None


class PowerFlowExecuteResponse(BaseModel):
    id: str
    status: str
    converged: bool | None = None
    iterations: int | None = None
    error_message: str | None = None


def _require_canonical_run(run_id: UUID) -> CanonicalRun:
    run = get_canonical_run(run_id)
    if run is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Run {run_id} not found",
        )
    return run


def _resolve_case_id(
    project_id: UUID,
    operating_case_id: UUID | None,
    uow_factory: Any,
) -> str:
    if operating_case_id is not None:
        return str(operating_case_id)
    with uow_factory() as uow:
        active_case = uow.cases.get_active_study_case(project_id)
    if active_case is None:
        raise ValueError(f"Brak aktywnego przypadku dla projektu {project_id}")
    return str(active_case.id)


def _build_export_bundle(run_id: UUID) -> dict[str, Any]:
    return build_power_flow_export_bundle(_require_canonical_run(run_id))


@router.get("/projects/{project_id}/power-flow-runs")
def list_power_flow_runs(
    project_id: UUID,
    status: str | None = Query(
        default=None,
        description="Filtruj po statusie (CREATED, RUNNING, FINISHED, FAILED)",
    ),
) -> dict[str, Any]:
    runs = [
        {
            "id": str(run.id),
            "project_id": run.project_id,
            "operating_case_id": run.case_id,
            "status": run.status,
            "result_status": run.result_status,
            "created_at": run.created_at.isoformat(),
            "finished_at": run.finished_at.isoformat() if run.finished_at else None,
            "input_hash": run.input_hash,
            "converged": ((run.raw_result or {}).get("result_v1") or {}).get("converged"),
            "iterations": ((run.raw_result or {}).get("result_v1") or {}).get("iterations_count"),
        }
        for run in list_canonical_runs_for_project(str(project_id), analysis_type="PF")
        if status is None or run.status == status
    ]
    runs.sort(key=lambda run: run.get("created_at") or "", reverse=True)
    return canonicalize_json({"runs": runs, "total": len(runs)})


@router.post("/projects/{project_id}/power-flow-runs")
def create_power_flow_run(
    project_id: UUID,
    request: PowerFlowRunCreateRequest,
    uow_factory=Depends(get_uow_factory),
) -> dict[str, Any]:
    try:
        options = dict(request.options or {})
        options.setdefault("trace_level", "summary")
        case_id = _resolve_case_id(project_id, request.operating_case_id, uow_factory)
        run = create_canonical_run(
            case_id=case_id,
            project_id=str(project_id),
            analysis_type="PF",
            options=options,
        )
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc

    return canonicalize_json(
        {
            "id": str(run.id),
            "deterministic_id": run.input_hash,
            "project_id": str(project_id),
            "operating_case_id": case_id,
            "analysis_type": run.analysis_type,
            "status": run.status,
            "result_status": run.result_status,
            "created_at": run.created_at.isoformat() if run.created_at else None,
            "input_hash": run.input_hash,
        }
    )


@router.post("/power-flow-runs/{run_id}/execute")
def execute_power_flow_run(run_id: UUID) -> dict[str, Any]:
    _require_canonical_run(run_id)
    run = execute_canonical_run(run_id)
    result_v1 = ((run.raw_result or {}).get("result_v1") or {})
    return canonicalize_json(
        {
            "id": str(run.id),
            "status": run.status,
            "converged": result_v1.get("converged"),
            "iterations": result_v1.get("iterations_count"),
            "error_message": run.error_message,
        }
    )


@router.get("/power-flow-runs/{run_id}")
def get_power_flow_run(run_id: UUID) -> dict[str, Any]:
    try:
        return canonicalize_json(build_power_flow_run_header(_require_canonical_run(run_id)))
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc


@router.get("/power-flow-runs/{run_id}/results")
def get_power_flow_results(run_id: UUID) -> dict[str, Any]:
    try:
        return canonicalize_json(get_canonical_power_flow_result(_require_canonical_run(run_id)))
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc


@router.get("/power-flow-runs/{run_id}/trace")
def get_power_flow_trace(run_id: UUID) -> dict[str, Any]:
    try:
        return canonicalize_json(get_canonical_power_flow_trace(_require_canonical_run(run_id)))
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc


@router.get("/power-flow-runs/{run_id}/export/json")
def export_power_flow_run_json(run_id: UUID):
    from fastapi.responses import Response
    import json

    try:
        bundle = _build_export_bundle(run_id)
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc

    export_payload = {
        "report_type": "power_flow_result",
        "report_version": "1.0.0",
        "metadata": bundle["metadata"],
        "result": bundle["result"],
        "trace_summary": {
            "solver_version": bundle["trace"].get("solver_version"),
            "input_hash": bundle["trace"].get("input_hash"),
            "converged": bundle["trace"].get("converged"),
            "final_iterations_count": bundle["trace"].get("final_iterations_count"),
        },
    }
    json_content = json.dumps(export_payload, indent=2, ensure_ascii=False, sort_keys=True)
    return Response(
        content=json_content,
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="power_flow_run_{run_id}.json"'},
    )


@router.get("/power-flow-runs/{run_id}/export/docx")
def export_power_flow_run_docx(run_id: UUID):
    from fastapi.responses import Response
    import io

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="Eksport DOCX wymaga python-docx. Zainstaluj: pip install python-docx",
        )

    bundle = _build_export_bundle(run_id)
    result = bundle["result"]
    metadata = bundle["metadata"]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading = doc.add_heading("Raport rozplywu mocy", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    status_parts = [
        f"Status: {'Zbiezny' if result.get('converged') else 'Niezbiezny'}",
        f"Iteracje: {result.get('iterations_count', '—')}",
        f"Run: {metadata.get('run_id', '—')[:8]}...",
    ]
    subtitle = doc.add_paragraph(" | ".join(status_parts))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("Podsumowanie", level=1)
    summary = result.get("summary", {})
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "Wartosc"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    def add_row(label: str, value: Any) -> None:
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = str(value) if value is not None else "—"

    add_row("Status zbieznosci", "Zbiezny" if result.get("converged") else "Niezbiezny")
    add_row("Liczba iteracji", result.get("iterations_count"))
    add_row("Wezel bilansujacy", result.get("slack_bus_id"))
    add_row("Calkowite straty P [MW]", f"{summary.get('total_losses_p_mw', 0):.4g}")
    add_row("Calkowite straty Q [Mvar]", f"{summary.get('total_losses_q_mvar', 0):.4g}")
    add_row("Min. napiecie [pu]", f"{summary.get('min_v_pu', 0):.4g}")
    add_row("Max. napiecie [pu]", f"{summary.get('max_v_pu', 0):.4g}")

    doc.add_paragraph()
    doc.add_heading("Wyniki wezlowe (szyny)", level=1)
    bus_results = result.get("bus_results", [])
    if bus_results:
        bus_table = doc.add_table(rows=1, cols=5)
        bus_table.style = "Table Grid"
        header = bus_table.rows[0].cells
        header[0].text = "ID szyny"
        header[1].text = "V [pu]"
        header[2].text = "Kat [deg]"
        header[3].text = "P_inj [MW]"
        header[4].text = "Q_inj [Mvar]"
        for cell in header:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for bus in bus_results[:30]:
            row = bus_table.add_row().cells
            row[0].text = str(bus.get("bus_id", "—"))[:16]
            row[1].text = f"{bus.get('v_pu', 0):.4g}"
            row[2].text = f"{bus.get('angle_deg', 0):.2f}"
            row[3].text = f"{bus.get('p_injected_mw', 0):.3g}"
            row[4].text = f"{bus.get('q_injected_mvar', 0):.3g}"
        if len(bus_results) > 30:
            doc.add_paragraph(f"... oraz {len(bus_results) - 30} dodatkowych wezlow")
    else:
        doc.add_paragraph("Brak wynikow wezlowych.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="power_flow_run_{run_id}.docx"'},
    )


@router.get("/power-flow-runs/{run_id}/export/pdf")
def export_power_flow_run_pdf(run_id: UUID):
    from fastapi.responses import Response
    import io

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="Eksport PDF wymaga reportlab. Zainstaluj: pip install reportlab",
        )

    bundle = _build_export_bundle(run_id)
    result = bundle["result"]
    metadata = bundle["metadata"]

    buffer = io.BytesIO()
    canvas_obj = canvas.Canvas(buffer, pagesize=A4)
    page_width, page_height = A4
    left_margin = 25 * mm
    top_margin = page_height - 25 * mm
    y = top_margin
    line_height = 5 * mm

    canvas_obj.setFont("Helvetica-Bold", 16)
    title = "Raport rozplywu mocy"
    canvas_obj.drawString(
        (page_width - canvas_obj.stringWidth(title, "Helvetica-Bold", 16)) / 2,
        y,
        title,
    )
    y -= 10 * mm

    canvas_obj.setFont("Helvetica", 10)
    status_text = (
        f"Status: {'Zbiezny' if result.get('converged') else 'Niezbiezny'} | "
        f"Iteracje: {result.get('iterations_count', '—')} | "
        f"Run: {metadata.get('run_id', '—')[:8]}..."
    )
    canvas_obj.drawString(left_margin, y, status_text)
    y -= 8 * mm

    canvas_obj.setFont("Helvetica-Bold", 14)
    canvas_obj.drawString(left_margin, y, "Podsumowanie")
    y -= 6 * mm

    canvas_obj.setFont("Helvetica", 10)
    summary = result.get("summary", {})
    summary_lines = [
        f"Wezel bilansujacy: {result.get('slack_bus_id', '—')}",
        f"Calkowite straty P: {summary.get('total_losses_p_mw', 0):.4g} MW",
        f"Calkowite straty Q: {summary.get('total_losses_q_mvar', 0):.4g} Mvar",
        f"Min. napiecie: {summary.get('min_v_pu', 0):.4g} pu",
        f"Max. napiecie: {summary.get('max_v_pu', 0):.4g} pu",
    ]
    for line in summary_lines:
        canvas_obj.drawString(left_margin, y, line)
        y -= line_height

    y -= 5 * mm
    canvas_obj.setFont("Helvetica-Bold", 12)
    canvas_obj.drawString(left_margin, y, "Wyniki wezlowe (top 20)")
    y -= 5 * mm

    canvas_obj.setFont("Helvetica", 9)
    for bus in result.get("bus_results", [])[:20]:
        text = (
            f"{str(bus.get('bus_id', '—'))[:12]}: "
            f"V={bus.get('v_pu', 0):.4g} pu, "
            f"kat={bus.get('angle_deg', 0):.2f} deg"
        )
        canvas_obj.drawString(left_margin, y, text)
        y -= line_height
        if y < 30 * mm:
            canvas_obj.showPage()
            y = top_margin

    canvas_obj.save()
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="power_flow_run_{run_id}.pdf"'},
    )


@router.get("/power-flow-runs/{run_id}/export/proof/json")
def export_power_flow_proof_json(run_id: UUID) -> dict[str, Any]:
    _ = run_id
    raise HTTPException(
        status_code=status.HTTP_410_GONE,
        detail=(
            "Eksport proof dla power-flow został wycofany z toru produkcyjnego do czasu "
            "pełnego przepięcia na kanoniczny White Box."
        ),
    )


@router.get("/power-flow-runs/{run_id}/export/proof/latex")
def export_power_flow_proof_latex(run_id: UUID) -> dict[str, Any]:
    _ = run_id
    raise HTTPException(
        status_code=status.HTTP_410_GONE,
        detail=(
            "Eksport proof dla power-flow został wycofany z toru produkcyjnego do czasu "
            "pełnego przepięcia na kanoniczny White Box."
        ),
    )


@router.get("/power-flow-runs/{run_id}/export/proof/pdf")
def export_power_flow_proof_pdf(run_id: UUID) -> dict[str, Any]:
    _ = run_id
    raise HTTPException(
        status_code=status.HTTP_410_GONE,
        detail=(
            "Eksport proof dla power-flow został wycofany z toru produkcyjnego do czasu "
            "pełnego przepięcia na kanoniczny White Box."
        ),
    )


_interpretation_cache: dict[str, dict[str, Any]] = {}


@router.get("/power-flow-runs/{run_id}/interpretation")
def get_power_flow_interpretation(run_id: UUID) -> dict[str, Any]:
    run_id_str = str(run_id)
    if run_id_str in _interpretation_cache:
        return canonicalize_json(_interpretation_cache[run_id_str])

    try:
        result_dict = build_power_flow_interpretation(_require_canonical_run(run_id))
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc
    _interpretation_cache[run_id_str] = result_dict
    return canonicalize_json(result_dict)


@router.post("/power-flow-runs/{run_id}/interpretation")
def create_power_flow_interpretation(run_id: UUID) -> dict[str, Any]:
    return get_power_flow_interpretation(run_id)
