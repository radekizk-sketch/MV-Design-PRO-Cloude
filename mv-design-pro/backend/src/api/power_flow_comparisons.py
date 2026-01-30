"""
Power Flow Comparison API — P20c/P20d (A/B)

REST API endpoints for comparing two power flow analysis runs.
100% read-only — no physics calculations, no state mutations.

Endpoints:
- POST /power-flow-comparisons — Create/execute comparison
- GET /power-flow-comparisons/{id} — Get comparison metadata
- GET /power-flow-comparisons/{id}/results — Get comparison results
- GET /power-flow-comparisons/{id}/trace — Get comparison trace
- GET /power-flow-comparisons/{id}/export/json — Export to JSON (P20d)
- GET /power-flow-comparisons/{id}/export/docx — Export to DOCX (P20d)
- GET /power-flow-comparisons/{id}/export/pdf — Export to PDF (P20d)

CANONICAL ALIGNMENT:
- P20c: Power Flow A/B Comparison
- P20d: Export endpoints
- Read-only comparison endpoint
- Deterministic response
"""

from __future__ import annotations

from typing import Any

from fastapi import APIRouter, Depends, HTTPException, status
from pydantic import BaseModel, Field

from api.dependencies import get_uow_factory
from application.power_flow_comparison import PowerFlowComparisonService
from domain.power_flow_comparison import (
    PowerFlowComparisonError,
    PowerFlowComparisonNotFoundError,
    PowerFlowProjectMismatchError,
    PowerFlowResultNotFoundError,
    PowerFlowRunNotFinishedError,
    PowerFlowRunNotFoundError,
)


router = APIRouter(prefix="/power-flow-comparisons", tags=["power-flow-comparison"])


# =============================================================================
# REQUEST/RESPONSE MODELS
# =============================================================================


class CreatePowerFlowComparisonRequest(BaseModel):
    """Request to create a power flow comparison."""
    power_flow_run_id_a: str = Field(
        ...,
        description="UUID pierwszego PowerFlowRun (baseline)",
    )
    power_flow_run_id_b: str = Field(
        ...,
        description="UUID drugiego PowerFlowRun (porownanie)",
    )


class BusDiffRowResponse(BaseModel):
    """Single bus diff row."""
    bus_id: str
    v_pu_a: float
    v_pu_b: float
    angle_deg_a: float
    angle_deg_b: float
    p_injected_mw_a: float
    p_injected_mw_b: float
    q_injected_mvar_a: float
    q_injected_mvar_b: float
    delta_v_pu: float
    delta_angle_deg: float
    delta_p_mw: float
    delta_q_mvar: float


class BranchDiffRowResponse(BaseModel):
    """Single branch diff row."""
    branch_id: str
    p_from_mw_a: float
    p_from_mw_b: float
    q_from_mvar_a: float
    q_from_mvar_b: float
    p_to_mw_a: float
    p_to_mw_b: float
    q_to_mvar_a: float
    q_to_mvar_b: float
    losses_p_mw_a: float
    losses_p_mw_b: float
    losses_q_mvar_a: float
    losses_q_mvar_b: float
    delta_p_from_mw: float
    delta_q_from_mvar: float
    delta_p_to_mw: float
    delta_q_to_mvar: float
    delta_losses_p_mw: float
    delta_losses_q_mvar: float


class RankingIssueResponse(BaseModel):
    """Single ranking issue."""
    issue_code: str
    severity: int
    element_ref: str
    description_pl: str
    evidence_ref: int


class ComparisonSummaryResponse(BaseModel):
    """Comparison summary statistics."""
    total_buses: int
    total_branches: int
    converged_a: bool
    converged_b: bool
    total_losses_p_mw_a: float
    total_losses_p_mw_b: float
    delta_total_losses_p_mw: float
    max_delta_v_pu: float
    max_delta_angle_deg: float
    total_issues: int
    critical_issues: int
    major_issues: int
    moderate_issues: int
    minor_issues: int


class PowerFlowComparisonResultResponse(BaseModel):
    """Full comparison result response."""
    comparison_id: str
    run_a_id: str
    run_b_id: str
    project_id: str
    bus_diffs: list[BusDiffRowResponse]
    branch_diffs: list[BranchDiffRowResponse]
    ranking: list[RankingIssueResponse]
    summary: ComparisonSummaryResponse
    input_hash: str
    created_at: str


class TraceStepResponse(BaseModel):
    """Single trace step."""
    step: str
    description_pl: str
    inputs: dict[str, Any]
    outputs: dict[str, Any]


class PowerFlowComparisonTraceResponse(BaseModel):
    """Full comparison trace response."""
    comparison_id: str
    run_a_id: str
    run_b_id: str
    snapshot_id_a: str | None
    snapshot_id_b: str | None
    input_hash_a: str
    input_hash_b: str
    solver_version: str
    ranking_thresholds: dict[str, float]
    steps: list[TraceStepResponse]
    created_at: str


class PowerFlowComparisonMetadataResponse(BaseModel):
    """Comparison metadata (without full results)."""
    comparison_id: str
    run_a_id: str
    run_b_id: str
    project_id: str
    summary: ComparisonSummaryResponse
    input_hash: str
    created_at: str


# =============================================================================
# SERVICE FACTORY
# =============================================================================


def _build_service(uow_factory: Any) -> PowerFlowComparisonService:
    return PowerFlowComparisonService(uow_factory)


# =============================================================================
# ENDPOINTS
# =============================================================================


@router.post(
    "",
    status_code=status.HTTP_201_CREATED,
    response_model=PowerFlowComparisonResultResponse,
    summary="Utworz porownanie dwoch analiz rozplywu mocy",
    description="""
P20c: Porownuje dwa PowerFlowRun i generuje deterministyczny ranking problemow.

**Walidacje:**
- Oba runy musza istniec
- Oba runy musza miec status FINISHED
- Oba runy musza nalezec do tego samego projektu

**Zwraca:**
- PowerFlowComparisonResult z:
  - bus_diffs: porownanie per szyna (posortowane po bus_id)
  - branch_diffs: porownanie per galaz (posortowane po branch_id)
  - ranking: lista problemow posortowana wg severity (5->1)
  - summary: statystyki porownania

**Cache:**
- Ta sama para (A, B) -> ten sam comparison_id
- A->B != B->A (kierunkowe)
""",
)
def create_power_flow_comparison(
    request: CreatePowerFlowComparisonRequest,
    uow_factory=Depends(get_uow_factory),
) -> dict[str, Any]:
    """
    Create a power flow comparison between two runs.

    P20c: Main comparison endpoint.

    INVARIANTS:
    - Read-only: Zero physics calculations, zero state mutations
    - Same Project: Both runs must belong to the same project
    - Finished Only: Both runs must be FINISHED
    - Deterministic: Same inputs produce identical comparison output
    """
    service = _build_service(uow_factory)

    try:
        result = service.compare(
            run_a_id=request.power_flow_run_id_a,
            run_b_id=request.power_flow_run_id_b,
        )
        return result.to_dict()

    except PowerFlowRunNotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Power flow run nie znaleziony: {e.run_id}",
        ) from e
    except PowerFlowRunNotFinishedError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Power flow run nie zakonczony (status: {e.status}): {e.run_id}",
        ) from e
    except PowerFlowProjectMismatchError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Runs naleza do roznych projektow: {e.run_a_project} vs {e.run_b_project}",
        ) from e
    except PowerFlowResultNotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Wyniki power flow nie znalezione dla run: {e.run_id}",
        ) from e
    except PowerFlowComparisonError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        ) from e


@router.get(
    "/{comparison_id}",
    response_model=PowerFlowComparisonMetadataResponse,
    summary="Pobierz metadane porownania",
)
def get_power_flow_comparison(
    comparison_id: str,
    uow_factory=Depends(get_uow_factory),
) -> dict[str, Any]:
    """
    Get power flow comparison metadata.

    Returns comparison ID, run IDs, summary statistics, and timestamps.
    Does not include full bus_diffs/branch_diffs/ranking (use /results endpoint for that).
    """
    service = _build_service(uow_factory)

    try:
        result = service.get_comparison(comparison_id)
        return {
            "comparison_id": result.comparison_id,
            "run_a_id": result.run_a_id,
            "run_b_id": result.run_b_id,
            "project_id": result.project_id,
            "summary": result.summary.to_dict(),
            "input_hash": result.input_hash,
            "created_at": result.created_at.isoformat(),
        }

    except PowerFlowComparisonNotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Power flow comparison nie znalezione: {e.comparison_id}",
        ) from e


@router.get(
    "/{comparison_id}/results",
    response_model=PowerFlowComparisonResultResponse,
    summary="Pobierz pelne wyniki porownania",
)
def get_power_flow_comparison_results(
    comparison_id: str,
    uow_factory=Depends(get_uow_factory),
) -> dict[str, Any]:
    """
    Get full power flow comparison results.

    Returns all bus_diffs, branch_diffs, ranking issues, and summary.
    """
    service = _build_service(uow_factory)

    try:
        result = service.get_comparison(comparison_id)
        return result.to_dict()

    except PowerFlowComparisonNotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Power flow comparison nie znalezione: {e.comparison_id}",
        ) from e


@router.get(
    "/{comparison_id}/trace",
    response_model=PowerFlowComparisonTraceResponse,
    summary="Pobierz slad porownania (audyt)",
)
def get_power_flow_comparison_trace(
    comparison_id: str,
    uow_factory=Depends(get_uow_factory),
) -> dict[str, Any]:
    """
    Get power flow comparison trace for audit.

    Returns all comparison steps with inputs/outputs for reproducibility.
    Includes explicit ranking thresholds used.
    """
    service = _build_service(uow_factory)

    try:
        trace = service.get_comparison_trace(comparison_id)
        return trace.to_dict()

    except PowerFlowComparisonNotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Power flow comparison nie znalezione: {e.comparison_id}",
        ) from e


# =============================================================================
# P20d: Export Endpoints
# =============================================================================


@router.get(
    "/{comparison_id}/export/json",
    summary="Eksportuj porownanie do JSON",
)
def export_power_flow_comparison_json(
    comparison_id: str,
    uow_factory=Depends(get_uow_factory),
):
    """P20d: Export power flow comparison to JSON file."""
    from fastapi.responses import Response
    import json

    service = _build_service(uow_factory)

    try:
        result = service.get_comparison(comparison_id)
        comparison_dict = result.to_dict()
    except PowerFlowComparisonNotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Power flow comparison nie znalezione: {e.comparison_id}",
        ) from e

    # Build export payload
    export_payload = {
        "report_type": "power_flow_comparison",
        "report_version": "1.0.0",
        "comparison": comparison_dict,
    }

    json_content = json.dumps(export_payload, indent=2, ensure_ascii=False, sort_keys=True)

    return Response(
        content=json_content,
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="power_flow_comparison_{comparison_id}.json"'},
    )


@router.get(
    "/{comparison_id}/export/docx",
    summary="Eksportuj porownanie do DOCX",
)
def export_power_flow_comparison_docx(
    comparison_id: str,
    uow_factory=Depends(get_uow_factory),
):
    """P20d: Export power flow comparison to DOCX file."""
    from fastapi.responses import Response
    import io

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="DOCX export requires python-docx. Install with: pip install python-docx",
        )

    service = _build_service(uow_factory)

    try:
        result = service.get_comparison(comparison_id)
        comparison = result.to_dict()
    except PowerFlowComparisonNotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Power flow comparison nie znalezione: {e.comparison_id}",
        ) from e

    # Create DOCX document
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    heading = doc.add_heading("Raport porownania rozplywu mocy", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle_parts = [
        f"Run A: {comparison.get('run_a_id', '—')[:8]}...",
        f"Run B: {comparison.get('run_b_id', '—')[:8]}...",
    ]
    subtitle = doc.add_paragraph(" | ".join(subtitle_parts))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Summary section
    doc.add_heading("Podsumowanie", level=1)
    summary = comparison.get("summary", {})
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"

    hdr = summary_table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "Wartosc"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    def add_row(table, label, value):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value) if value is not None else "—"

    add_row(summary_table, "Liczba szyn", summary.get("total_buses"))
    add_row(summary_table, "Liczba galezi", summary.get("total_branches"))
    add_row(summary_table, "Zbieznosc A", "Tak" if summary.get("converged_a") else "Nie")
    add_row(summary_table, "Zbieznosc B", "Tak" if summary.get("converged_b") else "Nie")
    add_row(summary_table, "Delta strat P [MW]", f"{summary.get('delta_total_losses_p_mw', 0):.4g}")
    add_row(summary_table, "Max delta V [pu]", f"{summary.get('max_delta_v_pu', 0):.4g}")
    add_row(summary_table, "Liczba problemow", summary.get("total_issues"))
    add_row(summary_table, "Krytyczne", summary.get("critical_issues"))
    add_row(summary_table, "Powazne", summary.get("major_issues"))

    doc.add_paragraph()

    # Ranking section
    doc.add_heading("Ranking problemow", level=1)
    ranking = comparison.get("ranking", [])
    severity_labels = {5: "Krytyczny", 4: "Powazny", 3: "Sredni", 2: "Drobny", 1: "Info"}

    if ranking:
        rank_table = doc.add_table(rows=1, cols=4)
        rank_table.style = "Table Grid"
        rhdr = rank_table.rows[0].cells
        rhdr[0].text = "Priorytet"
        rhdr[1].text = "Kod"
        rhdr[2].text = "Element"
        rhdr[3].text = "Opis"
        for cell in rhdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        for issue in ranking[:30]:
            row = rank_table.add_row().cells
            row[0].text = severity_labels.get(issue.get("severity", 1), "?")
            row[1].text = issue.get("issue_code", "—")
            row[2].text = str(issue.get("element_ref", "—"))[:16]
            row[3].text = issue.get("description_pl", "—")[:50]

        if len(ranking) > 30:
            doc.add_paragraph(f"... oraz {len(ranking) - 30} dodatkowych problemow")
    else:
        doc.add_paragraph("Brak wykrytych problemow.")

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="power_flow_comparison_{comparison_id}.docx"'},
    )


@router.get(
    "/{comparison_id}/export/pdf",
    summary="Eksportuj porownanie do PDF",
)
def export_power_flow_comparison_pdf(
    comparison_id: str,
    uow_factory=Depends(get_uow_factory),
):
    """P20d: Export power flow comparison to PDF file."""
    from fastapi.responses import Response
    import io

    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="PDF export requires reportlab. Install with: pip install reportlab",
        )

    service = _build_service(uow_factory)

    try:
        result = service.get_comparison(comparison_id)
        comparison = result.to_dict()
    except PowerFlowComparisonNotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Power flow comparison nie znalezione: {e.comparison_id}",
        ) from e

    # Create PDF in memory
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    page_width, page_height = A4

    left_margin = 25 * mm
    top_margin = page_height - 25 * mm
    y = top_margin
    line_height = 5 * mm

    # Title
    c.setFont("Helvetica-Bold", 16)
    title = "Raport porownania rozplywu mocy"
    c.drawString((page_width - c.stringWidth(title, "Helvetica-Bold", 16)) / 2, y, title)
    y -= 10 * mm

    # Subtitle
    c.setFont("Helvetica", 10)
    subtitle = f"Run A: {comparison.get('run_a_id', '—')[:8]}... | Run B: {comparison.get('run_b_id', '—')[:8]}..."
    c.drawString(left_margin, y, subtitle)
    y -= 8 * mm

    # Summary
    c.setFont("Helvetica-Bold", 14)
    c.drawString(left_margin, y, "Podsumowanie")
    y -= 6 * mm

    c.setFont("Helvetica", 10)
    summary = comparison.get("summary", {})
    summary_lines = [
        f"Liczba szyn: {summary.get('total_buses', '—')}",
        f"Liczba galezi: {summary.get('total_branches', '—')}",
        f"Zbieznosc A: {'Tak' if summary.get('converged_a') else 'Nie'}",
        f"Zbieznosc B: {'Tak' if summary.get('converged_b') else 'Nie'}",
        f"Delta strat P: {summary.get('delta_total_losses_p_mw', 0):.4g} MW",
        f"Liczba problemow: {summary.get('total_issues', 0)}",
    ]
    for line in summary_lines:
        c.drawString(left_margin, y, line)
        y -= line_height

    y -= 5 * mm

    # Ranking
    c.setFont("Helvetica-Bold", 12)
    c.drawString(left_margin, y, "Ranking problemow (top 15)")
    y -= 5 * mm

    ranking = comparison.get("ranking", [])
    severity_labels = {5: "Krytyczny", 4: "Powazny", 3: "Sredni", 2: "Drobny", 1: "Info"}

    c.setFont("Helvetica", 9)
    for issue in ranking[:15]:
        severity = severity_labels.get(issue.get("severity", 1), "?")
        text = f"[{severity}] {issue.get('issue_code', '—')}: {issue.get('description_pl', '—')[:50]}"
        c.drawString(left_margin, y, text)
        y -= line_height
        if y < 30 * mm:
            c.showPage()
            y = top_margin

    if not ranking:
        c.drawString(left_margin, y, "Brak wykrytych problemow.")
        y -= line_height

    c.save()
    buffer.seek(0)

    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="power_flow_comparison_{comparison_id}.pdf"'},
    )
