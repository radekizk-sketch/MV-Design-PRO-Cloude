"""
Raportowanie wzorców odniesienia — eksport do DOCX i PDF.

Moduł generuje raporty diagnostyczne dla wzorców odniesienia A i C
w formatach DOCX i PDF. Raporty są przeznaczone do:
- załączników diagnostycznych „do podpisu",
- archiwizacji projektu,
- materiału audytowego (ślad obliczeń).

CANONICAL ALIGNMENT:
- Warstwa prezentacji — NIE liczy fizyki
- Pełna powtarzalność binarna dla DOCX (identyczne wejście → identyczne bajty)
- Treści po polsku, bez nazw kodowych

DETERMINIZM:
- DOCX: pełna powtarzalność binarna (SHA256)
- PDF: stabilne metadane (może wymagać dodatkowej normalizacji)
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from .base import ReferencePatternResult, VERDICT_DESCRIPTIONS_PL

# DOCX availability check
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# PDF availability check
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm

    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False

# Import DOCX determinism utilities
from network_model.reporting.docx_determinism import make_docx_deterministic


# =============================================================================
# METADATA TYPES
# =============================================================================


@dataclass(frozen=True)
class ReportMetadata:
    """
    Metadane raportu wzorca odniesienia.

    Wszystkie pola są wymagane do zapewnienia pełnej identyfikowalności raportu.
    """

    project_name: str  # Nazwa projektu
    case_name: str  # Nazwa przypadku obliczeniowego
    execution_date: str  # Data wykonania (format: YYYY-MM-DD)
    author: str | None = None  # Autor raportu (opcjonalnie)
    version: str = "1.0"  # Wersja raportu


# Stała data dla raportów (dla powtarzalności)
_FIXED_REPORT_DATE = "2000-01-01"


# =============================================================================
# FORMATOWANIE WARTOŚCI
# =============================================================================


def _format_value(value: Any) -> str:
    """Formatuje wartość do wyświetlenia w raporcie."""
    if value is None:
        return "—"
    if isinstance(value, float):
        # Formatowanie liczb zmiennoprzecinkowych
        if abs(value) >= 1000:
            return f"{value:,.1f}".replace(",", " ")
        elif abs(value) < 0.01 and value != 0:
            return f"{value:.4g}"
        else:
            return f"{value:.2f}"
    if isinstance(value, bool):
        return "Tak" if value else "Nie"
    if isinstance(value, (list, tuple)):
        return f"[{len(value)} elementów]"
    if isinstance(value, dict):
        return f"{{słownik, {len(value)} kluczy}}"
    return str(value)


def _verdict_color_docx(verdict: str) -> tuple[int, int, int]:
    """Zwraca kolor RGB dla werdyktu (do stylizacji DOCX)."""
    colors = {
        "ZGODNE": (34, 139, 34),  # Zielony (ForestGreen)
        "GRANICZNE": (255, 165, 0),  # Pomarańczowy (Orange)
        "NIEZGODNE": (178, 34, 34),  # Czerwony (Firebrick)
    }
    return colors.get(verdict, (0, 0, 0))


# =============================================================================
# DOCX GENERATOR
# =============================================================================


def export_reference_pattern_to_docx(
    result: ReferencePatternResult,
    path: str | Path,
    metadata: ReportMetadata,
    *,
    include_trace: bool = True,
) -> Path:
    """
    Eksportuje wynik wzorca odniesienia do pliku DOCX.

    Generuje raport diagnostyczny zawierający:
    - Stronę tytułową (projekt, data stała, nazwa wzorca)
    - Streszczenie z werdyktem
    - Dane wejściowe (tabela)
    - Sprawdzenia (tabela z wynikami)
    - Wartości pośrednie (artefakty)
    - Ślad obliczeń (opcjonalnie)
    - Informację o powtarzalności

    Args:
        result: Wynik walidacji wzorca odniesienia
        path: Ścieżka do pliku wyjściowego
        metadata: Metadane raportu
        include_trace: Czy dołączyć ślad obliczeń (domyślnie: tak)

    Returns:
        Ścieżka do wygenerowanego pliku DOCX

    Raises:
        ImportError: Gdy python-docx nie jest zainstalowany
    """
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx jest wymagany do eksportu DOCX. "
            "Zainstaluj: pip install python-docx"
        )

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Utwórz dokument
    doc = Document()

    # Konfiguracja stylu domyślnego
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # 1) STRONA TYTUŁOWA
    _add_docx_title_page(doc, result, metadata)

    # 2) STRESZCZENIE
    _add_docx_summary_section(doc, result)

    # 3) DANE WEJŚCIOWE (artefakty jako dane wejściowe)
    _add_docx_input_data_section(doc, result)

    # 4) SPRAWDZENIA
    _add_docx_checks_section(doc, result)

    # 5) WARTOŚCI POŚREDNIE
    _add_docx_artifacts_section(doc, result)

    # 6) ŚLAD OBLICZEŃ (opcjonalnie)
    if include_trace:
        _add_docx_trace_section(doc, result)

    # 7) INFORMACJA O POWTARZALNOŚCI
    _add_docx_reproducibility_section(doc, metadata)

    # Zapisz dokument
    doc.save(str(output_path))

    # Normalizacja dla powtarzalności binarnej
    make_docx_deterministic(output_path)

    return output_path


def _add_docx_title_page(
    doc: Document,
    result: ReferencePatternResult,
    metadata: ReportMetadata,
) -> None:
    """Dodaje stronę tytułową do dokumentu DOCX."""
    # Tytuł główny
    title = doc.add_heading("Raport wzorca odniesienia", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Nazwa wzorca
    pattern_heading = doc.add_heading(result.name_pl, level=1)
    pattern_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Metadane projektu (tabela)
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"

    meta_rows = [
        ("Projekt:", metadata.project_name),
        ("Przypadek:", metadata.case_name),
        ("Data wykonania:", _FIXED_REPORT_DATE),  # Stała data dla powtarzalności
        ("Identyfikator wzorca:", result.pattern_id),
    ]

    for i, (label, value) in enumerate(meta_rows):
        cells = meta_table.rows[i].cells
        cells[0].text = label
        cells[1].text = str(value)
        # Pogrubienie etykiety
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()


def _add_docx_summary_section(doc: Document, result: ReferencePatternResult) -> None:
    """Dodaje sekcję streszczenia z werdyktem."""
    doc.add_heading("Streszczenie", level=1)

    # Werdykt
    verdict_para = doc.add_paragraph()
    verdict_para.add_run("Werdykt: ").bold = True
    verdict_run = verdict_para.add_run(result.verdict)
    verdict_run.bold = True

    # Opis werdyktu
    verdict_desc = VERDICT_DESCRIPTIONS_PL.get(result.verdict, result.verdict)
    doc.add_paragraph(f"({verdict_desc})")

    # Podsumowanie
    doc.add_paragraph()
    summary_para = doc.add_paragraph()
    summary_para.add_run("Podsumowanie: ").bold = True
    doc.add_paragraph(result.summary_pl)

    doc.add_paragraph()


def _add_docx_input_data_section(
    doc: Document,
    result: ReferencePatternResult,
) -> None:
    """Dodaje sekcję danych wejściowych."""
    doc.add_heading("Dane wejściowe", level=1)

    # Wybierz kluczowe dane wejściowe z artefaktów
    input_keys = [
        "line_id",
        "line_name",
        "ct_ratio",
        "punkt_zabezpieczenia_id",
        "punkt_zabezpieczenia_nazwa",
        "szyny_id",
        "szyny_nazwa",
        "liczba_zrodel_generacji",
        "sumaryczna_moc_generacji_kw",
    ]

    input_data = [
        (k, result.artifacts.get(k))
        for k in input_keys
        if result.artifacts.get(k) is not None
    ]

    if not input_data:
        doc.add_paragraph("Brak szczegółowych danych wejściowych w artefaktach.")
    else:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Nagłówek
        hdr = table.rows[0].cells
        hdr[0].text = "Parametr"
        hdr[1].text = "Wartość"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        # Dane
        for key, value in input_data:
            row = table.add_row().cells
            row[0].text = _translate_key(key)
            row[1].text = _format_value(value)

    doc.add_paragraph()


def _add_docx_checks_section(doc: Document, result: ReferencePatternResult) -> None:
    """Dodaje sekcję sprawdzeń."""
    doc.add_heading("Sprawdzenia", level=1)

    if not result.checks:
        doc.add_paragraph("Brak sprawdzeń do wyświetlenia.")
        return

    # Tabela sprawdzeń
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Nagłówek
    hdr = table.rows[0].cells
    hdr[0].text = "Sprawdzenie"
    hdr[1].text = "Wynik"
    hdr[2].text = "Uzasadnienie"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Sprawdzenia (sortowane deterministycznie)
    for check in result.checks:
        row = table.add_row().cells
        row[0].text = check.get("name_pl", "—")
        row[1].text = check.get("status_pl", check.get("status", "—"))
        row[2].text = check.get("description_pl", "—")

    doc.add_paragraph()


def _add_docx_artifacts_section(
    doc: Document,
    result: ReferencePatternResult,
) -> None:
    """Dodaje sekcję wartości pośrednich (artefaktów)."""
    doc.add_heading("Wartości pośrednie", level=1)

    if not result.artifacts:
        doc.add_paragraph("Brak wartości pośrednich.")
        return

    # Filtruj artefakty numeryczne (pomijając już wyświetlone dane wejściowe)
    numeric_keys = [
        "tk_total_s",
        "ithn_a",
        "ithdop_a",
        "i_min_sel_primary_a",
        "i_max_sens_primary_a",
        "i_max_th_primary_a",
        "window_i_min_primary_a",
        "window_i_max_primary_a",
        "window_valid",
        "recommended_setting_secondary_a",
        "zmiana_pradu_zwarciowego_pct",
        "wklad_generacji_max_a",
        "i_wyzszy_stopien_a",
        "i_nizszy_stopien_a",
        "ik_bez_generacji_3f_a",
        "ik_z_generacja_max_3f_a",
    ]

    artifacts_to_show = [
        (k, result.artifacts.get(k))
        for k in numeric_keys
        if result.artifacts.get(k) is not None
    ]

    if not artifacts_to_show:
        # Pokaż wszystkie artefakty jeśli żaden z kluczy nie pasuje
        artifacts_to_show = list(result.artifacts.items())

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Nagłówek
    hdr = table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "Wartość"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Wartości
    for key, value in artifacts_to_show:
        row = table.add_row().cells
        row[0].text = _translate_key(key)
        row[1].text = _format_value(value)

    doc.add_paragraph()


def _add_docx_trace_section(doc: Document, result: ReferencePatternResult) -> None:
    """Dodaje sekcję śladu obliczeń."""
    doc.add_heading("Ślad obliczeń (skrót)", level=1)

    if not result.trace:
        doc.add_paragraph("Brak śladu obliczeń.")
        return

    # Pokaż tylko kluczowe kroki (max 10)
    max_steps = 10
    trace_to_show = list(result.trace)[:max_steps]

    for i, step in enumerate(trace_to_show, 1):
        step_name = step.get("step", f"krok_{i}")
        description = step.get("description_pl", "—")

        # Nagłówek kroku
        step_para = doc.add_paragraph()
        step_para.add_run(f"{i}. {step_name}: ").bold = True
        step_para.add_run(description)

        # Formuła (jeśli dostępna)
        formula = step.get("formula")
        if formula:
            formula_para = doc.add_paragraph()
            formula_para.add_run("   Wzór: ").italic = True
            formula_para.add_run(formula)

        # Wynik (jeśli dostępny)
        outputs = step.get("outputs", {})
        if outputs:
            outputs_str = ", ".join(f"{k}={_format_value(v)}" for k, v in outputs.items())
            output_para = doc.add_paragraph()
            output_para.add_run("   Wynik: ").italic = True
            output_para.add_run(outputs_str)

    if len(result.trace) > max_steps:
        doc.add_paragraph(f"... oraz {len(result.trace) - max_steps} dodatkowych kroków.")

    doc.add_paragraph()


def _add_docx_reproducibility_section(
    doc: Document,
    metadata: ReportMetadata,
) -> None:
    """Dodaje sekcję informacji o powtarzalności."""
    doc.add_heading("Informacje o raporcie", level=1)

    info_para = doc.add_paragraph()
    info_para.add_run("Powtarzalność: ").bold = True
    info_para.add_run(
        "Raport jest generowany w sposób deterministyczny. "
        "Identyczne dane wejściowe zawsze generują identyczny plik (SHA256)."
    )

    doc.add_paragraph()

    scope_para = doc.add_paragraph()
    scope_para.add_run("Zakres raportu: ").bold = True
    scope_para.add_run(
        "Raport zawiera wyniki walidacji wzorca odniesienia, "
        "sprawdzenia kryteriów oraz skrócony ślad obliczeń. "
        "Pełny ślad dostępny jest w formacie JSON."
    )

    doc.add_paragraph()

    version_para = doc.add_paragraph()
    version_para.add_run("Wersja raportu: ").bold = True
    version_para.add_run(metadata.version)


def _translate_key(key: str) -> str:
    """Tłumaczy klucz artefaktu na polski."""
    translations = {
        "tk_total_s": "Czas zwarciowy całkowity [s]",
        "ithn_a": "Prąd termiczny znamionowy [A]",
        "ithdop_a": "Prąd termiczny dopuszczalny [A]",
        "i_min_sel_primary_a": "I_min (selektywność) [A]",
        "i_max_sens_primary_a": "I_max (czułość) [A]",
        "i_max_th_primary_a": "I_max (cieplne) [A]",
        "window_i_min_primary_a": "Okno nastaw: I_min [A]",
        "window_i_max_primary_a": "Okno nastaw: I_max [A]",
        "window_valid": "Okno nastaw prawidłowe",
        "recommended_setting_secondary_a": "Zalecana nastawa (wtórna) [A]",
        "line_id": "ID linii",
        "line_name": "Nazwa linii",
        "ct_ratio": "Przekładnia BI",
        "punkt_zabezpieczenia_id": "ID punktu zabezpieczenia",
        "punkt_zabezpieczenia_nazwa": "Nazwa punktu zabezpieczenia",
        "szyny_id": "ID szyn",
        "szyny_nazwa": "Nazwa szyn",
        "liczba_zrodel_generacji": "Liczba źródeł generacji",
        "sumaryczna_moc_generacji_kw": "Sumaryczna moc generacji [kW]",
        "zmiana_pradu_zwarciowego_pct": "Zmiana prądu zwarciowego [%]",
        "wklad_generacji_max_a": "Wkład generacji max [A]",
        "i_wyzszy_stopien_a": "I>> (wyższy stopień) [A]",
        "i_nizszy_stopien_a": "I> (niższy stopień) [A]",
        "ik_bez_generacji_3f_a": "Ik (bez generacji) [A]",
        "ik_z_generacja_max_3f_a": "Ik (z generacją max) [A]",
        "limiting_criterion_min": "Kryterium ograniczające (min)",
        "limiting_criterion_max": "Kryterium ograniczające (max)",
    }
    return translations.get(key, key)


# =============================================================================
# PDF GENERATOR
# =============================================================================


def export_reference_pattern_to_pdf(
    result: ReferencePatternResult,
    path: str | Path,
    metadata: ReportMetadata,
    *,
    include_trace: bool = True,
) -> Path:
    """
    Eksportuje wynik wzorca odniesienia do pliku PDF.

    Generuje raport diagnostyczny zawierający:
    - Stronę tytułową (projekt, data stała, nazwa wzorca)
    - Streszczenie z werdyktem
    - Dane wejściowe
    - Sprawdzenia
    - Wartości pośrednie
    - Ślad obliczeń (opcjonalnie)
    - Informację o powtarzalności

    UWAGA: PDF może nie być w pełni binarnie powtarzalny ze względu na
    ograniczenia biblioteki reportlab. W takim przypadku test powtarzalności
    zostanie oznaczony jako oczekiwane niepowodzenie z uzasadnieniem.

    Args:
        result: Wynik walidacji wzorca odniesienia
        path: Ścieżka do pliku wyjściowego
        metadata: Metadane raportu
        include_trace: Czy dołączyć ślad obliczeń (domyślnie: tak)

    Returns:
        Ścieżka do wygenerowanego pliku PDF

    Raises:
        ImportError: Gdy reportlab nie jest zainstalowany
    """
    if not _PDF_AVAILABLE:
        raise ImportError(
            "reportlab jest wymagany do eksportu PDF. "
            "Zainstaluj: pip install reportlab"
        )

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Utwórz canvas PDF z stabilnymi metadanymi
    c = canvas.Canvas(str(output_path), pagesize=A4)

    # Ustaw stałe metadane dokumentu (dla powtarzalności)
    c.setTitle(f"Raport wzorca odniesienia: {result.name_pl}")
    c.setAuthor(metadata.author or "MV-DESIGN-PRO")
    c.setSubject(f"Wzorzec: {result.pattern_id}")
    c.setCreator("MV-DESIGN-PRO Reference Pattern Reporter")

    page_width, page_height = A4

    # Marginesy i układ
    left_margin = 25 * mm
    right_margin = page_width - 25 * mm
    top_margin = page_height - 25 * mm
    bottom_margin = 30 * mm
    line_height = 5 * mm
    section_spacing = 8 * mm

    y = top_margin

    def check_page_break(needed_height: float = 20 * mm) -> float:
        nonlocal y
        if y - needed_height < bottom_margin:
            c.showPage()
            return top_margin
        return y

    def draw_text(
        text: str, x: float, font_size: int = 10, bold: bool = False
    ) -> None:
        nonlocal y
        font_name = "Helvetica-Bold" if bold else "Helvetica"
        c.setFont(font_name, font_size)
        c.drawString(x, y, text)
        y -= line_height

    def draw_wrapped_text(
        text: str, x: float, max_width: float, font_size: int = 10
    ) -> None:
        nonlocal y
        c.setFont("Helvetica", font_size)
        words = text.split()
        current_line = ""
        for word in words:
            test_line = f"{current_line} {word}".strip()
            if c.stringWidth(test_line, "Helvetica", font_size) < max_width:
                current_line = test_line
            else:
                if current_line:
                    y = check_page_break(line_height)
                    c.drawString(x, y, current_line)
                    y -= line_height
                current_line = word
        if current_line:
            y = check_page_break(line_height)
            c.drawString(x, y, current_line)
            y -= line_height

    # 1) STRONA TYTUŁOWA
    # Tytuł główny
    c.setFont("Helvetica-Bold", 18)
    title = "Raport wzorca odniesienia"
    title_width = c.stringWidth(title, "Helvetica-Bold", 18)
    c.drawString((page_width - title_width) / 2, y, title)
    y -= 12 * mm

    # Nazwa wzorca
    c.setFont("Helvetica-Bold", 14)
    pattern_name = result.name_pl
    if c.stringWidth(pattern_name, "Helvetica-Bold", 14) > (right_margin - left_margin):
        # Jeśli za długa, użyj mniejszej czcionki
        c.setFont("Helvetica-Bold", 12)
    name_width = c.stringWidth(pattern_name, "Helvetica-Bold", 12)
    c.drawString((page_width - name_width) / 2, y, pattern_name)
    y -= 10 * mm

    # Metadane
    c.setFont("Helvetica", 10)
    meta_lines = [
        f"Projekt: {metadata.project_name}",
        f"Przypadek: {metadata.case_name}",
        f"Data wykonania: {_FIXED_REPORT_DATE}",
        f"Identyfikator wzorca: {result.pattern_id}",
    ]
    for line in meta_lines:
        y = check_page_break(line_height)
        line_width = c.stringWidth(line, "Helvetica", 10)
        c.drawString((page_width - line_width) / 2, y, line)
        y -= line_height

    y -= section_spacing

    # 2) STRESZCZENIE
    y = check_page_break(30 * mm)
    draw_text("Streszczenie", left_margin, font_size=14, bold=True)
    y -= 3 * mm

    # Werdykt
    y = check_page_break(line_height * 2)
    c.setFont("Helvetica-Bold", 11)
    c.drawString(left_margin, y, f"Werdykt: {result.verdict}")
    y -= line_height

    # Opis werdyktu
    verdict_desc = VERDICT_DESCRIPTIONS_PL.get(result.verdict, result.verdict)
    c.setFont("Helvetica", 10)
    c.drawString(left_margin, y, f"({verdict_desc})")
    y -= line_height * 2

    # Podsumowanie
    c.setFont("Helvetica-Bold", 10)
    c.drawString(left_margin, y, "Podsumowanie:")
    y -= line_height
    draw_wrapped_text(result.summary_pl, left_margin, right_margin - left_margin, 10)

    y -= section_spacing

    # 3) SPRAWDZENIA
    y = check_page_break(30 * mm)
    draw_text("Sprawdzenia", left_margin, font_size=14, bold=True)
    y -= 3 * mm

    for check in result.checks:
        y = check_page_break(line_height * 4)

        # Nazwa sprawdzenia i status
        name = check.get("name_pl", "—")
        status = check.get("status_pl", check.get("status", "—"))

        c.setFont("Helvetica-Bold", 10)
        c.drawString(left_margin, y, f"{name}: {status}")
        y -= line_height

        # Uzasadnienie
        description = check.get("description_pl", "")
        if description:
            draw_wrapped_text(
                description, left_margin + 5 * mm, right_margin - left_margin - 5 * mm, 9
            )

        y -= 3 * mm

    y -= section_spacing

    # 4) WARTOŚCI POŚREDNIE
    y = check_page_break(30 * mm)
    draw_text("Wartości pośrednie", left_margin, font_size=14, bold=True)
    y -= 3 * mm

    # Kluczowe artefakty
    key_artifacts = [
        ("tk_total_s", "Czas zwarciowy [s]"),
        ("ithdop_a", "I_th,dop [A]"),
        ("window_i_min_primary_a", "Okno I_min [A]"),
        ("window_i_max_primary_a", "Okno I_max [A]"),
        ("window_valid", "Okno prawidłowe"),
        ("zmiana_pradu_zwarciowego_pct", "Zmiana Ik [%]"),
        ("wklad_generacji_max_a", "Wkład generacji [A]"),
    ]

    label_x = left_margin
    value_x = left_margin + 60 * mm

    for key, label in key_artifacts:
        value = result.artifacts.get(key)
        if value is not None:
            y = check_page_break(line_height)
            c.setFont("Helvetica", 10)
            c.drawString(label_x, y, label + ":")
            c.drawString(value_x, y, _format_value(value))
            y -= line_height

    y -= section_spacing

    # 5) ŚLAD OBLICZEŃ (skrót)
    if include_trace and result.trace:
        y = check_page_break(30 * mm)
        draw_text("Ślad obliczeń (skrót)", left_margin, font_size=14, bold=True)
        y -= 3 * mm

        max_steps = 5
        for i, step in enumerate(list(result.trace)[:max_steps], 1):
            y = check_page_break(line_height * 2)

            step_name = step.get("step", f"krok_{i}")
            description = step.get("description_pl", "—")

            c.setFont("Helvetica-Bold", 9)
            c.drawString(left_margin, y, f"{i}. {step_name}")
            y -= line_height

            c.setFont("Helvetica", 9)
            # Skróć opis jeśli za długi
            if len(description) > 80:
                description = description[:77] + "..."
            c.drawString(left_margin + 5 * mm, y, description)
            y -= line_height

        if len(result.trace) > max_steps:
            y = check_page_break(line_height)
            c.setFont("Helvetica-Oblique", 9)
            c.drawString(
                left_margin, y, f"... oraz {len(result.trace) - max_steps} dodatkowych kroków"
            )
            y -= line_height

    y -= section_spacing

    # 6) INFORMACJA O POWTARZALNOŚCI
    y = check_page_break(30 * mm)
    draw_text("Informacje o raporcie", left_margin, font_size=14, bold=True)
    y -= 3 * mm

    c.setFont("Helvetica", 9)
    info_text = (
        "Raport generowany przez MV-DESIGN-PRO. "
        "Dane wejściowe określają deterministycznie treść raportu."
    )
    draw_wrapped_text(info_text, left_margin, right_margin - left_margin, 9)

    # Zapisz PDF
    c.save()

    return output_path


# =============================================================================
# PUBLIC API
# =============================================================================


__all__ = [
    "export_reference_pattern_to_docx",
    "export_reference_pattern_to_pdf",
    "ReportMetadata",
]
