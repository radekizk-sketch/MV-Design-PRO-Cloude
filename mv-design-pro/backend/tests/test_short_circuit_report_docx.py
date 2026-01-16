"""
Tests for short-circuit result DOCX report generator.

Verifies:
- Correct file writing and parent directory creation
- Valid DOCX document creation (openable by python-docx)
- Presence of key sections (title, results, white box)
- Proper handling of include_white_box=False
- No numerical value assertions (only structure/presence checks)
"""

from __future__ import annotations

from pathlib import Path

import pytest

from network_model.core.branch import BranchType, LineBranch, TransformerBranch
from network_model.core.graph import NetworkGraph
from network_model.core.node import Node, NodeType
from network_model.reporting.short_circuit_report_docx import (
    export_short_circuit_result_to_docx,
)
from network_model.solvers.short_circuit_iec60909 import (
    ShortCircuitIEC60909Solver,
)

# Check if python-docx is available for tests
try:
    from docx import Document

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


pytestmark = pytest.mark.skipif(
    not _DOCX_AVAILABLE,
    reason="python-docx is not installed",
)


# -----------------------------------------------------------------------------
# Test Helpers (reused from existing tests)
# -----------------------------------------------------------------------------
def create_pq_node(node_id: str, voltage_level: float) -> Node:
    """Create a PQ node with the given voltage level."""
    return Node(
        id=node_id,
        name=f"Node {node_id}",
        node_type=NodeType.PQ,
        voltage_level=voltage_level,
        active_power=5.0,
        reactive_power=2.0,
    )


def create_reference_node(node_id: str, voltage_level: float) -> Node:
    """Create a reference (ground) node."""
    return Node(
        id=node_id,
        name=f"Reference {node_id}",
        node_type=NodeType.PQ,
        voltage_level=voltage_level,
        active_power=0.0,
        reactive_power=0.0,
    )


def create_transformer_branch(
    branch_id: str,
    from_node_id: str,
    to_node_id: str,
    rated_power_mva: float,
    voltage_hv_kv: float,
    voltage_lv_kv: float,
    uk_percent: float,
    pk_kw: float,
) -> TransformerBranch:
    """Create a transformer branch with specified parameters."""
    return TransformerBranch(
        id=branch_id,
        name=f"Transformer {branch_id}",
        branch_type=BranchType.TRANSFORMER,
        from_node_id=from_node_id,
        to_node_id=to_node_id,
        in_service=True,
        rated_power_mva=rated_power_mva,
        voltage_hv_kv=voltage_hv_kv,
        voltage_lv_kv=voltage_lv_kv,
        uk_percent=uk_percent,
        pk_kw=pk_kw,
        i0_percent=0.0,
        p0_kw=0.0,
        vector_group="Dyn11",
        tap_position=0,
        tap_step_percent=2.5,
    )


def create_reference_branch(
    branch_id: str,
    from_node_id: str,
    to_node_id: str,
    r_ohm: float,
) -> LineBranch:
    """Create a reference (very high impedance) line branch."""
    return LineBranch(
        id=branch_id,
        name=f"Reference {branch_id}",
        branch_type=BranchType.LINE,
        from_node_id=from_node_id,
        to_node_id=to_node_id,
        r_ohm_per_km=r_ohm,
        x_ohm_per_km=0.0,
        b_us_per_km=0.0,
        length_km=1.0,
        rated_current_a=0.0,
    )


def build_transformer_only_graph() -> NetworkGraph:
    """Build a minimal test graph with a single transformer."""
    graph = NetworkGraph()
    graph.add_node(create_pq_node("A", 110.0))
    graph.add_node(create_pq_node("B", 20.0))
    graph.add_node(create_reference_node("GND", 20.0))

    transformer = create_transformer_branch(
        "T1",
        "A",
        "B",
        rated_power_mva=25.0,
        voltage_hv_kv=110.0,
        voltage_lv_kv=20.0,
        uk_percent=10.0,
        pk_kw=120.0,
    )
    graph.add_branch(transformer)
    graph.add_branch(create_reference_branch("REF", "B", "GND", r_ohm=1e9))
    return graph


def _get_document_text(doc: Document) -> str:
    """Extract all text from a Document for assertion checks."""
    text_parts = []
    for para in doc.paragraphs:
        text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    return "\n".join(text_parts)


# -----------------------------------------------------------------------------
# Tests for export_short_circuit_result_to_docx
# -----------------------------------------------------------------------------
class TestExportShortCircuitResultToDocx:
    """Tests for DOCX report generation."""

    def test_export_creates_file(self, tmp_path: Path) -> None:
        """Verify that export creates a DOCX file at the specified path."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "report.docx"
        returned_path = export_short_circuit_result_to_docx(result, output_file)

        assert output_file.exists()
        assert returned_path == output_file

    def test_export_creates_non_empty_file(self, tmp_path: Path) -> None:
        """Verify that the created file has size > 0."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "report.docx"
        export_short_circuit_result_to_docx(result, output_file)

        assert output_file.stat().st_size > 0

    def test_export_creates_parent_directories(self, tmp_path: Path) -> None:
        """Verify that export creates parent directories if they don't exist."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "nested" / "deep" / "report.docx"
        export_short_circuit_result_to_docx(result, output_file)

        assert output_file.exists()
        assert output_file.parent.exists()

    def test_export_produces_openable_docx(self, tmp_path: Path) -> None:
        """Verify that the exported file can be opened by python-docx."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "report.docx"
        export_short_circuit_result_to_docx(result, output_file)

        # Should not raise any exception
        doc = Document(output_file)
        assert doc is not None

    def test_export_contains_default_title(self, tmp_path: Path) -> None:
        """Verify that document contains the default title."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "report.docx"
        export_short_circuit_result_to_docx(result, output_file)

        doc = Document(output_file)
        full_text = _get_document_text(doc)

        assert "Raport zwarciowy IEC 60909" in full_text

    def test_export_contains_results_section(self, tmp_path: Path) -> None:
        """Verify that document contains the Results section."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "report.docx"
        export_short_circuit_result_to_docx(result, output_file)

        doc = Document(output_file)
        full_text = _get_document_text(doc)

        assert "Wyniki" in full_text

    def test_export_contains_key_result_labels(self, tmp_path: Path) -> None:
        """Verify that document contains key result labels like Ik'' and Sk."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "report.docx"
        export_short_circuit_result_to_docx(result, output_file)

        doc = Document(output_file)
        full_text = _get_document_text(doc)

        assert "Ik''" in full_text
        assert "Sk" in full_text

    def test_export_contains_white_box_section(self, tmp_path: Path) -> None:
        """Verify that document contains White Box section when included."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "report.docx"
        export_short_circuit_result_to_docx(result, output_file, include_white_box=True)

        doc = Document(output_file)
        full_text = _get_document_text(doc)

        assert "White Box" in full_text

    def test_export_with_custom_title(self, tmp_path: Path) -> None:
        """Verify that custom title is used when provided."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        custom_title = "Custom Short-Circuit Report"
        output_file = tmp_path / "report.docx"
        export_short_circuit_result_to_docx(result, output_file, title=custom_title)

        doc = Document(output_file)
        full_text = _get_document_text(doc)

        assert custom_title in full_text

    def test_export_accepts_string_path(self, tmp_path: Path) -> None:
        """Verify that export accepts both str and Path for the path argument."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = str(tmp_path / "report.docx")
        returned_path = export_short_circuit_result_to_docx(result, output_file)

        assert Path(output_file).exists()
        assert isinstance(returned_path, Path)


class TestExportWithoutWhiteBox:
    """Tests for DOCX export with include_white_box=False."""

    def test_export_without_white_box_creates_file(self, tmp_path: Path) -> None:
        """Verify that export without white box creates a valid file."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "report.docx"
        export_short_circuit_result_to_docx(result, output_file, include_white_box=False)

        assert output_file.exists()
        assert output_file.stat().st_size > 0

    def test_export_without_white_box_contains_results(self, tmp_path: Path) -> None:
        """Verify that document contains Results section even without White Box."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "report.docx"
        export_short_circuit_result_to_docx(result, output_file, include_white_box=False)

        doc = Document(output_file)
        full_text = _get_document_text(doc)

        assert "Wyniki" in full_text

    def test_export_without_white_box_omits_white_box_section(
        self, tmp_path: Path
    ) -> None:
        """Verify that White Box section is not present when include_white_box=False."""
        graph = build_transformer_only_graph()
        result = ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
            graph=graph,
            fault_node_id="B",
            c_factor=1.0,
            tk_s=1.0,
        )

        output_file = tmp_path / "report.docx"
        export_short_circuit_result_to_docx(result, output_file, include_white_box=False)

        doc = Document(output_file)
        full_text = _get_document_text(doc)

        assert "White Box" not in full_text


class TestExportErrorHandling:
    """Tests for error handling in DOCX export."""

    def test_export_validates_to_dict_returns_dict(self, tmp_path: Path) -> None:
        """Verify that export raises ValueError if to_dict doesn't return dict."""

        class BadResult:
            def to_dict(self) -> str:
                return "not a dict"

        output_file = tmp_path / "report.docx"

        with pytest.raises(ValueError, match="must return a dict"):
            export_short_circuit_result_to_docx(BadResult(), output_file)  # type: ignore
