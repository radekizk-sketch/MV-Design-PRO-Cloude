"""FIX-12C: E2E Determinism Tests for Protection Coordination Export (PDF/DOCX).

Ten moduł testuje deterministyczność eksportów PDF i DOCX dla wyników
analizy koordynacji zabezpieczeń nadprądowych.

Scenariusze testowe:
1. Ten sam wynik koordynacji -> 2x eksport DOCX -> identyczny binarnie (SHA256)
2. Ten sam wynik koordynacji -> 2x eksport PDF -> identyczny binarnie (SHA256)
3. Walidacja zawartości sekcji w raporcie

Wymagania:
- Export DOCX deterministyczny dla tego samego wejścia
- PDF deterministyczny dzięki reportlab invariant mode
- SHA256 plików identyczne między run1/run2

CANONICAL ALIGNMENT:
- NOT-A-SOLVER: Testy nie modyfikują solverów ani eksporterów
- Tylko weryfikacja deterministyczności eksportów
- 100% Polish labels w raportach
"""
from __future__ import annotations

import hashlib
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import UUID

import pytest

# Check for optional dependencies
try:
    from network_model.reporting.protection_report_pdf import (
        export_protection_coordination_to_pdf,
        _PDF_AVAILABLE,
    )
except ImportError:
    _PDF_AVAILABLE = False
    export_protection_coordination_to_pdf = None  # type: ignore

try:
    from network_model.reporting.protection_report_docx import (
        export_protection_coordination_to_docx,
        _DOCX_AVAILABLE,
    )
except ImportError:
    _DOCX_AVAILABLE = False
    export_protection_coordination_to_docx = None  # type: ignore


# =============================================================================
# Test Fixtures - Deterministic Protection Coordination Result
# =============================================================================


def _create_deterministic_protection_result() -> dict[str, Any]:
    """
    Tworzy deterministyczny wynik analizy koordynacji zabezpieczeń.

    Wszystkie wartości są stałe, bez użycia datetime.now() ani random.
    Sortowanie jest stabilne po device_id/name.
    """
    # Fixed timestamp for determinism
    fixed_timestamp = "2024-01-01T00:00:00+00:00"

    # Fixed device IDs (sorted alphabetically)
    device_a_id = "DEV-A-001"
    device_b_id = "DEV-B-002"

    return {
        "run_id": "run_deterministic_e2e_test_001",
        "project_id": "proj_deterministic_001",
        "created_at": fixed_timestamp,
        "overall_verdict": "PASS",
        "devices": [
            {
                "id": str(UUID(int=1)),
                "name": "Przekaźnik A",
                "device_type": "RELAY",
                "manufacturer": "ABB",
                "model": "REF615",
                "location_element_id": "BUS-001",
                "location_description": "Rozdzielnia główna",
                "settings": {
                    "stage_51": {
                        "enabled": True,
                        "pickup_current_a": 100.0,
                        "time_s": None,
                        "curve_settings": {
                            "standard": "IEC",
                            "variant": "SI",
                            "pickup_current_a": 100.0,
                            "time_multiplier": 0.3,
                            "definite_time_s": None,
                            "reset_time_s": 0.0,
                        },
                        "directional": False,
                    },
                    "stage_50": None,
                    "stage_50_high": None,
                    "stage_51n": None,
                    "stage_50n": None,
                },
                "ct_ratio": "400/5",
                "rated_current_a": 400.0,
                "created_at": fixed_timestamp,
            },
            {
                "id": str(UUID(int=2)),
                "name": "Przekaźnik B",
                "device_type": "RELAY",
                "manufacturer": "Siemens",
                "model": "7SJ82",
                "location_element_id": "BUS-002",
                "location_description": "Rozdzielnia odbiorcza",
                "settings": {
                    "stage_51": {
                        "enabled": True,
                        "pickup_current_a": 80.0,
                        "time_s": None,
                        "curve_settings": {
                            "standard": "IEC",
                            "variant": "VI",
                            "pickup_current_a": 80.0,
                            "time_multiplier": 0.2,
                            "definite_time_s": None,
                            "reset_time_s": 0.0,
                        },
                        "directional": False,
                    },
                    "stage_50": None,
                    "stage_50_high": None,
                    "stage_51n": None,
                    "stage_50n": None,
                },
                "ct_ratio": "200/5",
                "rated_current_a": 200.0,
                "created_at": fixed_timestamp,
            },
        ],
        "sensitivity_checks": [
            {
                "device_id": device_a_id,
                "i_fault_min_a": 500.0,
                "i_pickup_a": 100.0,
                "margin_percent": 400.0,
                "verdict": "PASS",
                "verdict_pl": "Prawidłowa",
                "notes_pl": "Margines czułości wystarczający (400%)",
            },
            {
                "device_id": device_b_id,
                "i_fault_min_a": 300.0,
                "i_pickup_a": 80.0,
                "margin_percent": 275.0,
                "verdict": "PASS",
                "verdict_pl": "Prawidłowa",
                "notes_pl": "Margines czułości wystarczający (275%)",
            },
        ],
        "selectivity_checks": [
            {
                "upstream_device_id": device_a_id,
                "downstream_device_id": device_b_id,
                "analysis_current_a": 1000.0,
                "t_upstream_s": 0.8,
                "t_downstream_s": 0.3,
                "margin_s": 0.5,
                "required_margin_s": 0.3,
                "verdict": "PASS",
                "verdict_pl": "Prawidłowa",
                "notes_pl": "Margines czasowy 500ms > wymagane 300ms",
            },
        ],
        "overload_checks": [
            {
                "device_id": device_a_id,
                "i_operating_a": 50.0,
                "i_pickup_a": 100.0,
                "margin_percent": 100.0,
                "verdict": "PASS",
                "verdict_pl": "Prawidłowa",
                "notes_pl": "Prąd roboczy znacznie poniżej progu rozruchowego",
            },
            {
                "device_id": device_b_id,
                "i_operating_a": 40.0,
                "i_pickup_a": 80.0,
                "margin_percent": 100.0,
                "verdict": "PASS",
                "verdict_pl": "Prawidłowa",
                "notes_pl": "Prąd roboczy znacznie poniżej progu rozruchowego",
            },
        ],
        "tcc_curves": [
            {
                "device_name": "Przekaźnik A",
                "curve_type": "SI",
                "pickup_current_a": 100.0,
                "time_multiplier": 0.3,
            },
            {
                "device_name": "Przekaźnik B",
                "curve_type": "VI",
                "pickup_current_a": 80.0,
                "time_multiplier": 0.2,
            },
        ],
        "summary": {
            "total_devices": 2,
            "total_checks": 5,
            "overall_verdict_pl": "Prawidłowa",
            "sensitivity": {"pass": 2, "fail": 0, "marginal": 0},
            "selectivity": {"pass": 1, "fail": 0, "marginal": 0},
            "overload": {"pass": 2, "fail": 0, "marginal": 0},
        },
    }


def _compute_file_hash(file_path: Path) -> str:
    """Oblicza SHA-256 hash pliku."""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256_hash.update(chunk)
    return sha256_hash.hexdigest()


# =============================================================================
# E2E Test Classes - DOCX Export
# =============================================================================


@pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx not installed")
class TestProtectionDOCXExportDeterminism:
    """E2E: Deterministyczność eksportu DOCX dla koordynacji zabezpieczeń."""

    def test_docx_export_identical_twice(self) -> None:
        """2x eksport DOCX -> identyczne pliki (SHA256)."""
        result = _create_deterministic_protection_result()

        with tempfile.TemporaryDirectory() as tmpdir:
            path_1 = Path(tmpdir) / "protection_report_1.docx"
            path_2 = Path(tmpdir) / "protection_report_2.docx"

            # Export 1
            export_protection_coordination_to_docx(
                result,
                path_1,
                title="Raport koordynacji zabezpieczeń - Test E2E",
                metadata={"project_name": "Projekt testowy", "created_at": "2024-01-01T00:00:00"},
                deterministic=True,
            )

            # Export 2
            export_protection_coordination_to_docx(
                result,
                path_2,
                title="Raport koordynacji zabezpieczeń - Test E2E",
                metadata={"project_name": "Projekt testowy", "created_at": "2024-01-01T00:00:00"},
                deterministic=True,
            )

            # Compare hashes
            hash_1 = _compute_file_hash(path_1)
            hash_2 = _compute_file_hash(path_2)

            assert hash_1 == hash_2, (
                f"DOCX export nie deterministyczny\n"
                f"Hash 1: {hash_1}\nHash 2: {hash_2}"
            )

    def test_docx_export_without_metadata_deterministic(self) -> None:
        """DOCX bez metadata jest deterministyczny."""
        result = _create_deterministic_protection_result()

        with tempfile.TemporaryDirectory() as tmpdir:
            path_1 = Path(tmpdir) / "no_meta_1.docx"
            path_2 = Path(tmpdir) / "no_meta_2.docx"

            export_protection_coordination_to_docx(result, path_1, deterministic=True)
            export_protection_coordination_to_docx(result, path_2, deterministic=True)

            hash_1 = _compute_file_hash(path_1)
            hash_2 = _compute_file_hash(path_2)

            assert hash_1 == hash_2, "DOCX bez metadata nie deterministyczny"

    def test_docx_export_empty_checks_deterministic(self) -> None:
        """DOCX z pustymi sprawdzeniami jest deterministyczny."""
        result = _create_deterministic_protection_result()
        # Clear all checks
        result["sensitivity_checks"] = []
        result["selectivity_checks"] = []
        result["overload_checks"] = []
        result["tcc_curves"] = []
        result["devices"] = []
        result["summary"]["total_devices"] = 0
        result["summary"]["total_checks"] = 0

        with tempfile.TemporaryDirectory() as tmpdir:
            path_1 = Path(tmpdir) / "empty_1.docx"
            path_2 = Path(tmpdir) / "empty_2.docx"

            export_protection_coordination_to_docx(result, path_1, deterministic=True)
            export_protection_coordination_to_docx(result, path_2, deterministic=True)

            hash_1 = _compute_file_hash(path_1)
            hash_2 = _compute_file_hash(path_2)

            assert hash_1 == hash_2, "DOCX z pustymi danymi nie deterministyczny"


# =============================================================================
# E2E Test Classes - PDF Export
# =============================================================================


@pytest.mark.skipif(not _PDF_AVAILABLE, reason="reportlab not installed")
class TestProtectionPDFExportDeterminism:
    """E2E: Deterministyczność eksportu PDF dla koordynacji zabezpieczeń.

    PDF używa reportlab invariant mode dla deterministycznego wyjścia.
    """

    def test_pdf_export_identical_twice(self) -> None:
        """2x eksport PDF -> identyczne pliki (SHA256)."""
        result = _create_deterministic_protection_result()

        with tempfile.TemporaryDirectory() as tmpdir:
            path_1 = Path(tmpdir) / "protection_report_1.pdf"
            path_2 = Path(tmpdir) / "protection_report_2.pdf"

            # Export 1
            export_protection_coordination_to_pdf(
                result,
                path_1,
                title="Raport koordynacji zabezpieczeń - Test E2E",
                metadata={"project_name": "Projekt testowy", "created_at": "2024-01-01T00:00:00"},
            )

            # Export 2
            export_protection_coordination_to_pdf(
                result,
                path_2,
                title="Raport koordynacji zabezpieczeń - Test E2E",
                metadata={"project_name": "Projekt testowy", "created_at": "2024-01-01T00:00:00"},
            )

            # Compare hashes
            hash_1 = _compute_file_hash(path_1)
            hash_2 = _compute_file_hash(path_2)

            assert hash_1 == hash_2, (
                f"PDF export nie deterministyczny\n"
                f"Hash 1: {hash_1}\nHash 2: {hash_2}"
            )

    def test_pdf_export_without_metadata_deterministic(self) -> None:
        """PDF bez metadata jest deterministyczny."""
        result = _create_deterministic_protection_result()

        with tempfile.TemporaryDirectory() as tmpdir:
            path_1 = Path(tmpdir) / "no_meta_1.pdf"
            path_2 = Path(tmpdir) / "no_meta_2.pdf"

            export_protection_coordination_to_pdf(result, path_1)
            export_protection_coordination_to_pdf(result, path_2)

            hash_1 = _compute_file_hash(path_1)
            hash_2 = _compute_file_hash(path_2)

            assert hash_1 == hash_2, "PDF bez metadata nie deterministyczny"

    def test_pdf_export_empty_checks_deterministic(self) -> None:
        """PDF z pustymi sprawdzeniami jest deterministyczny."""
        result = _create_deterministic_protection_result()
        # Clear all checks
        result["sensitivity_checks"] = []
        result["selectivity_checks"] = []
        result["overload_checks"] = []
        result["tcc_curves"] = []
        result["devices"] = []
        result["summary"]["total_devices"] = 0
        result["summary"]["total_checks"] = 0

        with tempfile.TemporaryDirectory() as tmpdir:
            path_1 = Path(tmpdir) / "empty_1.pdf"
            path_2 = Path(tmpdir) / "empty_2.pdf"

            export_protection_coordination_to_pdf(result, path_1)
            export_protection_coordination_to_pdf(result, path_2)

            hash_1 = _compute_file_hash(path_1)
            hash_2 = _compute_file_hash(path_2)

            assert hash_1 == hash_2, "PDF z pustymi danymi nie deterministyczny"


# =============================================================================
# E2E Test Classes - Content Validation
# =============================================================================


@pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx not installed")
class TestProtectionDOCXContentValidation:
    """E2E: Walidacja zawartości sekcji w raporcie DOCX."""

    def test_content_sections_present(self) -> None:
        """Wszystkie wymagane sekcje są obecne w raporcie DOCX."""
        from docx import Document as DocxDocument

        result = _create_deterministic_protection_result()

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "content_test.docx"
            export_protection_coordination_to_docx(result, path, deterministic=True)

            # Read DOCX and check sections
            doc = DocxDocument(str(path))
            full_text = "\n".join([p.text for p in doc.paragraphs])

            # Check for required Polish section headers
            required_sections = [
                "Podsumowanie",
                "Tabela urządzeń i nastaw",
                "Sprawdzenie czułości",
                "Sprawdzenie selektywności",
                "Sprawdzenie przeciążalności",
                "Krzywe czasowo-prądowe",
            ]

            for section in required_sections:
                assert section in full_text, (
                    f"Brak sekcji '{section}' w raporcie DOCX"
                )

    def test_no_codenames_in_docx(self) -> None:
        """Raport DOCX nie zawiera nazw kodowych projektu."""
        from docx import Document as DocxDocument

        result = _create_deterministic_protection_result()

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "codename_test.docx"
            export_protection_coordination_to_docx(result, path, deterministic=True)

            doc = DocxDocument(str(path))
            full_text = "\n".join([p.text for p in doc.paragraphs])

            # Forbidden codenames
            forbidden = ["P7", "P11", "P14", "P17", "P20", "FIX-12"]
            for codename in forbidden:
                assert codename not in full_text, (
                    f"Znaleziono niedozwoloną nazwę kodową '{codename}' w raporcie"
                )


@pytest.mark.skipif(not _PDF_AVAILABLE, reason="reportlab not installed")
class TestProtectionPDFContentValidation:
    """E2E: Walidacja zawartości PDF."""

    def test_pdf_file_valid_and_nonzero(self) -> None:
        """PDF jest poprawnym plikiem o niezerowym rozmiarze."""
        result = _create_deterministic_protection_result()

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "valid_test.pdf"
            export_protection_coordination_to_pdf(result, path)

            # Check file exists and has content
            assert path.exists(), "Plik PDF nie został utworzony"
            assert path.stat().st_size > 0, "Plik PDF jest pusty"

            # Check PDF magic bytes
            with open(path, "rb") as f:
                magic = f.read(5)
            assert magic == b"%PDF-", "Plik nie jest poprawnym PDF-em"


# =============================================================================
# E2E Integration Test - Full Workflow
# =============================================================================


class TestProtectionExportFullWorkflow:
    """E2E: Pełny workflow eksportu obu formatów."""

    @pytest.mark.skipif(
        not (_PDF_AVAILABLE and _DOCX_AVAILABLE),
        reason="reportlab or python-docx not installed"
    )
    def test_both_exports_from_same_result(self) -> None:
        """PDF i DOCX z tego samego wyniku są deterministyczne."""
        result = _create_deterministic_protection_result()

        with tempfile.TemporaryDirectory() as tmpdir:
            # Export PDF twice
            pdf_1 = Path(tmpdir) / "report_1.pdf"
            pdf_2 = Path(tmpdir) / "report_2.pdf"
            export_protection_coordination_to_pdf(result, pdf_1)
            export_protection_coordination_to_pdf(result, pdf_2)

            # Export DOCX twice
            docx_1 = Path(tmpdir) / "report_1.docx"
            docx_2 = Path(tmpdir) / "report_2.docx"
            export_protection_coordination_to_docx(result, docx_1, deterministic=True)
            export_protection_coordination_to_docx(result, docx_2, deterministic=True)

            # Verify all exports are deterministic
            assert _compute_file_hash(pdf_1) == _compute_file_hash(pdf_2), (
                "PDF export nie deterministyczny"
            )
            assert _compute_file_hash(docx_1) == _compute_file_hash(docx_2), (
                "DOCX export nie deterministyczny"
            )

            # Verify files are different formats (not accidentally same)
            assert _compute_file_hash(pdf_1) != _compute_file_hash(docx_1), (
                "PDF i DOCX mają ten sam hash - to niemożliwe"
            )

    @pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx not installed")
    def test_sorted_order_in_checks(self) -> None:
        """Sprawdzenia są sortowane deterministycznie po device_id."""
        from docx import Document as DocxDocument

        # Create result with unsorted checks
        result = _create_deterministic_protection_result()

        # Swap order of sensitivity checks (B before A)
        result["sensitivity_checks"] = [
            result["sensitivity_checks"][1],  # DEV-B-002
            result["sensitivity_checks"][0],  # DEV-A-001
        ]

        with tempfile.TemporaryDirectory() as tmpdir:
            path_1 = Path(tmpdir) / "sorted_1.docx"
            path_2 = Path(tmpdir) / "sorted_2.docx"

            export_protection_coordination_to_docx(result, path_1, deterministic=True)
            export_protection_coordination_to_docx(result, path_2, deterministic=True)

            # Should still be identical due to internal sorting
            hash_1 = _compute_file_hash(path_1)
            hash_2 = _compute_file_hash(path_2)

            assert hash_1 == hash_2, (
                "Eksport z różną kolejnością wejściową daje różne wyniki"
            )
