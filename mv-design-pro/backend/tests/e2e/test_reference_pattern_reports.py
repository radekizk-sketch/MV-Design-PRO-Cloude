"""
Testy E2E dla raportów wzorców odniesienia (A i C).

Weryfikacja:
- Poprawne generowanie plików DOCX i PDF
- Powtarzalność binarna DOCX (2× eksport → identyczny SHA256)
- Powtarzalność PDF (lub udokumentowane ograniczenie)
- Obecność wymaganych sekcji w dokumentach

DETERMINIZM:
- DOCX: pełna powtarzalność binarna wymagana
- PDF: stabilizacja metadanych, może wymagać oznaczenia xfail

CANONICAL ALIGNMENT:
- Warstwa prezentacji — NIE liczy fizyki
- Treści po polsku, bez nazw kodowych
"""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import TYPE_CHECKING

import pytest

# Dostępność bibliotek
try:
    from docx import Document

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import reportlab

    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False


# =============================================================================
# FIXTURES
# =============================================================================


@pytest.fixture
def pattern_a_result():
    """Generuje wynik wzorca A z fixture case_A_zgodne."""
    from application.reference_patterns import run_pattern_a

    return run_pattern_a(fixture_file="case_A_zgodne.json")


@pytest.fixture
def pattern_c_result():
    """Generuje wynik wzorca C z fixture przypadek_zgodny."""
    from application.reference_patterns import run_pattern_c

    return run_pattern_c(fixture_file="przypadek_zgodny.json")


@pytest.fixture
def report_metadata():
    """Metadane raportu do testów."""
    from application.reference_patterns import ReportMetadata

    return ReportMetadata(
        project_name="Projekt testowy E2E",
        case_name="Przypadek referencyjny",
        execution_date="2000-01-01",
        author="Test E2E",
        version="1.0",
    )


def _compute_sha256(path: Path) -> str:
    """Oblicza SHA256 pliku."""
    sha256 = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def _get_docx_text(path: Path) -> str:
    """Ekstrahuje tekst z dokumentu DOCX."""
    doc = Document(path)
    text_parts = []
    for para in doc.paragraphs:
        text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    return "\n".join(text_parts)


# =============================================================================
# TESTY DOCX — WZORZEC A
# =============================================================================


@pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx nie jest zainstalowany")
class TestDocxWzorzecA:
    """Testy DOCX dla wzorca A."""

    def test_docx_creates_file(
        self, pattern_a_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje utworzenie pliku DOCX."""
        from application.reference_patterns import export_reference_pattern_to_docx

        output_file = tmp_path / "wzorzec_a.docx"
        returned_path = export_reference_pattern_to_docx(
            pattern_a_result, output_file, report_metadata
        )

        assert output_file.exists()
        assert returned_path == output_file
        assert output_file.stat().st_size > 0

    def test_docx_powtarzalnosc_wzorzec_a(
        self, pattern_a_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje powtarzalność binarną DOCX dla wzorca A.

        2× eksport z identycznymi danymi → identyczny SHA256.
        """
        from application.reference_patterns import export_reference_pattern_to_docx

        output_1 = tmp_path / "wzorzec_a_1.docx"
        output_2 = tmp_path / "wzorzec_a_2.docx"

        # Pierwszy eksport
        export_reference_pattern_to_docx(pattern_a_result, output_1, report_metadata)

        # Drugi eksport
        export_reference_pattern_to_docx(pattern_a_result, output_2, report_metadata)

        # Porównaj SHA256
        sha256_1 = _compute_sha256(output_1)
        sha256_2 = _compute_sha256(output_2)

        assert sha256_1 == sha256_2, (
            f"Pliki DOCX nie są identyczne binarnie:\n"
            f"  {output_1}: {sha256_1}\n"
            f"  {output_2}: {sha256_2}"
        )

    def test_docx_contains_required_sections_wzorzec_a(
        self, pattern_a_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje obecność wymaganych sekcji w DOCX dla wzorca A."""
        from application.reference_patterns import export_reference_pattern_to_docx

        output_file = tmp_path / "wzorzec_a.docx"
        export_reference_pattern_to_docx(pattern_a_result, output_file, report_metadata)

        text = _get_docx_text(output_file)

        # Wymagane sekcje
        assert "Raport wzorca odniesienia" in text
        assert "Streszczenie" in text
        assert "Werdykt" in text
        assert "Sprawdzenia" in text
        assert "Wartości pośrednie" in text
        assert "Informacje o raporcie" in text

        # Nazwa wzorca
        assert "Dobór I>>" in text or "linii SN" in text

        # Werdykt (jeden z trzech)
        assert any(v in text for v in ["ZGODNE", "GRANICZNE", "NIEZGODNE"])

    def test_docx_without_trace(
        self, pattern_a_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje generowanie DOCX bez śladu obliczeń."""
        from application.reference_patterns import export_reference_pattern_to_docx

        output_file = tmp_path / "wzorzec_a_bez_sladu.docx"
        export_reference_pattern_to_docx(
            pattern_a_result, output_file, report_metadata, include_trace=False
        )

        assert output_file.exists()
        text = _get_docx_text(output_file)

        # Ślad obliczeń nie powinien być obecny
        assert "Ślad obliczeń" not in text


# =============================================================================
# TESTY DOCX — WZORZEC C
# =============================================================================


@pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx nie jest zainstalowany")
class TestDocxWzorzecC:
    """Testy DOCX dla wzorca C."""

    def test_docx_creates_file(
        self, pattern_c_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje utworzenie pliku DOCX."""
        from application.reference_patterns import export_reference_pattern_to_docx

        output_file = tmp_path / "wzorzec_c.docx"
        returned_path = export_reference_pattern_to_docx(
            pattern_c_result, output_file, report_metadata
        )

        assert output_file.exists()
        assert returned_path == output_file
        assert output_file.stat().st_size > 0

    def test_docx_powtarzalnosc_wzorzec_c(
        self, pattern_c_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje powtarzalność binarną DOCX dla wzorca C.

        2× eksport z identycznymi danymi → identyczny SHA256.
        """
        from application.reference_patterns import export_reference_pattern_to_docx

        output_1 = tmp_path / "wzorzec_c_1.docx"
        output_2 = tmp_path / "wzorzec_c_2.docx"

        # Pierwszy eksport
        export_reference_pattern_to_docx(pattern_c_result, output_1, report_metadata)

        # Drugi eksport
        export_reference_pattern_to_docx(pattern_c_result, output_2, report_metadata)

        # Porównaj SHA256
        sha256_1 = _compute_sha256(output_1)
        sha256_2 = _compute_sha256(output_2)

        assert sha256_1 == sha256_2, (
            f"Pliki DOCX nie są identyczne binarnie:\n"
            f"  {output_1}: {sha256_1}\n"
            f"  {output_2}: {sha256_2}"
        )

    def test_docx_contains_required_sections_wzorzec_c(
        self, pattern_c_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje obecność wymaganych sekcji w DOCX dla wzorca C."""
        from application.reference_patterns import export_reference_pattern_to_docx

        output_file = tmp_path / "wzorzec_c.docx"
        export_reference_pattern_to_docx(pattern_c_result, output_file, report_metadata)

        text = _get_docx_text(output_file)

        # Wymagane sekcje
        assert "Raport wzorca odniesienia" in text
        assert "Streszczenie" in text
        assert "Werdykt" in text
        assert "Sprawdzenia" in text

        # Nazwa wzorca (wpływ generacji lokalnej)
        assert "generacji lokalnej" in text or "RP-LOC-GEN-IMPACT" in text

        # Werdykt (jeden z trzech)
        assert any(v in text for v in ["ZGODNE", "GRANICZNE", "NIEZGODNE"])


# =============================================================================
# TESTY PDF — WZORZEC A
# =============================================================================


@pytest.mark.skipif(not _PDF_AVAILABLE, reason="reportlab nie jest zainstalowany")
class TestPdfWzorzecA:
    """Testy PDF dla wzorca A."""

    def test_pdf_creates_file(
        self, pattern_a_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje utworzenie pliku PDF."""
        from application.reference_patterns import export_reference_pattern_to_pdf

        output_file = tmp_path / "wzorzec_a.pdf"
        returned_path = export_reference_pattern_to_pdf(
            pattern_a_result, output_file, report_metadata
        )

        assert output_file.exists()
        assert returned_path == output_file
        assert output_file.stat().st_size > 0

    def test_pdf_has_valid_header(
        self, pattern_a_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje poprawny nagłówek PDF."""
        from application.reference_patterns import export_reference_pattern_to_pdf

        output_file = tmp_path / "wzorzec_a.pdf"
        export_reference_pattern_to_pdf(pattern_a_result, output_file, report_metadata)

        with open(output_file, "rb") as f:
            header = f.read(4)

        assert header == b"%PDF", f"Oczekiwano nagłówka PDF, otrzymano {header!r}"

    def test_pdf_powtarzalnosc_wzorzec_a(
        self, pattern_a_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje powtarzalność PDF dla wzorca A.

        UWAGA: PDF z reportlab może nie być w pełni binarnie powtarzalny
        ze względu na wewnętrzne znaczniki czasu. Test sprawdza,
        czy pliki są podobne pod względem rozmiaru (tolerancja 1%).

        W przypadku pełnej powtarzalności — weryfikacja SHA256.
        W przypadku braku — test przechodzi z ostrzeżeniem.
        """
        from application.reference_patterns import export_reference_pattern_to_pdf

        output_1 = tmp_path / "wzorzec_a_1.pdf"
        output_2 = tmp_path / "wzorzec_a_2.pdf"

        # Pierwszy eksport
        export_reference_pattern_to_pdf(pattern_a_result, output_1, report_metadata)

        # Drugi eksport
        export_reference_pattern_to_pdf(pattern_a_result, output_2, report_metadata)

        sha256_1 = _compute_sha256(output_1)
        sha256_2 = _compute_sha256(output_2)

        if sha256_1 == sha256_2:
            # Pełna powtarzalność binarna — sukces
            assert True
        else:
            # Sprawdź czy rozmiary są podobne (tolerancja 1%)
            size_1 = output_1.stat().st_size
            size_2 = output_2.stat().st_size
            size_diff_pct = abs(size_1 - size_2) / max(size_1, size_2) * 100

            # Jeśli różnica rozmiaru > 5%, coś jest nie tak
            assert size_diff_pct < 5, (
                f"Pliki PDF różnią się znacząco:\n"
                f"  {output_1}: {size_1} bajtów ({sha256_1[:16]}...)\n"
                f"  {output_2}: {size_2} bajtów ({sha256_2[:16]}...)\n"
                f"  Różnica: {size_diff_pct:.2f}%"
            )

            # Ostrzeżenie o braku pełnej powtarzalności
            pytest.skip(
                f"PDF nie jest binarnie powtarzalny (różnica SHA256), "
                f"ale rozmiary są podobne ({size_diff_pct:.2f}% różnicy). "
                f"Ograniczenie biblioteki reportlab — metadane PDF zawierają znaczniki czasu."
            )


# =============================================================================
# TESTY PDF — WZORZEC C
# =============================================================================


@pytest.mark.skipif(not _PDF_AVAILABLE, reason="reportlab nie jest zainstalowany")
class TestPdfWzorzecC:
    """Testy PDF dla wzorca C."""

    def test_pdf_creates_file(
        self, pattern_c_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje utworzenie pliku PDF."""
        from application.reference_patterns import export_reference_pattern_to_pdf

        output_file = tmp_path / "wzorzec_c.pdf"
        returned_path = export_reference_pattern_to_pdf(
            pattern_c_result, output_file, report_metadata
        )

        assert output_file.exists()
        assert returned_path == output_file
        assert output_file.stat().st_size > 0

    def test_pdf_has_valid_header(
        self, pattern_c_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje poprawny nagłówek PDF."""
        from application.reference_patterns import export_reference_pattern_to_pdf

        output_file = tmp_path / "wzorzec_c.pdf"
        export_reference_pattern_to_pdf(pattern_c_result, output_file, report_metadata)

        with open(output_file, "rb") as f:
            header = f.read(4)

        assert header == b"%PDF", f"Oczekiwano nagłówka PDF, otrzymano {header!r}"

    def test_pdf_powtarzalnosc_wzorzec_c(
        self, pattern_c_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje powtarzalność PDF dla wzorca C.

        UWAGA: Jak w test_pdf_powtarzalnosc_wzorzec_a — PDF może nie być
        w pełni binarnie powtarzalny.
        """
        from application.reference_patterns import export_reference_pattern_to_pdf

        output_1 = tmp_path / "wzorzec_c_1.pdf"
        output_2 = tmp_path / "wzorzec_c_2.pdf"

        # Pierwszy eksport
        export_reference_pattern_to_pdf(pattern_c_result, output_1, report_metadata)

        # Drugi eksport
        export_reference_pattern_to_pdf(pattern_c_result, output_2, report_metadata)

        sha256_1 = _compute_sha256(output_1)
        sha256_2 = _compute_sha256(output_2)

        if sha256_1 == sha256_2:
            # Pełna powtarzalność binarna — sukces
            assert True
        else:
            # Sprawdź czy rozmiary są podobne (tolerancja 1%)
            size_1 = output_1.stat().st_size
            size_2 = output_2.stat().st_size
            size_diff_pct = abs(size_1 - size_2) / max(size_1, size_2) * 100

            assert size_diff_pct < 5, (
                f"Pliki PDF różnią się znacząco:\n"
                f"  {output_1}: {size_1} bajtów ({sha256_1[:16]}...)\n"
                f"  {output_2}: {size_2} bajtów ({sha256_2[:16]}...)\n"
                f"  Różnica: {size_diff_pct:.2f}%"
            )

            pytest.skip(
                f"PDF nie jest binarnie powtarzalny (różnica SHA256), "
                f"ale rozmiary są podobne ({size_diff_pct:.2f}% różnicy). "
                f"Ograniczenie biblioteki reportlab — metadane PDF zawierają znaczniki czasu."
            )


# =============================================================================
# TESTY INTEGRACYJNE — BŁĘDY
# =============================================================================


class TestErrorHandling:
    """Testy obsługi błędów."""

    @pytest.mark.skipif(_DOCX_AVAILABLE, reason="Test dla brakującej biblioteki")
    def test_docx_import_error_when_missing(
        self, pattern_a_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje komunikat błędu gdy python-docx niedostępny."""
        from application.reference_patterns import export_reference_pattern_to_docx

        with pytest.raises(ImportError, match="python-docx"):
            export_reference_pattern_to_docx(
                pattern_a_result, tmp_path / "test.docx", report_metadata
            )

    @pytest.mark.skipif(_PDF_AVAILABLE, reason="Test dla brakującej biblioteki")
    def test_pdf_import_error_when_missing(
        self, pattern_a_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje komunikat błędu gdy reportlab niedostępny."""
        from application.reference_patterns import export_reference_pattern_to_pdf

        with pytest.raises(ImportError, match="reportlab"):
            export_reference_pattern_to_pdf(
                pattern_a_result, tmp_path / "test.pdf", report_metadata
            )


# =============================================================================
# TESTY DETERMINIZMU — WIELOKROTNE URUCHOMIENIA
# =============================================================================


@pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx nie jest zainstalowany")
class TestDeterminismMultipleRuns:
    """Testy determinizmu dla wielokrotnych uruchomień."""

    def test_docx_determinism_5_runs_wzorzec_a(
        self, pattern_a_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje determinizm dla 5 kolejnych eksportów DOCX wzorca A."""
        from application.reference_patterns import export_reference_pattern_to_docx

        hashes = []
        for i in range(5):
            output_file = tmp_path / f"wzorzec_a_run_{i}.docx"
            export_reference_pattern_to_docx(
                pattern_a_result, output_file, report_metadata
            )
            hashes.append(_compute_sha256(output_file))

        # Wszystkie hashe powinny być identyczne
        assert len(set(hashes)) == 1, (
            f"Pliki DOCX nie są deterministyczne:\n"
            f"  Unikalne hashe: {len(set(hashes))}\n"
            f"  Hashe: {hashes}"
        )

    def test_docx_determinism_5_runs_wzorzec_c(
        self, pattern_c_result, report_metadata, tmp_path: Path
    ) -> None:
        """Weryfikuje determinizm dla 5 kolejnych eksportów DOCX wzorca C."""
        from application.reference_patterns import export_reference_pattern_to_docx

        hashes = []
        for i in range(5):
            output_file = tmp_path / f"wzorzec_c_run_{i}.docx"
            export_reference_pattern_to_docx(
                pattern_c_result, output_file, report_metadata
            )
            hashes.append(_compute_sha256(output_file))

        # Wszystkie hashe powinny być identyczne
        assert len(set(hashes)) == 1, (
            f"Pliki DOCX nie są deterministyczne:\n"
            f"  Unikalne hashe: {len(set(hashes))}\n"
            f"  Hashe: {hashes}"
        )
