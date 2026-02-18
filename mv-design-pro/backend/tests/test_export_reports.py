"""
Tests for the professional export report layer.

Verifies:
- test_sc_docx_generation: SC DOCX report created, non-empty, openable, has sections
- test_pf_docx_generation: PF DOCX report created, non-empty, openable, has sections
- test_sc_pdf_generation: SC PDF report created, non-empty, starts with %PDF
- test_pf_pdf_generation: PF PDF report created, non-empty, starts with %PDF
- test_jsonl_trace_export: Trace JSONL created, correct line count, valid JSON per line
- test_export_manifest_determinism: Same inputs produce identical manifest export_id
"""

from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

import pytest

from network_model.core.branch import BranchType, LineBranch, TransformerBranch
from network_model.core.graph import NetworkGraph
from network_model.core.node import Node, NodeType
from network_model.solvers.short_circuit_iec60909 import (
    ShortCircuitIEC60909Solver,
    ShortCircuitResult,
)
from network_model.solvers.power_flow_result import (
    PowerFlowBranchResult,
    PowerFlowBusResult,
    PowerFlowResultV1,
    PowerFlowSummary,
)

# Check for optional dependencies
try:
    from docx import Document

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas  # noqa: F401

    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False


# =============================================================================
# Test Helpers
# =============================================================================

def _create_pq_node(node_id: str, voltage_level: float) -> Node:
    return Node(
        id=node_id,
        name=f"Node {node_id}",
        node_type=NodeType.PQ,
        voltage_level=voltage_level,
        active_power=5.0,
        reactive_power=2.0,
    )


def _create_reference_node(node_id: str, voltage_level: float) -> Node:
    return Node(
        id=node_id,
        name=f"Ref {node_id}",
        node_type=NodeType.PQ,
        voltage_level=voltage_level,
        active_power=0.0,
        reactive_power=0.0,
    )


def _create_transformer_branch(
    branch_id: str,
    from_node: str,
    to_node: str,
) -> TransformerBranch:
    return TransformerBranch(
        id=branch_id,
        name=f"Transformer {branch_id}",
        branch_type=BranchType.TRANSFORMER,
        from_node_id=from_node,
        to_node_id=to_node,
        in_service=True,
        rated_power_mva=25.0,
        voltage_hv_kv=110.0,
        voltage_lv_kv=20.0,
        uk_percent=10.0,
        pk_kw=120.0,
        i0_percent=0.0,
        p0_kw=0.0,
        vector_group="Dyn11",
        tap_position=0,
        tap_step_percent=2.5,
    )


def _create_reference_branch(
    branch_id: str,
    from_node: str,
    to_node: str,
) -> LineBranch:
    return LineBranch(
        id=branch_id,
        name=f"Ref {branch_id}",
        branch_type=BranchType.LINE,
        from_node_id=from_node,
        to_node_id=to_node,
        r_ohm_per_km=1e9,
        x_ohm_per_km=0.0,
        b_us_per_km=0.0,
        length_km=1.0,
        rated_current_a=0.0,
    )


def _build_test_graph() -> NetworkGraph:
    """Build a minimal test graph with a single transformer."""
    graph = NetworkGraph()
    graph.add_node(_create_pq_node("A", 110.0))
    graph.add_node(_create_pq_node("B", 20.0))
    graph.add_node(_create_reference_node("GND", 20.0))
    graph.add_branch(_create_transformer_branch("T1", "A", "B"))
    graph.add_branch(_create_reference_branch("REF", "B", "GND"))
    return graph


def _build_sc_result() -> ShortCircuitResult:
    """Compute a SC result using the real solver on a minimal graph."""
    graph = _build_test_graph()
    return ShortCircuitIEC60909Solver.compute_3ph_short_circuit(
        graph=graph,
        fault_node_id="B",
        c_factor=1.0,
        tk_s=1.0,
    )


def _build_pf_result() -> PowerFlowResultV1:
    """Build a synthetic PowerFlowResultV1 for testing."""
    bus_results = (
        PowerFlowBusResult(
            bus_id="BUS_A",
            v_pu=1.00,
            angle_deg=0.0,
            p_injected_mw=10.0,
            q_injected_mvar=3.0,
        ),
        PowerFlowBusResult(
            bus_id="BUS_B",
            v_pu=0.98,
            angle_deg=-2.5,
            p_injected_mw=-8.0,
            q_injected_mvar=-2.5,
        ),
        PowerFlowBusResult(
            bus_id="BUS_C",
            v_pu=0.93,
            angle_deg=-5.0,
            p_injected_mw=-1.5,
            q_injected_mvar=-0.3,
        ),
    )
    branch_results = (
        PowerFlowBranchResult(
            branch_id="BR_1",
            p_from_mw=5.0,
            q_from_mvar=1.5,
            p_to_mw=-4.8,
            q_to_mvar=-1.4,
            losses_p_mw=0.2,
            losses_q_mvar=0.1,
        ),
        PowerFlowBranchResult(
            branch_id="BR_2",
            p_from_mw=3.0,
            q_from_mvar=1.0,
            p_to_mw=-2.9,
            q_to_mvar=-0.9,
            losses_p_mw=0.1,
            losses_q_mvar=0.1,
        ),
    )
    summary = PowerFlowSummary(
        total_losses_p_mw=0.3,
        total_losses_q_mvar=0.2,
        min_v_pu=0.93,
        max_v_pu=1.00,
        slack_p_mw=10.0,
        slack_q_mvar=3.0,
    )
    return PowerFlowResultV1(
        result_version="1.0.0",
        converged=True,
        iterations_count=4,
        tolerance_used=1e-6,
        base_mva=100.0,
        slack_bus_id="BUS_A",
        bus_results=bus_results,
        branch_results=branch_results,
        summary=summary,
    )


def _build_test_snapshot() -> dict:
    """Build a minimal test snapshot dict."""
    return {
        "metadata": {
            "snapshot_id": "snap_001",
            "created_at": "2025-01-15T10:30:00Z",
        },
        "nodes": [
            {"id": "A", "node_type": "PQ", "voltage_level": 110.0},
            {"id": "B", "node_type": "PQ", "voltage_level": 20.0},
            {"id": "GND", "node_type": "PQ", "voltage_level": 20.0},
        ],
        "branches": [
            {
                "id": "T1",
                "branch_type": "TRANSFORMER",
                "from_node_id": "A",
                "to_node_id": "B",
                "uk_percent": 10.0,
            },
            {
                "id": "REF",
                "branch_type": "LINE",
                "from_node_id": "B",
                "to_node_id": "GND",
                "z_ohm": 1e9,
            },
        ],
    }


def _build_test_white_box_trace() -> list[dict]:
    """Build a minimal white-box trace for testing JSONL export."""
    return [
        {
            "key": "Zk",
            "title": "Impedancja zastpcza",
            "formula_latex": "Z_k = Z_1",
            "inputs": {
                "z1_ohm": complex(0.5, 1.2),
                "fault_node_id": "B",
            },
            "substitution": "0.5+j1.2",
            "result": {"z_equiv_ohm": complex(0.5, 1.2)},
            "notes": None,
        },
        {
            "key": "Ikss",
            "title": "Prad poczatkowy",
            "formula_latex": "I_{k}'' = (c * U_n) / |Z_k|",
            "inputs": {
                "c_factor": 1.0,
                "un_v": 20000.0,
                "z_equiv_abs_ohm": 1.3,
            },
            "substitution": "(1.0 * 20000) / 1.3",
            "result": {"ikss_a": 15384.6},
            "notes": None,
        },
        {
            "key": "kappa",
            "title": "Wspolczynnik udaru",
            "formula_latex": "kappa = 1.02 + 0.98 * exp(-3*R/X)",
            "inputs": {"r_ohm": 0.5, "x_ohm": 1.2, "rx_ratio": 0.417},
            "substitution": "1.02 + 0.98 * exp(-3 * 0.417)",
            "result": {"kappa": 1.296},
            "notes": None,
        },
    ]


def _get_docx_text(doc: Document) -> str:
    """Extract all text from a DOCX Document."""
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# =============================================================================
# SC DOCX Tests
# =============================================================================

@pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx not installed")
class TestSCDocxGeneration:
    """test_sc_docx_generation: verify file created, non-empty, correct content."""

    def test_creates_file(self, tmp_path: Path) -> None:
        """DOCX file is created at specified path."""
        from network_model.reporting.export_docx import generate_sc_report_docx

        result = _build_sc_result()
        snapshot = _build_test_snapshot()
        out = tmp_path / "sc_report.docx"

        returned = generate_sc_report_docx(result, snapshot, out)

        assert out.exists()
        assert returned == out

    def test_file_is_non_empty(self, tmp_path: Path) -> None:
        """DOCX file has size > 0."""
        from network_model.reporting.export_docx import generate_sc_report_docx

        result = _build_sc_result()
        out = tmp_path / "sc_report.docx"
        generate_sc_report_docx(result, None, out)

        assert out.stat().st_size > 0

    def test_openable_and_has_content(self, tmp_path: Path) -> None:
        """DOCX file is openable and contains expected sections."""
        from network_model.reporting.export_docx import generate_sc_report_docx

        result = _build_sc_result()
        snapshot = _build_test_snapshot()
        out = tmp_path / "sc_report.docx"
        generate_sc_report_docx(result, snapshot, out)

        doc = Document(str(out))
        text = _get_docx_text(doc)

        # Title
        assert "Raport zwarciowy IEC 60909" in text
        # Input data section
        assert "Dane wejsciowe" in text
        # Equations
        assert "Rownania IEC 60909" in text
        # White box trace
        assert "White Box" in text
        # Results
        assert "Wyniki" in text
        assert "Ik''" in text
        # Contributions
        assert "Udzialy zrodel" in text

    def test_with_project_info(self, tmp_path: Path) -> None:
        """Custom project info appears in the report."""
        from network_model.reporting.export_docx import generate_sc_report_docx

        result = _build_sc_result()
        out = tmp_path / "sc_report.docx"
        generate_sc_report_docx(
            result,
            None,
            out,
            project_info={"name": "Test Project", "author": "Tester"},
        )

        doc = Document(str(out))
        text = _get_docx_text(doc)
        assert "Test Project" in text

    def test_creates_parent_dirs(self, tmp_path: Path) -> None:
        """Parent directories are created automatically."""
        from network_model.reporting.export_docx import generate_sc_report_docx

        result = _build_sc_result()
        out = tmp_path / "nested" / "deep" / "sc_report.docx"
        generate_sc_report_docx(result, None, out)

        assert out.exists()

    def test_without_snapshot(self, tmp_path: Path) -> None:
        """Report works when snapshot is None."""
        from network_model.reporting.export_docx import generate_sc_report_docx

        result = _build_sc_result()
        out = tmp_path / "sc_no_snap.docx"
        generate_sc_report_docx(result, None, out)

        assert out.exists()
        assert out.stat().st_size > 0


# =============================================================================
# PF DOCX Tests
# =============================================================================

@pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx not installed")
class TestPFDocxGeneration:
    """test_pf_docx_generation: verify PF DOCX report."""

    def test_creates_file(self, tmp_path: Path) -> None:
        """PF DOCX file is created at specified path."""
        from network_model.reporting.export_docx import generate_pf_report_docx

        result = _build_pf_result()
        out = tmp_path / "pf_report.docx"
        returned = generate_pf_report_docx(result, None, out)

        assert out.exists()
        assert returned == out

    def test_file_is_non_empty(self, tmp_path: Path) -> None:
        """PF DOCX file has size > 0."""
        from network_model.reporting.export_docx import generate_pf_report_docx

        result = _build_pf_result()
        out = tmp_path / "pf_report.docx"
        generate_pf_report_docx(result, None, out)

        assert out.stat().st_size > 0

    def test_openable_and_has_sections(self, tmp_path: Path) -> None:
        """PF DOCX is openable and contains expected sections."""
        from network_model.reporting.export_docx import generate_pf_report_docx

        result = _build_pf_result()
        snapshot = _build_test_snapshot()
        out = tmp_path / "pf_report.docx"
        generate_pf_report_docx(result, snapshot, out)

        doc = Document(str(out))
        text = _get_docx_text(doc)

        # Title
        assert "Raport rozplywu mocy" in text
        # Summary
        assert "Podsumowanie" in text
        # Bus voltages
        assert "Napiecia wezlowe" in text
        assert "BUS_A" in text
        # Branch flows
        assert "Przeplywy galeziowe" in text
        assert "BR_1" in text
        # Violations (BUS_C at 0.93 pu < 0.95 limit)
        assert "Naruszenia" in text
        assert "Podnapiecie" in text

    def test_convergence_ok(self, tmp_path: Path) -> None:
        """Converged result shows no convergence violation."""
        from network_model.reporting.export_docx import generate_pf_report_docx

        result = _build_pf_result()
        out = tmp_path / "pf_conv.docx"
        generate_pf_report_docx(result, None, out)

        doc = Document(str(out))
        text = _get_docx_text(doc)
        assert "NIE zbiegly" not in text


# =============================================================================
# SC PDF Tests
# =============================================================================

@pytest.mark.skipif(not _PDF_AVAILABLE, reason="reportlab not installed")
class TestSCPDFGeneration:
    """test_sc_pdf_generation: verify SC PDF report."""

    def test_creates_file(self, tmp_path: Path) -> None:
        """SC PDF file is created at specified path."""
        from network_model.reporting.export_pdf import generate_sc_report_pdf

        result = _build_sc_result()
        out = tmp_path / "sc_report.pdf"
        returned = generate_sc_report_pdf(result, None, out)

        assert out.exists()
        assert returned == out

    def test_file_is_non_empty(self, tmp_path: Path) -> None:
        """SC PDF file has size > 0."""
        from network_model.reporting.export_pdf import generate_sc_report_pdf

        result = _build_sc_result()
        out = tmp_path / "sc_report.pdf"
        generate_sc_report_pdf(result, None, out)

        assert out.stat().st_size > 0

    def test_is_valid_pdf(self, tmp_path: Path) -> None:
        """SC PDF starts with %PDF header."""
        from network_model.reporting.export_pdf import generate_sc_report_pdf

        result = _build_sc_result()
        snapshot = _build_test_snapshot()
        out = tmp_path / "sc_report.pdf"
        generate_sc_report_pdf(result, snapshot, out)

        content = out.read_bytes()
        assert content[:5] == b"%PDF-"

    def test_creates_parent_dirs(self, tmp_path: Path) -> None:
        """Parent directories are created automatically."""
        from network_model.reporting.export_pdf import generate_sc_report_pdf

        result = _build_sc_result()
        out = tmp_path / "nested" / "sc.pdf"
        generate_sc_report_pdf(result, None, out)

        assert out.exists()

    def test_with_snapshot(self, tmp_path: Path) -> None:
        """SC PDF with snapshot is larger than without."""
        from network_model.reporting.export_pdf import generate_sc_report_pdf

        result = _build_sc_result()
        snapshot = _build_test_snapshot()

        out_no_snap = tmp_path / "no_snap.pdf"
        generate_sc_report_pdf(result, None, out_no_snap)

        out_snap = tmp_path / "with_snap.pdf"
        generate_sc_report_pdf(result, snapshot, out_snap)

        # With snapshot should generally be larger (more content)
        assert out_snap.stat().st_size > 0
        assert out_no_snap.stat().st_size > 0


# =============================================================================
# PF PDF Tests
# =============================================================================

@pytest.mark.skipif(not _PDF_AVAILABLE, reason="reportlab not installed")
class TestPFPDFGeneration:
    """test_pf_pdf_generation: verify PF PDF report."""

    def test_creates_file(self, tmp_path: Path) -> None:
        """PF PDF file is created."""
        from network_model.reporting.export_pdf import generate_pf_report_pdf

        result = _build_pf_result()
        out = tmp_path / "pf_report.pdf"
        returned = generate_pf_report_pdf(result, None, out)

        assert out.exists()
        assert returned == out

    def test_file_is_non_empty(self, tmp_path: Path) -> None:
        """PF PDF file has size > 0."""
        from network_model.reporting.export_pdf import generate_pf_report_pdf

        result = _build_pf_result()
        out = tmp_path / "pf_report.pdf"
        generate_pf_report_pdf(result, None, out)

        assert out.stat().st_size > 0

    def test_is_valid_pdf(self, tmp_path: Path) -> None:
        """PF PDF starts with %PDF header."""
        from network_model.reporting.export_pdf import generate_pf_report_pdf

        result = _build_pf_result()
        out = tmp_path / "pf_report.pdf"
        generate_pf_report_pdf(result, None, out)

        content = out.read_bytes()
        assert content[:5] == b"%PDF-"


# =============================================================================
# JSONL Trace Export Tests
# =============================================================================

class TestJSONLTraceExport:
    """test_jsonl_trace_export: verify JSONL trace export."""

    def test_creates_file(self, tmp_path: Path) -> None:
        """JSONL file is created at specified path."""
        from network_model.reporting.export_jsonl import export_trace_jsonl

        trace = _build_test_white_box_trace()
        out = tmp_path / "trace.jsonl"
        returned = export_trace_jsonl(trace, out)

        assert out.exists()
        assert returned == out

    def test_correct_line_count(self, tmp_path: Path) -> None:
        """JSONL has one line per trace step."""
        from network_model.reporting.export_jsonl import export_trace_jsonl

        trace = _build_test_white_box_trace()
        out = tmp_path / "trace.jsonl"
        export_trace_jsonl(trace, out)

        content = out.read_text(encoding="utf-8").strip()
        lines = content.split("\n")
        assert len(lines) == len(trace)

    def test_each_line_is_valid_json(self, tmp_path: Path) -> None:
        """Each JSONL line parses as valid JSON."""
        from network_model.reporting.export_jsonl import export_trace_jsonl

        trace = _build_test_white_box_trace()
        out = tmp_path / "trace.jsonl"
        export_trace_jsonl(trace, out)

        content = out.read_text(encoding="utf-8").strip()
        for line_num, line in enumerate(content.split("\n"), 1):
            parsed = json.loads(line)
            assert isinstance(parsed, dict), f"Line {line_num} is not a dict"

    def test_step_fields_present(self, tmp_path: Path) -> None:
        """Each line contains required fields."""
        from network_model.reporting.export_jsonl import export_trace_jsonl

        trace = _build_test_white_box_trace()
        out = tmp_path / "trace.jsonl"
        export_trace_jsonl(trace, out)

        content = out.read_text(encoding="utf-8").strip()
        required_fields = {"step_number", "step_name", "formula", "values", "result", "unit"}

        for line in content.split("\n"):
            parsed = json.loads(line)
            assert required_fields.issubset(set(parsed.keys()))

    def test_chronological_order(self, tmp_path: Path) -> None:
        """Steps are numbered sequentially starting from 1."""
        from network_model.reporting.export_jsonl import export_trace_jsonl

        trace = _build_test_white_box_trace()
        out = tmp_path / "trace.jsonl"
        export_trace_jsonl(trace, out)

        content = out.read_text(encoding="utf-8").strip()
        step_numbers = []
        for line in content.split("\n"):
            parsed = json.loads(line)
            step_numbers.append(parsed["step_number"])

        assert step_numbers == [1, 2, 3]

    def test_step_names_match(self, tmp_path: Path) -> None:
        """Step names match the keys from the trace."""
        from network_model.reporting.export_jsonl import export_trace_jsonl

        trace = _build_test_white_box_trace()
        out = tmp_path / "trace.jsonl"
        export_trace_jsonl(trace, out)

        content = out.read_text(encoding="utf-8").strip()
        names = []
        for line in content.split("\n"):
            parsed = json.loads(line)
            names.append(parsed["step_name"])

        assert names == ["Zk", "Ikss", "kappa"]

    def test_complex_values_serialized(self, tmp_path: Path) -> None:
        """Complex values are serialized as {re, im} dicts."""
        from network_model.reporting.export_jsonl import export_trace_jsonl

        trace = _build_test_white_box_trace()
        out = tmp_path / "trace.jsonl"
        export_trace_jsonl(trace, out)

        content = out.read_text(encoding="utf-8").strip()
        first_line = json.loads(content.split("\n")[0])
        z1 = first_line["values"]["z1_ohm"]
        assert "re" in z1
        assert "im" in z1

    def test_empty_trace(self, tmp_path: Path) -> None:
        """Empty trace produces empty file."""
        from network_model.reporting.export_jsonl import export_trace_jsonl

        out = tmp_path / "empty.jsonl"
        export_trace_jsonl([], out)

        assert out.exists()
        assert out.read_text(encoding="utf-8") == ""

    def test_snapshot_jsonl(self, tmp_path: Path) -> None:
        """Snapshot JSONL export creates valid output."""
        from network_model.reporting.export_jsonl import export_snapshot_jsonl

        snapshot = _build_test_snapshot()
        out = tmp_path / "snapshot.jsonl"
        returned = export_snapshot_jsonl(snapshot, out)

        assert out.exists()
        assert returned == out

        content = out.read_text(encoding="utf-8").strip()
        lines = content.split("\n")

        # 1 metadata + 3 nodes + 2 branches = 6 lines
        assert len(lines) == 6

        # First line is metadata
        meta = json.loads(lines[0])
        assert meta["type"] == "metadata"
        assert meta["node_count"] == 3
        assert meta["branch_count"] == 2

        # Verify all lines are valid JSON
        for line in lines:
            parsed = json.loads(line)
            assert "type" in parsed

    def test_from_real_solver_trace(self, tmp_path: Path) -> None:
        """Export works with real solver white_box_trace."""
        from network_model.reporting.export_jsonl import export_trace_jsonl

        sc_result = _build_sc_result()
        trace = sc_result.to_dict()["white_box_trace"]
        out = tmp_path / "real_trace.jsonl"
        export_trace_jsonl(trace, out)

        assert out.exists()
        assert out.stat().st_size > 0

        content = out.read_text(encoding="utf-8").strip()
        lines = content.split("\n")
        assert len(lines) == len(trace)

        for line in lines:
            parsed = json.loads(line)
            assert "step_number" in parsed


# =============================================================================
# Export Manifest Determinism Tests
# =============================================================================

class TestExportManifestDeterminism:
    """test_export_manifest_determinism: same inputs produce identical export_id."""

    def test_manifest_from_files(self, tmp_path: Path) -> None:
        """Manifest is created from files with correct attributes."""
        from network_model.reporting.export_manifest import build_export_manifest

        # Create test files
        f1 = tmp_path / "report.docx"
        f1.write_text("test docx content")
        f2 = tmp_path / "trace.jsonl"
        f2.write_text('{"step": 1}\n')

        manifest = build_export_manifest(
            files=[f1, f2],
            snapshot_hash="abc123",
            run_id="run_001",
        )

        assert manifest.export_id  # non-empty
        assert len(manifest.export_id) == 64  # SHA-256 hex
        assert manifest.snapshot_hash == "abc123"
        assert manifest.run_id == "run_001"
        assert len(manifest.files) == 2
        assert manifest.created_at is not None

    def test_deterministic_export_id(self, tmp_path: Path) -> None:
        """Same files produce identical export_id."""
        from network_model.reporting.export_manifest import build_export_manifest

        f1 = tmp_path / "report.pdf"
        f1.write_bytes(b"pdf content here")
        f2 = tmp_path / "data.jsonl"
        f2.write_text('{"test": true}\n')

        ts = datetime(2025, 1, 15, 12, 0, 0, tzinfo=timezone.utc)

        manifest_a = build_export_manifest(
            files=[f1, f2],
            snapshot_hash="hash_abc",
            run_id="run_42",
            created_at=ts,
        )

        manifest_b = build_export_manifest(
            files=[f1, f2],
            snapshot_hash="hash_abc",
            run_id="run_42",
            created_at=ts,
        )

        assert manifest_a.export_id == manifest_b.export_id
        assert manifest_a.to_json() == manifest_b.to_json()

    def test_different_files_different_id(self, tmp_path: Path) -> None:
        """Different file contents produce different export_id."""
        from network_model.reporting.export_manifest import build_export_manifest

        f1 = tmp_path / "a.txt"
        f1.write_text("content A")

        f2 = tmp_path / "b.txt"
        f2.write_text("content B")

        ts = datetime(2025, 1, 15, 12, 0, 0, tzinfo=timezone.utc)

        manifest_a = build_export_manifest(
            files=[f1],
            snapshot_hash="same",
            run_id="same",
            created_at=ts,
        )
        manifest_b = build_export_manifest(
            files=[f2],
            snapshot_hash="same",
            run_id="same",
            created_at=ts,
        )

        assert manifest_a.export_id != manifest_b.export_id

    def test_file_format_inferred(self, tmp_path: Path) -> None:
        """File formats are correctly inferred from extensions."""
        from network_model.reporting.export_manifest import build_export_manifest

        docx_f = tmp_path / "report.docx"
        docx_f.write_text("docx")
        pdf_f = tmp_path / "report.pdf"
        pdf_f.write_text("pdf")
        jsonl_f = tmp_path / "data.jsonl"
        jsonl_f.write_text("jsonl")

        manifest = build_export_manifest(files=[docx_f, pdf_f, jsonl_f])

        formats = {f.format for f in manifest.files}
        assert formats == {"docx", "pdf", "jsonl"}

    def test_file_hashes_present(self, tmp_path: Path) -> None:
        """Each file in manifest has a non-empty SHA-256 hash."""
        from network_model.reporting.export_manifest import build_export_manifest

        f1 = tmp_path / "test.json"
        f1.write_text('{"key": "value"}')

        manifest = build_export_manifest(files=[f1])

        for ef in manifest.files:
            assert ef.hash
            assert len(ef.hash) == 64  # SHA-256 hex

    def test_frozen_dataclass(self, tmp_path: Path) -> None:
        """ExportManifest is immutable (frozen dataclass)."""
        from network_model.reporting.export_manifest import build_export_manifest

        f1 = tmp_path / "test.txt"
        f1.write_text("test")

        manifest = build_export_manifest(files=[f1])

        with pytest.raises(AttributeError):
            manifest.export_id = "modified"  # type: ignore

    def test_to_dict_serializable(self, tmp_path: Path) -> None:
        """to_dict() returns a JSON-serializable dict."""
        from network_model.reporting.export_manifest import build_export_manifest

        f1 = tmp_path / "test.pdf"
        f1.write_bytes(b"pdf bytes")

        manifest = build_export_manifest(
            files=[f1],
            snapshot_hash="test_hash",
            run_id="test_run",
        )

        d = manifest.to_dict()
        # Should be JSON-serializable
        json_str = json.dumps(d, ensure_ascii=False)
        parsed = json.loads(json_str)
        assert parsed["export_id"] == manifest.export_id
        assert parsed["snapshot_hash"] == "test_hash"
        assert parsed["run_id"] == "test_run"
        assert len(parsed["files"]) == 1

    def test_missing_file_raises(self, tmp_path: Path) -> None:
        """FileNotFoundError raised when file does not exist."""
        from network_model.reporting.export_manifest import build_export_manifest

        missing = tmp_path / "does_not_exist.txt"

        with pytest.raises(FileNotFoundError):
            build_export_manifest(files=[missing])

    def test_order_independence(self, tmp_path: Path) -> None:
        """File order does not affect export_id (files are sorted internally)."""
        from network_model.reporting.export_manifest import build_export_manifest

        f1 = tmp_path / "alpha.txt"
        f1.write_text("alpha")
        f2 = tmp_path / "beta.txt"
        f2.write_text("beta")

        ts = datetime(2025, 6, 1, 0, 0, 0, tzinfo=timezone.utc)

        m1 = build_export_manifest(
            files=[f1, f2], snapshot_hash="h", run_id="r", created_at=ts
        )
        m2 = build_export_manifest(
            files=[f2, f1], snapshot_hash="h", run_id="r", created_at=ts
        )

        assert m1.export_id == m2.export_id
